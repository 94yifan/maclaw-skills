#!/usr/bin/env python3
"""Convert liumang-report markdown to professional DOCX."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import os

SRC = os.path.expanduser("~/.openclaw/workspace/supermind/memory/liumang-report-20260718.md")
DST = os.path.expanduser("~/.openclaw/workspace/supermind/memory/liumang-report-20260718.docx")

FONT_NAME = "等线"
FONT_NAME_ASCII = "DengXian"
BODY_SIZE = Pt(10.5)
H1_SIZE = Pt(18)
H2_SIZE = Pt(14)
H3_SIZE = Pt(12)

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = BODY_SIZE
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)

# ── Heading styles ──
for level, size, bold in [(1, H1_SIZE, True), (2, H2_SIZE, True), (3, H3_SIZE, True)]:
    sname = f'Heading {level}'
    if sname in [s.name for s in doc.styles]:
        hs = doc.styles[sname]
    else:
        hs = doc.styles.add_style(sname, 1)  # WD_STYLE_TYPE.PARAGRAPH
    hs.font.name = FONT_NAME
    hs.font.size = size
    hs.font.bold = bold
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(18)
    hs.paragraph_format.space_after = Pt(10)

# ── Helpers ──
def add_cover_page():
    """Create a professional cover page."""
    # Add empty paragraphs for spacing
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("榴芒一刻 品牌竞品研究咨询报告")
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.paragraph_format.space_after = Pt(24)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("燃创咨询 | 2026-07-18")
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.paragraph_format.space_after = Pt(60)

    # Bottom line
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Supermind × 燃创 出品")
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Page break
    doc.add_page_break()

def add_styled_paragraph(text, style_name='Normal', bold=False):
    """Add a paragraph with proper font settings."""
    p = doc.add_paragraph(style=style_name)
    if not text:
        return p
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if bold:
        run.font.bold = True
    return p

def parse_inline_formatting(paragraph, text):
    """Parse bold markers (**text**) within a line and add runs."""
    # Split on bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = FONT_NAME
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run.font.size = BODY_SIZE
        else:
            # Handle inline code and other formatting
            run = paragraph.add_run(part)
            run.font.name = FONT_NAME
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run.font.size = BODY_SIZE

def add_table_from_lines(table_lines):
    """Convert markdown table lines to docx table."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            # Skip separator lines
            if re.match(r'^\|[\s\-:|]+\|$', line):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)

    if not rows:
        return

    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row formatting
    header_cells = table.rows[0].cells
    for cell in header_cells:
        # Gray background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        # Bold text
        for para in cell.paragraphs:
            para.paragraph_format.line_spacing = 1.2
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.bold = True
                run.font.name = FONT_NAME
                run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                run.font.size = Pt(9)
        if not cell.paragraphs[0].text.strip():
            run = cell.paragraphs[0].add_run(rows[0][header_cells.index(cell)])
            run.font.bold = True
            run.font.name = FONT_NAME
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run.font.size = Pt(9)

    # Data rows
    for i, row_data in enumerate(rows):
        if i == 0:
            continue
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = row.cells[j]
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.2
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after = Pt(1)
                # Clear default empty paragraph and add text
                if cell.paragraphs[0].text.strip() == '' or cell.paragraphs[0].text.strip() == cell_text:
                    pass
                if not cell.paragraphs[0].text.strip():
                    run = cell.paragraphs[0].add_run(cell_text)
                    run.font.name = FONT_NAME
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                    run.font.size = Pt(9)
                else:
                    # Already has text from header row template
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacer after table

def process_content():
    """Read markdown and convert to docx."""
    with open(SRC, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0

    in_table = False
    table_buffer = []
    in_code_block = False
    skip_until_heading1 = True  # Skip the title/frontmatter until first #

    while i < len(lines):
        line = lines[i]

        # Skip YAML frontmatter / title block before first H1
        if skip_until_heading1:
            if line.startswith('# ') and not line.startswith('## '):
                skip_until_heading1 = False
                # Add this H1
                text = line[2:].strip()
                add_styled_paragraph(text, 'Heading 1')
                i += 1
                continue
            i += 1
            continue

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            add_styled_paragraph(text, 'Heading 1')
            i += 1
            continue

        # H2
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            add_styled_paragraph(text, 'Heading 2')
            i += 1
            continue

        # H3
        if line.startswith('### '):
            text = line[4:].strip()
            add_styled_paragraph(text, 'Heading 3')
            i += 1
            continue

        # H4
        if line.startswith('#### '):
            text = line[5:].strip()
            p = add_styled_paragraph(text, 'Normal')
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(11)
            i += 1
            continue

        # Code block
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        # Table
        if line.startswith('|') and line.endswith('|') and not in_code_block:
            in_table = True
            table_buffer.append(line)
            i += 1
            continue

        if in_table:
            # End of table
            add_table_from_lines(table_buffer)
            table_buffer = []
            in_table = False
            # Don't advance i, process current line below
            continue

        # Quote block
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.italic = True
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            # Add a thin border as horizontal rule
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Bold title line (numbered items like **01 榴莲半熟芝士**)
        bold_match = re.match(r'^\*\*(.+?)\*\*\s*—(.+)$', line)
        if bold_match:
            p = doc.add_paragraph()
            run = p.add_run(bold_match.group(1).strip())
            run.font.bold = True
            run.font.name = FONT_NAME
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run.font.size = Pt(11)
            run2 = p.add_run(' — ' + bold_match.group(2).strip())
            run2.font.name = FONT_NAME
            run2.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run2.font.size = BODY_SIZE
            i += 1
            continue

        # Regular text line
        if line.strip():
            p = doc.add_paragraph()
            parse_inline_formatting(p, line)
        else:
            # Blank line
            pass

        i += 1

    # Handle trailing table
    if in_table and table_buffer:
        add_table_from_lines(table_buffer)


# ── Main ──
add_cover_page()
process_content()

# Set narrow margins for all sections
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

doc.save(DST)
print(f"✅ DOCX saved to {DST}")
print(f"   Size: {os.path.getsize(DST):,} bytes")
