#!/usr/bin/env python3
"""Generate final DOCX reports with embedded ECharts chart screenshots."""

import re, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CHARTS_DIR = os.path.join(os.path.dirname(__file__), "charts")
REPORTS_DIR = os.path.join(os.path.dirname(__file__), "reports")
CONTENT_DIR = os.path.join(os.path.dirname(__file__), "content")

os.makedirs(REPORTS_DIR, exist_ok=True)

def parse_markdown_line(line):
    parts = re.split(r'(\*\*.*?\*\*)', line)
    runs = []
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            runs.append((part[2:-2], True))
        else:
            runs.append((part, False))
    return runs

def add_formatted_paragraph(doc, text, font_size=None, bold_all=False, 
                           alignment=None, space_after=None, space_before=None, color=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    runs_data = parse_markdown_line(text)
    for run_text, is_bold in runs_data:
        if not run_text:
            continue
        run = p.add_run(run_text)
        if bold_all or is_bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
        run.font.name = '微软雅黑'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def insert_chart_image(doc, chart_name, width_inches=5.5):
    """Insert a chart PNG image into the document."""
    png_path = os.path.join(CHARTS_DIR, f"{chart_name}.png")
    if os.path.exists(png_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(png_path, width=Inches(width_inches))
        # Add caption
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"来源：天猫旗舰店，2026-07-18实时抓取")
        cap_run.font.size = Pt(8)
        cap_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        cap_run.font.name = '微软雅黑'

def convert_md_to_docx_liumang():
    """Convert 榴芒一刻 report to DOCX with charts."""
    md_path = os.path.join(CONTENT_DIR, "report_final.md")
    docx_path = os.path.join(REPORTS_DIR, "liumang_final.docx")
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle image embeds
        img_match = re.search(r'!\[.*\]\(.*charts/([^)]+\.png)\)', line)
        if img_match:
            chart_name = img_match.group(1).replace('.png', '')
            insert_chart_image(doc, chart_name)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single', qn('w:sz'): '6',
                qn('w:space'): '1', qn('w:color'): '999999',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Skip HTML comment lines like <!-- -->
        if line.strip().startswith('<!--'):
            i += 1
            continue
        
        # Handle blockquotes
        if line.startswith('>'):
            bq_lines = []
            while i < len(lines) and lines[i].rstrip().startswith('>'):
                bq_lines.append(lines[i].rstrip()[1:].strip())
                i += 1
            for bq_line in bq_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(bq_line)
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.italic = True
                run.font.name = '微软雅黑'
            continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            
            data_rows = [tl for tl in table_lines if not re.match(r'^\|[\s\-:|]+\|$', tl)]
            if len(data_rows) >= 1:
                cols = len([c for c in data_rows[0].split('|') if c.strip()])
                table = doc.add_table(rows=len(data_rows), cols=cols)
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_line in enumerate(data_rows):
                    cells = [c.strip() for c in row_line.split('|')][1:-1]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.paragraphs[0].clear()
                            runs_data = parse_markdown_line(cell_text)
                            for run_text, is_bold in runs_data:
                                if not run_text:
                                    continue
                                run = cell.paragraphs[0].add_run(run_text)
                                run.font.size = Pt(9)
                                run.font.name = '微软雅黑'
                                if is_bold or row_idx == 0:
                                    run.bold = True
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        
        # Handle headings
        if line.startswith('# '):
            add_formatted_paragraph(doc, line[2:], bold_all=True, font_size=20,
                                   space_before=18, space_after=10)
            i += 1; continue
        elif line.startswith('## '):
            add_formatted_paragraph(doc, line[3:], bold_all=True, font_size=16,
                                   space_before=16, space_after=8)
            i += 1; continue
        elif line.startswith('### '):
            add_formatted_paragraph(doc, line[4:], bold_all=True, font_size=14,
                                   space_before=14, space_after=6)
            i += 1; continue
        elif line.startswith('#### '):
            add_formatted_paragraph(doc, line[5:], bold_all=True, font_size=12,
                                   space_before=10, space_after=4)
            i += 1; continue
        
        # Handle lists
        if re.match(r'^[\-\*]\s+', line):
            list_items = []
            while i < len(lines):
                cl = lines[i].rstrip()
                if re.match(r'^[\-\*]\s+', cl):
                    list_items.append(cl)
                    i += 1
                elif cl.strip() and not cl.startswith('#') and not cl.startswith('|') and not cl.startswith('>'):
                    if list_items:
                        list_items[-1] += ' ' + cl.strip()
                    else:
                        break
                    i += 1
                else:
                    break
            
            for item in list_items:
                item_text = re.sub(r'^[\-\*]\s+', '', item)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.space_after = Pt(2)
                bullet_run = p.add_run('• ')
                bullet_run.font.size = Pt(11)
                runs_data = parse_markdown_line(item_text)
                for run_text, is_bold in runs_data:
                    if not run_text: continue
                    run = p.add_run(run_text)
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
                    if is_bold: run.bold = True
            continue
        
        if re.match(r'^\d+\.\s+', line):
            list_items = []
            while i < len(lines):
                cl = lines[i].rstrip()
                if re.match(r'^\d+\.\s+', cl):
                    list_items.append(cl)
                    i += 1
                elif cl.strip() and not cl.startswith('#') and not cl.startswith('|'):
                    if list_items:
                        list_items[-1] += ' ' + cl.strip()
                    else:
                        break
                    i += 1
                else:
                    break
            
            for idx, item in enumerate(list_items):
                item_text = re.sub(r'^\d+\.\s+', '', item)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.space_after = Pt(2)
                num_run = p.add_run(f'{idx+1}. ')
                num_run.font.size = Pt(11)
                runs_data = parse_markdown_line(item_text)
                for run_text, is_bold in runs_data:
                    if not run_text: continue
                    run = p.add_run(run_text)
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
                    if is_bold: run.bold = True
            continue
        
        if not line.strip():
            i += 1; continue
        
        # Regular paragraph - strip inline markdown
        para_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
        para_text = re.sub(r'`([^`]+)`', r'\1', para_text)
        if para_text.strip():
            add_formatted_paragraph(doc, para_text, font_size=11, space_after=6)
        i += 1
    
    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) / 1024
    print(f'榴芒一刻 DOCX saved: {docx_path}')
    print(f'  Size: {size_kb:.1f} KB')
    return docx_path

def convert_md_to_docx_beiwei47():
    """Convert 北纬47度 report to DOCX with charts."""
    md_path = os.path.join(CONTENT_DIR, "report_final_beiwei47.md")
    docx_path = os.path.join(REPORTS_DIR, "beiwei47_final.docx")
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle image embeds
        img_match = re.search(r'!\[.*\]\(.*charts/([^)]+\.png)\)', line)
        if img_match:
            chart_name = img_match.group(1).replace('.png', '')
            insert_chart_image(doc, chart_name)
            i += 1
            continue
        
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single', qn('w:sz'): '6',
                qn('w:space'): '1', qn('w:color'): '999999',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        if line.startswith('>'):
            bq_lines = []
            while i < len(lines) and lines[i].rstrip().startswith('>'):
                bq_lines.append(lines[i].rstrip()[1:].strip())
                i += 1
            for bq_line in bq_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(bq_line)
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.italic = True
                run.font.name = '微软雅黑'
            continue
        
        if '|' in line and line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            
            data_rows = [tl for tl in table_lines if not re.match(r'^\|[\s\-:|]+\|$', tl)]
            if len(data_rows) >= 1:
                cols = len([c for c in data_rows[0].split('|') if c.strip()])
                table = doc.add_table(rows=len(data_rows), cols=cols)
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_line in enumerate(data_rows):
                    cells = [c.strip() for c in row_line.split('|')][1:-1]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.paragraphs[0].clear()
                            runs_data = parse_markdown_line(cell_text)
                            for run_text, is_bold in runs_data:
                                if not run_text: continue
                                run = cell.paragraphs[0].add_run(run_text)
                                run.font.size = Pt(9)
                                run.font.name = '微软雅黑'
                                if is_bold or row_idx == 0:
                                    run.bold = True
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        
        if line.startswith('# '):
            add_formatted_paragraph(doc, line[2:], bold_all=True, font_size=20, space_before=18, space_after=10)
            i += 1; continue
        elif line.startswith('## '):
            add_formatted_paragraph(doc, line[3:], bold_all=True, font_size=16, space_before=16, space_after=8)
            i += 1; continue
        elif line.startswith('### '):
            add_formatted_paragraph(doc, line[4:], bold_all=True, font_size=14, space_before=14, space_after=6)
            i += 1; continue
        elif line.startswith('#### '):
            add_formatted_paragraph(doc, line[5:], bold_all=True, font_size=12, space_before=10, space_after=4)
            i += 1; continue
        
        if re.match(r'^[\-\*]\s+', line):
            list_items = []
            while i < len(lines):
                cl = lines[i].rstrip()
                if re.match(r'^[\-\*]\s+', cl):
                    list_items.append(cl)
                    i += 1
                elif cl.strip() and not cl.startswith('#') and not cl.startswith('|') and not cl.startswith('>'):
                    if list_items:
                        list_items[-1] += ' ' + cl.strip()
                    else:
                        break
                    i += 1
                else:
                    break
            for item in list_items:
                item_text = re.sub(r'^[\-\*]\s+', '', item)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.space_after = Pt(2)
                bullet_run = p.add_run('• ')
                bullet_run.font.size = Pt(11)
                runs_data = parse_markdown_line(item_text)
                for run_text, is_bold in runs_data:
                    if not run_text: continue
                    run = p.add_run(run_text)
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
                    if is_bold: run.bold = True
            continue
        
        if not line.strip():
            i += 1; continue
        
        para_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
        para_text = re.sub(r'`([^`]+)`', r'\1', para_text)
        if para_text.strip():
            add_formatted_paragraph(doc, para_text, font_size=11, space_after=6)
        i += 1
    
    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) / 1024
    print(f'北纬47度 DOCX saved: {docx_path}')
    print(f'  Size: {size_kb:.1f} KB')
    return docx_path

if __name__ == '__main__':
    convert_md_to_docx_liumang()
    convert_md_to_docx_beiwei47()
    print('\n✓ All DOCX reports generated with embedded charts.')
