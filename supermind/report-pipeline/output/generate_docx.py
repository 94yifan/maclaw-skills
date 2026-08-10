#!/usr/bin/env python3
"""Generate DOCX from markdown report with embedded charts."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Paths
MD_PATH = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/content/report_final.md"
CHARTS_DIR = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/charts"
OUTPUT_PATH = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/liumang_final.docx"

# Image mapping: relative path -> absolute path
IMAGE_MAP = {
    "../charts/price_range_liumang.png": os.path.join(CHARTS_DIR, "price_range_liumang.png"),
    "../charts/tmall_top_sellers_liumang.png": os.path.join(CHARTS_DIR, "tmall_top_sellers_liumang.png"),
    "../charts/store_compare_liumang.png": os.path.join(CHARTS_DIR, "store_compare_liumang.png"),
    "../charts/category_coverage_liumang.png": os.path.join(CHARTS_DIR, "category_coverage_liumang.png"),
}

# Constants
FONT_NAME = "等线"
BODY_SIZE = Pt(10.5)
HEADER_GRAY = "D9D9D9"
IMAGE_WIDTH = Inches(6)

doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = FONT_NAME
font.size = BODY_SIZE
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.5

# Set margins for all sections
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


# --- Helper functions ---
def set_run_font(run, font_name=FONT_NAME, size=BODY_SIZE, bold=False):
    """Set font properties on a run."""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph_spacer(space_after_pt=0):
    """Add an empty spacer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_cover_page():
    """Add cover page."""
    for _ in range(7):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("榴芒一刻")
    run.font.size = Pt(38)
    run.bold = True
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)  # dark green

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("品牌竞品研究咨询报告")
    run2.font.size = Pt(26)
    run2.bold = True
    run2.font.name = FONT_NAME
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    for _ in range(4):
        add_paragraph_spacer()

    # Subtitle info
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("燃创咨询 BreaC Lab")
    run3.font.size = Pt(16)
    run3.font.name = FONT_NAME
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("2026-07-18")
    run4.font.size = Pt(14)
    run4.font.name = FONT_NAME
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    doc.add_page_break()


def add_heading_styled(text, level):
    """Add a heading with proper styling."""
    sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13), 4: Pt(12), 5: Pt(11)}
    before = {1: 18, 2: 14, 3: 10, 4: 8, 5: 6}
    after = {1: 12, 2: 8, 3: 6, 4: 4, 5: 4}
    bold_levels = {1: True, 2: True, 3: True, 4: True, 5: False}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before.get(level, 6))
    p.paragraph_format.space_after = Pt(after.get(level, 4))
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, BODY_SIZE), bold=bold_levels.get(level, False))
    return p


def add_image_centered(img_path, caption=None):
    """Add an image centered, with optional caption."""
    if not os.path.exists(img_path):
        print(f"WARNING: Image not found: {img_path}")
        return
    add_paragraph_spacer(6)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after = Pt(2)
    run = p_img.add_run()
    run.add_picture(img_path, width=IMAGE_WIDTH)

    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run(caption)
        set_run_font(run_cap, size=Pt(9), bold=False)
        p_cap.paragraph_format.space_after = Pt(8)


def parse_inline_runs(text):
    """Parse inline markdown into a list of (bold, text) tuples.
    Handles **bold**, *italic*, and `code`."""
    parts = []
    pattern = re.compile(
        r'(\*\*(.+?)\*\*|'  # bold **text**
        r'\*(.+?)\*|'       # italic *text*
        r'`(.+?)`)'         # code `text`
    )
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            parts.append((False, text[last_end:m.start()]))
        if m.group(2):  # bold
            parts.append((True, m.group(2)))
        elif m.group(3):  # italic
            parts.append((False, m.group(3)))
        elif m.group(4):  # code
            parts.append((False, m.group(4)))
        last_end = m.end()
    if last_end < len(text):
        parts.append((False, text[last_end:]))
    return parts if parts else [(False, text)]


def add_rich_paragraph(text, font_size=BODY_SIZE, alignment=None, indent_first=False,
                       left_indent_cm=None, right_indent_cm=None):
    """Add a paragraph with inline bold/italic support."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if left_indent_cm:
        p.paragraph_format.left_indent = Cm(left_indent_cm)
    if right_indent_cm:
        p.paragraph_format.right_indent = Cm(right_indent_cm)

    parts = parse_inline_runs(text)
    for is_bold, content in parts:
        run = p.add_run(content)
        set_run_font(run, size=font_size, bold=is_bold)
    return p


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Set all borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def strip_md(text):
    """Strip markdown formatting from text, returning plain text."""
    # Remove bold markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Remove code markers
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove link syntax [text](url)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def add_table_from_lines(lines):
    """Parse lines as markdown table and add to document."""
    if len(lines) < 2:
        return

    # Find separator line
    sep_idx = None
    for i, line in enumerate(lines):
        if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
            sep_idx = i
            break

    if sep_idx is None or sep_idx < 0:
        return

    # Get headers (before separator)
    if sep_idx > 0:
        header_line = lines[sep_idx - 1]
    else:
        header_line = lines[0]
    headers = [strip_md(h.strip()) for h in header_line.split('|')[1:-1]]
    cols = len(headers)
    if cols == 0:
        return

    # Get data rows (after separator)
    data_rows = []
    for line in lines[sep_idx + 1:]:
        cells = [strip_md(c.strip()) for c in line.split('|')[1:-1]]
        if len(cells) >= cols:
            data_rows.append(cells[:cols])
        elif len(cells) > 0:
            # Pad short rows
            cells += [''] * (cols - len(cells))
            data_rows.append(cells)

    if not data_rows:
        data_rows = [[''] * cols]

    # Create table
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(header)
        set_run_font(run, size=Pt(9.5), bold=True)
        set_cell_shading(cell, HEADER_GRAY)

    # Data rows
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(cell_text)
            set_run_font(run, size=Pt(9.5), bold=False)

    # Spacing after table
    add_paragraph_spacer(4)


def is_separator_line(line):
    """Check if a line is a markdown table separator."""
    return bool(re.match(r'^\|[\s\-:|]+\|$', line.strip()))


def is_list_bullet(line):
    """Check if a line starts with a list marker (- * + followed by space, not bold)."""
    stripped = line.strip()
    # Dash is always a list marker in this markdown
    if re.match(r'^-\s+', stripped) and not stripped.startswith('---'):
        return True
    # * or + followed by space and not start of bold
    if re.match(r'^[*+]\s+[^*]', stripped):
        return True
    return False


def get_list_content(line):
    """Extract content from a list item line."""
    stripped = line.strip()
    m = re.match(r'^[-*+]\s+(.+)', stripped)
    if m:
        return m.group(1)
    return stripped


def process_markdown():
    """Process the markdown file and build the DOCX."""
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    table_buffer = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ---- IMAGE ----
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if img_match:
            if in_table and table_buffer:
                add_table_from_lines(table_buffer)
                table_buffer = []
                in_table = False
            caption = img_match.group(1)
            rel_path = img_match.group(2)
            abs_path = IMAGE_MAP.get(rel_path)
            if abs_path:
                add_image_centered(abs_path, caption)
            else:
                print(f"WARNING: Image not in map: {rel_path}")
            i += 1
            continue

        # ---- HORIZONTAL RULE ----
        if stripped == '---':
            if in_table and table_buffer:
                add_table_from_lines(table_buffer)
                table_buffer = []
                in_table = False
            p = doc.add_paragraph()
            run = p.add_run('─' * 60)
            set_run_font(run, size=Pt(8))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ---- TABLE DETECTION ----
        if stripped.startswith('|'):
            if in_table:
                table_buffer.append(stripped)
                i += 1
                continue
            elif i + 1 < len(lines) and is_separator_line(lines[i + 1].strip()):
                # Start of table: check if previous line is also a table line (header)
                if i > 0 and lines[i - 1].strip().startswith('|'):
                    # Header was previous line
                    table_buffer = [lines[i - 1].strip(), stripped]
                    if i + 1 < len(lines):
                        table_buffer.append(lines[i + 1].strip())
                    i += 2
                else:
                    table_buffer = [stripped, lines[i + 1].strip()]
                    i += 2
                in_table = True
                continue
            else:
                # Not a table separator follows, treat as text
                pass

        # End table on non-table line
        if in_table and not stripped.startswith('|'):
            add_table_from_lines(table_buffer)
            table_buffer = []
            in_table = False

        # ---- HEADINGS ----
        h_match = re.match(r'^(#{1,5})\s+(.+)', line)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2).strip()
            add_heading_styled(heading_text, level)
            i += 1
            continue

        # ---- BLOCKQUOTE (> ...) ----
        if stripped.startswith('>'):
            quote_text = re.sub(r'^>\s*', '', stripped)
            if quote_text.strip():
                add_rich_paragraph(quote_text, font_size=Pt(10), left_indent_cm=1.0)
            else:
                add_paragraph_spacer()
            i += 1
            continue

        # ---- UNORDERED LIST ----
        if is_list_bullet(line):
            content = get_list_content(line)
            p = doc.add_paragraph()
            run = p.add_run('• ')
            set_run_font(run, size=BODY_SIZE, bold=False)
            parts = parse_inline_runs(content)
            for is_bold, text in parts:
                run2 = p.add_run(text)
                set_run_font(run2, size=BODY_SIZE, bold=is_bold)
            p.paragraph_format.left_indent = Cm(1.0)
            i += 1
            continue

        # ---- EMPTY LINE ----
        if not stripped:
            add_paragraph_spacer(2)
            i += 1
            continue

        # ---- NORMAL TEXT ----
        add_rich_paragraph(stripped)
        i += 1

    # Flush remaining table buffer
    if table_buffer:
        add_table_from_lines(table_buffer)


# --- MAIN EXECUTION ---
print("Creating cover page...")
add_cover_page()

print("Processing markdown content...")
process_markdown()

# Save
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
print(f"Saving to {OUTPUT_PATH}...")
doc.save(OUTPUT_PATH)

file_size = os.path.getsize(OUTPUT_PATH)
print(f"Done! File size: {file_size:,} bytes ({file_size / 1024:.1f} KB)")
if file_size > 100 * 1024:
    print("✓ File size > 100KB — PASS")
else:
    print("⚠ File size < 100KB — check output")
