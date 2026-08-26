#!/usr/bin/env python3
"""专业排版脚本：中国财神·诸神·山海经首饰文化研究 DOCX 重排。

相比 convert_md_to_docx.py 的改进：
1. 封面页（公司/主标题/副标题/委托方/框架/日期）
2. 全篇统一苹方字体 + eastAsia 字体（修复标题没有中文字体的 bug）
3. 标题加粗分级 + 藏青主题色
4. 正文 1.5 倍行距、两端对齐
5. 两个纯文字段落转表格：28个IP四维打分总表、数据来源逐条对照表
6. 页脚页码
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 主题色：藏青 + 金（珠宝/祈福主题）
NAVY = RGBColor(0x1F, 0x38, 0x64)
NAVY_MID = RGBColor(0x2F, 0x54, 0x96)
NAVY_LIGHT = RGBColor(0x44, 0x72, 0xC4)
GOLD = RGBColor(0xBF, 0x90, 0x00)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT = 'PingFang SC'

CENTER = WD_ALIGN_PARAGRAPH.CENTER


def set_eastasia(style_or_run_el, font=FONT):
    """给样式或 run 的 rPr 设置 eastAsia 字体。"""
    rpr = style_or_run_el.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:eastAsia'), font)


def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    set_eastasia(normal.element)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_cfg = [
        ('Title', 26, True, NAVY),
        ('Heading 1', 18, True, NAVY),
        ('Heading 2', 15, True, NAVY_MID),
        ('Heading 3', 13, True, NAVY_LIGHT),
        ('Heading 4', 12, True, DARK),
    ]
    for sname, size, bold, color in heading_cfg:
        st = doc.styles[sname]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = color
        set_eastasia(st.element)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = CENTER
    r = p.add_run('燃创咨询（BreaC）')
    r.font.name = FONT; r.font.size = Pt(14); r.bold = True; r.font.color.rgb = GOLD
    set_eastasia(r._element)

    p = doc.add_paragraph(); p.alignment = CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run('中国财神·诸神·山海经')
    r.font.name = FONT; r.font.size = Pt(28); r.bold = True; r.font.color.rgb = NAVY
    set_eastasia(r._element)
    r.add_break()
    r2 = p.add_run('首饰文化研究')
    r2.font.name = FONT; r2.font.size = Pt(28); r2.bold = True; r2.font.color.rgb = NAVY
    set_eastasia(r2._element)

    p = doc.add_paragraph(); p.alignment = CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run('研究报告')
    r.font.name = FONT; r.font.size = Pt(16); r.font.color.rgb = GRAY
    set_eastasia(r._element)

    p = doc.add_paragraph(); p.alignment = CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run('━━━━━━━━━━━━')
    r.font.size = Pt(11); r.font.color.rgb = GOLD

    info = [
        ('委托方', '逸凡（曼拾 / 燃创创始人）'),
        ('研究框架', 'IP资产盘点 → 符号价值评估 → 商业化现状 → 消费者洞察 → 机会地图与创品 → 切入建议'),
        ('报告日期', '2026-08-23'),
    ]
    for label, val in info:
        p = doc.add_paragraph(); p.alignment = CENTER
        p.paragraph_format.space_before = Pt(14)
        r = p.add_run(f'{label}：')
        r.font.name = FONT; r.font.size = Pt(11); r.bold = True; r.font.color.rgb = NAVY_MID
        set_eastasia(r._element)
        r = p.add_run(val)
        r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = DARK
        set_eastasia(r._element)

    doc.add_page_break()


def set_cell(cell, text, bold=False, size=Pt(9), color=None, bg=None):
    cell.text = ''
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(text)
    r.font.name = FONT; r.font.size = size; r.bold = bold
    set_eastasia(r._element)
    if color is not None:
        r.font.color.rgb = color
    if bg:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg)
        cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, headers, rows, header_bg='1F3864', zebra_bg='EDF1F8', col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=Pt(9.5), color=WHITE, bg=header_bg)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, val in enumerate(row):
            bg = zebra_bg if r_idx % 2 == 1 else None
            set_cell(cells[c_idx], val, size=Pt(9))
            if bg:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), bg)
                cells[c_idx]._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for row in table.rows:
            for c_idx, w in enumerate(col_widths):
                row.cells[c_idx].width = w
    doc.add_paragraph()
    return table


def parse_28_symbols(text):
    pattern = r'([\u4e00-\u9fa5·]+?)(寓意\d+、识别度\d+、叙事\d+、成熟度(?:半成熟|成熟|空白))'
    rows = []
    for name, scores in re.findall(pattern, text):
        m = re.match(r'寓意(\d+)、识别度(\d+)、叙事(\d+)、成熟度(半成熟|成熟|空白)', scores)
        if m:
            y, s, x, c = m.groups()
            rows.append([name, y, s, x, c])
    return rows


def parse_data_sources(paragraphs):
    rows = []
    for para in paragraphs:
        t = para.strip()
        if not t or ('，来源' not in t and '，证据层级' not in t):
            continue
        if '，来源' in t:
            dp, rest = t.split('，来源', 1)
            src = tier = rel = ''
            if '，证据层级' in rest:
                src, rest2 = rest.split('，证据层级', 1)
                if '，可靠度' in rest2:
                    tier, rel = rest2.split('，可靠度', 1)
                    rel = rel.rstrip('。')
                else:
                    tier = rest2.rstrip('。')
            else:
                src = rest.rstrip('。')
            rows.append([dp.strip(), src.strip(), tier.strip(), rel.strip()])
        else:
            rows.append([t, '', '', ''])
    return rows


def add_para(doc, text, style='Normal', size=Pt(11), bold=False, color=None,
             align=None, space_after=Pt(6)):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    # 解析 **bold** 标记
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.font.name = FONT; r.font.size = size; r.bold = True
            set_eastasia(r._element)
            if color is not None:
                r.font.color.rgb = color
        else:
            r = para.add_run(part)
            r.font.name = FONT; r.font.size = size; r.bold = bold
            set_eastasia(r._element)
            if color is not None:
                r.font.color.rgb = color
    return para


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = CENTER
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.name = FONT; run.font.size = Pt(9); run.font.color.rgb = GRAY


def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    add_page_number_footer(doc)
    add_cover(doc)

    lines = content.split('\n')
    i = 0
    first_h1_skipped = False

    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过文档主标题（封面已展示）
        if not first_h1_skipped and line.startswith('# ') and not line.startswith('## '):
            first_h1_skipped = True
            i += 1
            continue

        # 特殊段：28个IP四维打分总表
        if line.startswith('## ') and '四维打分总表' in line:
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            ordered = []  # (is_score, text)
            while i < len(lines) and not (lines[i].startswith('## ') or lines[i].startswith('# ')):
                cur = lines[i].strip()
                if not cur:
                    i += 1
                    continue
                is_score = bool(re.search(r'寓意\d+、识别度\d+、叙事\d+、成熟度', cur))
                ordered.append((is_score, cur))
                i += 1
            score_lines = [t for s, t in ordered if s]
            # 表格前的说明
            first_score_idx = next((k for k, (s, _) in enumerate(ordered) if s), None)
            last_score_idx = max((k for k, (s, _) in enumerate(ordered) if s), default=-1)
            for k, (s, t) in enumerate(ordered):
                if not s and (first_score_idx is None or k < first_score_idx):
                    add_para(doc, t)
            rows = parse_28_symbols('\n'.join(score_lines))
            if rows:
                add_table(doc,
                          ['符号', '寓意', '视觉识别度', '叙事厚度', '商业化成熟度'],
                          rows,
                          col_widths=[Cm(3.2), Cm(1.5), Cm(2.2), Cm(2.0), Cm(2.6)])
            # 表格后的结论
            for k, (s, t) in enumerate(ordered):
                if not s and k > last_score_idx:
                    add_para(doc, t)
            continue

        # 特殊段：数据来源逐条对照表
        if line.startswith('## ') and '数据来源逐条对照表' in line:
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            ordered = []  # (is_src, text)
            while i < len(lines) and not (lines[i].startswith('## ') or lines[i].startswith('# ')):
                cur = lines[i].strip()
                if not cur:
                    i += 1
                    continue
                is_src = ('，来源' in cur or '，证据层级' in cur)
                ordered.append((is_src, cur))
                i += 1
            src_lines = [t for s, t in ordered if s]
            first_src_idx = next((k for k, (s, _) in enumerate(ordered) if s), None)
            last_src_idx = max((k for k, (s, _) in enumerate(ordered) if s), default=-1)
            for k, (s, t) in enumerate(ordered):
                if not s and (first_src_idx is None or k < first_src_idx):
                    add_para(doc, t)
            rows = parse_data_sources(src_lines)
            if rows:
                add_table(doc,
                          ['数据点', '来源', '证据层级', '可靠度'],
                          rows,
                          col_widths=[Cm(5.0), Cm(4.0), Cm(2.0), Cm(4.0)])
            for k, (s, t) in enumerate(ordered):
                if not s and k > last_src_idx:
                    add_para(doc, t)
            continue

        # h1 部分标题
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # 引用块
        if line.startswith('> '):
            block = []
            while i < len(lines) and lines[i].startswith('> '):
                block.append(lines[i][2:])
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.8)
            parts = re.split(r'(\*\*.*?\*\*)', '\n'.join(block))
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = para.add_run(part[2:-2]); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GRAY; r.font.name = FONT; set_eastasia(r._element)
                else:
                    r = para.add_run(part); r.font.size = Pt(10); r.font.color.rgb = GRAY; r.font.name = FONT; set_eastasia(r._element)
            continue

        # h2
        if line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # h3
        if line.startswith('### ') and not line.startswith('#### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # h4
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # markdown 表格（通用）
        if i + 1 < len(lines) and lines[i].startswith('|') and lines[i + 1].startswith('|---'):
            header_cells = [c.strip() for c in lines[i].split('|')[1:-1]]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                cells = [c.strip() for c in lines[i].split('|')[1:-1]]
                if len(cells) == len(header_cells):
                    rows.append(cells)
                i += 1
            if header_cells and rows:
                add_table(doc, header_cells, rows)
            continue

        # 水平线
        if line.strip() == '---':
            p = doc.add_paragraph(); p.alignment = CENTER
            r = p.add_run('─' * 40); r.font.color.rgb = GOLD; r.font.size = Pt(8)
            i += 1
            continue

        # 整行加粗（**...**）
        if line.startswith('**') and line.rstrip().endswith('**'):
            add_para(doc, line.strip()[2:-2], size=Pt(11.5), bold=True)
            i += 1
            continue

        # 无序列表
        if line.startswith('- ') or line.startswith('* '):
            list_lines = []
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ')
                                      or (lines[i].startswith('  ') and not lines[i].strip().startswith('|'))):
                list_lines.append(lines[i])
                i += 1
            for ll in list_lines:
                t = ll.strip()
                if t.startswith('- '):
                    t = t[2:]
                elif t.startswith('* '):
                    t = t[2:]
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.line_spacing = 1.5
                parts = re.split(r'(\*\*.*?\*\*)', t)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = para.add_run(part[2:-2]); r.bold = True; r.font.size = Pt(11); r.font.name = FONT; set_eastasia(r._element)
                    else:
                        r = para.add_run(part); r.font.size = Pt(11); r.font.name = FONT; set_eastasia(r._element)
            continue

        # 有序列表
        if re.match(r'^\d+\.\s', line):
            list_lines = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i]):
                list_lines.append(lines[i])
                i += 1
            for ll in list_lines:
                t = re.sub(r'^\d+\.\s+', '', ll)
                para = doc.add_paragraph(style='List Number')
                para.paragraph_format.line_spacing = 1.5
                parts = re.split(r'(\*\*.*?\*\*)', t)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = para.add_run(part[2:-2]); r.bold = True; r.font.size = Pt(11); r.font.name = FONT; set_eastasia(r._element)
                    else:
                        r = para.add_run(part); r.font.size = Pt(11); r.font.name = FONT; set_eastasia(r._element)
            continue

        # 普通段落
        if line.strip():
            add_para(doc, line.strip())

        i += 1

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f'DOCX saved: {docx_path}')


if __name__ == '__main__':
    md = '/Users/yifansmacmini/.openclaw/workspace/supermind/memory/财神诸神山海经首饰文化研究-20260823.md'
    out = '/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/财神诸神山海经首饰文化研究-20260823-V2.docx'
    convert(md, out)
