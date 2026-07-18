#!/usr/bin/env python3
"""榴芒一刻品牌竞品研究咨询报告 DOCX 生成脚本"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

MARKDOWN_PATH = os.path.join(os.path.dirname(__file__), "output/content/report_final.md")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "output/reports/liumang_final.docx")

os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

doc = Document()

# ─── Page Setup ───
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ─── Style definitions ───
style = doc.styles['Normal']
style.font.name = '等线'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
rPr = style.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
rPr.append(rFonts)

for level in [1, 2, 3]:
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = '等线'
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    h_rPr = h_style.element.get_or_add_rPr()
    h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
    h_rPr.append(h_rFonts)
    if level == 1:
        h_style.font.size = Pt(18)
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(24)
        h_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h_style.font.size = Pt(14)
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(18)
        h_style.paragraph_format.space_after = Pt(8)
    elif level == 3:
        h_style.font.size = Pt(12)
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(6)

def set_cell_font(cell, text, bold=False, size=Pt(10), font_name='等线'):
    """Set cell text with proper font"""
    for paragraph in cell.paragraphs:
        paragraph.clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
    rPr.append(rFonts)

def set_cell_shading(cell, color='D9D9D9'):
    """Set cell background color"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table):
    """Add borders to all cells"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_paragraph_with_style(doc, text, style_name='Normal', bold=False, font_size=Pt(10.5)):
    """Add a paragraph with Chinese font support"""
    p = doc.add_paragraph(style=style_name)
    # Clear default runs
    for r in p.runs:
        r.text = ''
    
    # Process inline bold markers: **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        else:
            run = p.add_run(part)
        run.font.name = '等线'
        run.font.size = font_size
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
        rPr.append(rFonts)
    return p

def parse_markdown_to_docx(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table_rows = []
    is_header_row = False
    
    # Skip cover page generation for now - we'll add it at the end
    first_content_line = 0
    for idx, line in enumerate(lines):
        if line.startswith('# 第一部分'):
            first_content_line = idx
            break
    
    # Add title on first page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run('榴芒一刻')
    run.font.name = '等线'
    run.font.size = Pt(26)
    run.font.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
    rPr.append(rFonts)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run('品牌竞品研究咨询报告')
    run.font.name = '等线'
    run.font.size = Pt(22)
    run.font.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
    rPr.append(rFonts)
    
    # Cover info lines
    cover_info = [
        ('报告日期：2026年7月18日', Pt(12)),
        ('研究范围：榴芒一刻品牌全维度分析 + 竞品五维扫描 + 创始人深度画像 + 创品策略（含归经配伍研究）', Pt(10)),
        ('框架版本：Schema v1.3', Pt(10)),
        ('文档密级：保密', Pt(10)),
    ]
    for text, size in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = '等线'
        run.font.size = size
        run.font.color.rgb = RGBColor(100, 100, 100)
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
        rPr.append(rFonts)
    
    doc.add_page_break()
    
    # Process from first content line
    i = first_content_line
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle table rows
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
                is_header_row = True
            
            # Skip separator rows like |---|---|
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                is_header_row = False
                i += 1
                continue
            
            # Parse table row
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            table_rows.append((cells, is_header_row))
            is_header_row = False
            i += 1
            continue
        else:
            if in_table:
                # End of table - render it
                in_table = False
                if table_rows:
                    num_cols = max(len(row[0]) for row in table_rows)
                    num_rows = len(table_rows)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(table)
                    
                    for ri, (cells, is_header) in enumerate(table_rows):
                        for ci, cell_text in enumerate(cells):
                            if ci < num_cols:
                                cell = table.cell(ri, ci)
                                set_cell_font(cell, cell_text, bold=is_header, size=Pt(9))
                                if is_header:
                                    set_cell_shading(cell, 'D9D9D9')
                    
                    doc.add_paragraph()  # spacing after table
                    table_rows = []
            
            # Empty line
            if not line.strip():
                i += 1
                continue
            
            # Headings
            if line.startswith('# '):
                add_paragraph_with_style(doc, line[2:], 'Heading 1', font_size=Pt(18))
                i += 1
                continue
            elif line.startswith('## '):
                add_paragraph_with_style(doc, line[3:], 'Heading 2', font_size=Pt(14))
                i += 1
                continue
            elif line.startswith('### '):
                add_paragraph_with_style(doc, line[4:], 'Heading 3', font_size=Pt(12))
                i += 1
                continue
            elif line.startswith('#### '):
                add_paragraph_with_style(doc, line[5:], 'Heading 3', font_size=Pt(11))
                i += 1
                continue
            
            # Blockquotes (metadata / note lines)
            if line.startswith('> '):
                add_paragraph_with_style(doc, line[2:], font_size=Pt(9))
                i += 1
                continue
            
            # Horizontal rules
            if line.strip() == '---':
                doc.add_paragraph()
                i += 1
                continue
            
            # Bold standalone lines (section markers like **01 榴莲半熟芝士**)
            bold_match = re.match(r'^\*\*(.+?)\*\*$', line.strip())
            if bold_match:
                add_paragraph_with_style(doc, bold_match.group(1), bold=True, font_size=Pt(10.5))
                i += 1
                continue
            
            # Star-prefixed subpoints
            if line.strip().startswith('⭐'):
                add_paragraph_with_style(doc, line.strip(), font_size=Pt(9))
                i += 1
                continue
            
            # Regular paragraph
            add_paragraph_with_style(doc, line, font_size=Pt(10.5))
            i += 1
    
    # Handle any lingering table
    if in_table and table_rows:
        num_cols = max(len(row[0]) for row in table_rows)
        num_rows = len(table_rows)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)
        for ri, (cells, is_header) in enumerate(table_rows):
            for ci, cell_text in enumerate(cells):
                if ci < num_cols:
                    cell = table.cell(ri, ci)
                    set_cell_font(cell, cell_text, bold=is_header, size=Pt(9))
                    if is_header:
                        set_cell_shading(cell, 'D9D9D9')

parse_markdown_to_docx(MARKDOWN_PATH)

# ─── Add page numbers ───
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Simple page number
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)

doc.save(OUTPUT_PATH)
print(f"DOCX saved to: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH)} bytes ({os.path.getsize(OUTPUT_PATH)/1024:.1f} KB)")
