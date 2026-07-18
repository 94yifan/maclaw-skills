#!/usr/bin/env python3
"""Convert markdown report to DOCX with bold formatting."""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_bold_text(paragraph, text, bold=True, size=None, color=None, font_name=None):
    """Add a run with optional bold, size, color."""
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    return run

def add_paragraph_with_markdown(doc, text, style='Normal', font_size=Pt(11), font_name='Microsoft YaHei'):
    """Parse and add paragraph text, handling **bold** markers."""
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(6)
    
    # Parse **bold** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            if font_size:
                run.font.size = font_size
            if font_name:
                run.font.name = font_name
        else:
            run = para.add_run(part)
            if font_size:
                run.font.size = font_size
            if font_name:
                run.font.name = font_name
    
    return para

def set_cell_text(cell, text, bold=False, size=Pt(9), font_name='Microsoft YaHei'):
    """Set cell text with optional bold."""
    cell.text = ''
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    
    # Handle **bold** in cells
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.size = size
            run.font.name = font_name
        else:
            run = para.add_run(part)
            run.font.size = size
            run.font.name = font_name

def add_table_from_md_row(doc, headers, rows, header_bold=True):
    """Add a table from markdown-style headers and rows."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h.strip(), bold=True, size=Pt(9))
        set_cell_shading(cell, '2F5496')
        # Set header text color to white
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < len(headers):
                cell = table.rows[r_idx + 1].cells[c_idx]
                set_cell_text(cell, val.strip(), size=Pt(9))
                if r_idx % 2 == 1:
                    set_cell_shading(cell, 'D6E4F0')
    
    doc.add_paragraph()  # space after table
    return table

def convert_md_to_docx(md_path, docx_path):
    """Convert markdown report to formatted DOCX."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Title (h1 heading)
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            heading = doc.add_heading(title_text, level=0)
            for run in heading.runs:
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            i += 1
            continue
        
        # Blockquote
        if line.startswith('> '):
            block_lines = []
            while i < len(lines) and lines[i].startswith('> '):
                block_lines.append(lines[i][2:])
                i += 1
            block_text = '\n'.join(block_lines)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            # Parse **bold** in blockquote
            parts = re.split(r'(\*\*.*?\*\*)', block_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    run.font.name = 'Microsoft YaHei'
                else:
                    run = para.add_run(part)
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    run.font.name = 'Microsoft YaHei'
            continue
        
        # h2 heading (## )
        if line.startswith('## ') and not line.startswith('### '):
            heading_text = line[3:].strip()
            heading = doc.add_heading(heading_text, level=1)
            for run in heading.runs:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
            i += 1
            continue
        
        # h3 heading (### )
        if line.startswith('### ') and not line.startswith('#### '):
            heading_text = line[4:].strip()
            heading = doc.add_heading(heading_text, level=2)
            for run in heading.runs:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x37, 0x5F, 0x9E)
            i += 1
            continue
        
        # h4 heading (#### )
        if line.startswith('#### '):
            heading_text = line[5:].strip()
            heading = doc.add_heading(heading_text, level=3)
            for run in heading.runs:
                run.font.size = Pt(11.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue
        
        # Table detection (starts with |---|)
        if i + 1 < len(lines) and lines[i].startswith('|') and lines[i+1].startswith('|---'):
            table_lines = [lines[i]]
            i += 1  # skip separator
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                # Parse header
                header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                # Parse rows
                rows = []
                for tl in table_lines[1:]:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    if len(cells) == len(header_cells):
                        rows.append(cells)
                
                if header_cells and rows:
                    add_table_from_md_row(doc, header_cells, rows)
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('─' * 50)
            i += 1
            continue
        
        # Bold title text (e.g., **发现一：...**  or **核心锚点：**)
        if line.startswith('**') and line.rstrip().endswith('**'):
            text = line.strip()[2:-2]
            add_paragraph_with_markdown(doc, text, font_size=Pt(11.5))
            i += 1
            continue
        
        # List items
        if line.startswith('- ') or line.startswith('* '):
            list_lines = []
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ') or 
                                       (lines[i].startswith('  ') and not lines[i].strip().startswith('|'))):
                list_lines.append(lines[i])
                i += 1
            
            for ll in list_lines:
                text = ll.strip()
                if text.startswith('- '):
                    text = text[2:]
                elif text.startswith('* '):
                    text = text[2:]
                elif text.startswith('  - '):
                    text = text[4:]
                
                para = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.name = 'Microsoft YaHei'
                    else:
                        run = para.add_run(part)
                        run.font.size = Pt(11)
                        run.font.name = 'Microsoft YaHei'
            continue
        
        # Numbered items (1. , 2. , etc.)
        if re.match(r'^\d+\.\s', line):
            list_lines = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i]):
                list_lines.append(lines[i])
                i += 1
            
            for ll in list_lines:
                text = re.sub(r'^\d+\.\s+', '', ll)
                para = doc.add_paragraph(style='List Number')
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.name = 'Microsoft YaHei'
                    else:
                        run = para.add_run(part)
                        run.font.size = Pt(11)
                        run.font.name = 'Microsoft YaHei'
            continue
        
        # Regular text paragraph (non-empty)
        if line.strip():
            add_paragraph_with_markdown(doc, line.strip())
        
        i += 1
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f"DOCX saved to: {docx_path}")

if __name__ == '__main__':
    md_path = '/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/content/report_final_beiwei47.md'
    docx_path = '/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/beiwei47_final.docx'
    convert_md_to_docx(md_path, docx_path)
