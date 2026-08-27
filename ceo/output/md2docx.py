#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""md -> 排好版的 docx（中文方案文档专用）"""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1] if len(sys.argv) > 1 else "/Users/yifansmacmini/.openclaw/workspace/ceo/output/老字号新活力青年数字营销官孵化计划完整方案V3.md"
DST = sys.argv[2] if len(sys.argv) > 2 else "/Users/yifansmacmini/.openclaw/workspace/ceo/output/老字号新活力青年数字营销官孵化计划完整方案V3-排版版.docx"

SONG = "宋体"   # 正文
HEI = "黑体"    # 标题
GRAY = RGBColor(0x59, 0x59, 0x59)


def set_font(run, cn=SONG, size=12, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn)
    if color:
        run.font.color.rgb = color


def add_runs(par, text, cn=SONG, size=12, bold=False, italic=False, color=None):
    """解析 **加粗** 与 *斜体* 行内标记"""
    tokens = re.split(r"(\*\*.+?\*\*|\*[^*\n]+\*)", text)
    for t in tokens:
        if not t:
            continue
        if t.startswith("**") and t.endswith("**"):
            r = par.add_run(t[2:-2])
            set_font(r, cn, size, bold=True, color=color)
        elif t.startswith("*") and t.endswith("*") and len(t) > 2:
            r = par.add_run(t[1:-1])
            set_font(r, cn, size, bold=bold, italic=True, color=color)
        else:
            r = par.add_run(t)
            set_font(r, cn, size, bold=bold, italic=italic, color=color)


def para_body(doc, text, size=12, indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if indent:
        pf.first_line_indent = Pt(size * 2)
    add_runs(p, text, size=size)
    return p


def para_heading(doc, text, level):
    sizes = {1: 18, 2: 15, 3: 13}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    add_runs(p, text, cn=HEI, size=sizes.get(level, 12), bold=True)
    return p


def para_list(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.6 if level == 1 else 1.2)
    add_runs(p, text, size=12)
    return p


def para_quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.line_spacing = 1.5
    pf.space_after = Pt(8)
    add_runs(p, text, size=12, italic=True, color=GRAY)
    return p


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def mark_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ""
            txt = row[j] if j < len(row) else ""
            lines = txt.split("\n")
            first = True
            for ln in lines:
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                pf = p.paragraph_format
                pf.line_spacing = 1.3
                pf.space_after = Pt(2)
                add_runs(p, ln, size=10.5, bold=(i == 0))
            if i == 0:
                shade_cell(cell, "D9E2F3")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
        if i == 0:
            mark_header_row(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, "黄浦区\u201c老字号·新活力\u201d", cn=HEI, size=26, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, "青年数字营销官孵化计划", cn=HEI, size=26, bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, "完整执行方案", cn=HEI, size=18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, "（比赛 + 培养计划 + 后续人才计划）", cn=SONG, size=14)
    for _ in range(3):
        doc.add_paragraph()
    info = [
        "主办单位：黄浦区就业促进中心",
        "项目支持方：曼拾（广告代理/数字营销专业机构）",
        "核心平台方：小红书（办公地与注册地均在黄浦区）",
        "方案版本：V3.0（完整定稿版）",
        "编制日期：2026年8月",
    ]
    for line in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, line, size=13)
    doc.add_page_break()


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    set_font(run, SONG, 10)


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    add_cover(doc)
    add_page_number_footer(doc)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            para_heading(doc, m.group(2).strip(), level)
            i += 1
            continue
        # 表格
        if line.startswith("|") and i + 1 < n and re.match(r"^\|[\s:\-|]+\|?\s*$", lines[i + 1]):
            rows = []
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(header)
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue
        if line.startswith(">"):
            para_quote(doc, line.lstrip(">").strip())
            i += 1
            continue
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            level = 2 if len(m.group(1)) > 0 else 1
            para_list(doc, m.group(2).strip(), level)
            i += 1
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            level = 2 if len(m.group(1)) > 0 else 1
            para_list(doc, m.group(2).strip(), level)
            i += 1
            continue
        para_body(doc, line.strip())
        i += 1

    doc.save(DST)
    print("SAVED:", DST)


if __name__ == "__main__":
    main()
