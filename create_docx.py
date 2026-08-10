#!/usr/bin/env python3
"""从文本文件生成整理好的 Word 文档"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, re

doc = Document()

# Default font
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Title
title = doc.add_heading('梁文锋4小时谈话全文：比赚钱更重要的，是团队的愿景和价值观', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Meta
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('来源：笔记侠 · 本文内容根据网传PDF文字稿整理 | 2026年7月23日 · 16667字 · 35分钟阅读')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Notes
notes_para = doc.add_paragraph()
nr = notes_para.add_run(
    '本文内容根据网传PDF文字稿整理，未经梁文锋本人及DeepSeek方面确认，笔记侠未获原始录音核验。'
    '原稿由语音识别转写并经人工整理，在不改变原录音文字稿的基础上略做行文修改及增加备注；'
    '专有名词、数字及表述可能有误，请谨慎参考。'
)
nr.font.size = Pt(9)
nr.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# Read content file
text_path = os.path.expanduser('~/.openclaw/workspace/liangwenfeng_content.txt')
with open(text_path, 'r', encoding='utf-8') as f:
    text = f.read()

# Pattern for section headers: 一、二、... 八、
section_pat = re.compile(r'^[一二三四五六七八九十]+、')

lines = text.split('\n')
in_quote_block = False

for line in lines:
    stripped = line.strip()
    if not stripped:
        # Keep blank lines as paragraph breaks sparingly
        continue

    # Section header
    if section_pat.match(stripped):
        doc.add_heading(stripped, level=1)
        continue

    # Subsection: starts with digit and dot like "1. xxxx"
    if re.match(r'^\d+\.\s', stripped):
        doc.add_heading(stripped, level=2)
        continue

    # Normal paragraph
    p = doc.add_paragraph(stripped)

# Save
output = os.path.expanduser('~/Desktop/梁文锋4小时谈话全文_DeepSeek.docx')
doc.save(output)
print(f'Done: {output}')
