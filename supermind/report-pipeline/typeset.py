#!/usr/bin/env python3
"""通用报告排版模块 (typeset)。

将 Markdown 研究报告转换为专业排版的 DOCX 文件，适用于燃创咨询各类研究报告。

排版四件套：
  1. 封面页（公司名 + 主标题 + 副标题 + 委托方/研究框架/报告日期）
  2. 全篇 PingFang SC 中文字体 + eastAsia 正确设置（标题分级加粗配色）
  3. 数据表格化（从结构化数据生成 python-docx 真表格，带表头底色与斑马纹）
  4. 页脚页码

正文默认 1.5 倍行距、两端对齐。视觉风格：藏青主题色 + 金色点缀。

用法示例：
    from typeset import ReportTypesetter, ReportConfig

    config = ReportConfig(
        company='燃创咨询（BreaC）',
        title='中国财神·诸神·山海经首饰文化研究',
        subtitle='研究报告',
        client='逸凡（曼拾 / 燃创创始人）',
        framework='IP资产盘点 → 符号价值评估 → 商业化现状 → 消费者洞察 → 机会地图与创品 → 切入建议',
        date='2026-08-23',
    )
    ts = ReportTypesetter(config)
    ts.convert('input.md', 'output.docx')

也可通过 JSON 配置文件加载：
    config = ReportConfig.from_json('config.json')
"""

import re
import os
import json
from dataclasses import dataclass, field
from typing import List, Optional, Tuple, Dict, Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =====================================================================
# 主题色常量（藏青 + 金色，可由配置覆盖）
# =====================================================================

class ThemeColors:
    """藏青主题色 + 金色点缀。"""
    NAVY       = RGBColor(0x1F, 0x38, 0x64)
    NAVY_MID   = RGBColor(0x2F, 0x54, 0x96)
    NAVY_LIGHT = RGBColor(0x44, 0x72, 0xC4)
    GOLD       = RGBColor(0xBF, 0x90, 0x00)
    DARK       = RGBColor(0x33, 0x33, 0x33)
    GRAY       = RGBColor(0x66, 0x66, 0x66)
    WHITE      = RGBColor(0xFF, 0xFF, 0xFF)


# 默认中文字体
DEFAULT_FONT = 'PingFang SC'

# 标题样式配置：(样式名, 字号, 加粗, 颜色)
DEFAULT_HEADING_STYLES = [
    ('Title',     26, True, ThemeColors.NAVY),
    ('Heading 1', 18, True, ThemeColors.NAVY),
    ('Heading 2', 15, True, ThemeColors.NAVY_MID),
    ('Heading 3', 13, True, ThemeColors.NAVY_LIGHT),
    ('Heading 4', 12, True, ThemeColors.DARK),
]


# =====================================================================
# 数据表格定义
# =====================================================================

@dataclass
class DataTable:
    """结构化数据表格定义，用于在报告中插入真表格。

    Attributes:
        title: 表格标题（作为 h2 插入在表格上方）
        headers: 表头列名列表
        rows: 行数据列表，每行为一个列表，长度与 headers 一致
        col_widths_cm: 可选，各列宽度（厘米），None 则自动均分
        header_bg: 表头背景色 hex，默认藏青
        zebra_bg: 斑马纹背景色 hex，默认浅蓝灰
    """
    title: str
    headers: List[str]
    rows: List[List[str]]
    col_widths_cm: Optional[List[float]] = None
    header_bg: str = '1F3864'
    zebra_bg: str = 'EDF1F8'


# =====================================================================
# 报告配置
# =====================================================================

@dataclass
class ReportConfig:
    """报告排版配置。

    通过此数据类传入报告元信息，typeset 不写死任何特定报告内容。
    """
    # ---- 必填 / 核心字段 ----
    company: str = '燃创咨询（BreaC）'   # 公司名称（封面顶部）
    title: str = ''                       # 报告主标题（封面大字）
    subtitle: str = ''                    # 副标题（如"研究报告"）

    # ---- 封面信息 ----
    client: str = ''                      # 委托方
    framework: str = ''                   # 研究框架描述
    date: str = ''                        # 报告日期

    # ---- 排版选项 ----
    font: str = DEFAULT_FONT              # 中文字体
    line_spacing: float = 1.5             # 正文行距
    font_size_pt: int = 11                # 正文字号

    # ---- 页边距（厘米）----
    margin_top: float = 2.2
    margin_bottom: float = 2.2
    margin_left: float = 2.4
    margin_right: float = 2.4

    # ---- 数据表格（可选）----
    tables: List[DataTable] = field(default_factory=list)

    # ---- 主题色覆盖（可选，保持 None 则用默认藏青+金）----
    color_navy: Optional[RGBColor] = None
    color_navy_mid: Optional[RGBColor] = None
    color_navy_light: Optional[RGBColor] = None
    color_gold: Optional[RGBColor] = None

    @classmethod
    def from_json(cls, path: str) -> 'ReportConfig':
        """从 JSON 文件加载配置。

        JSON 格式示例：
        {
            "company": "燃创咨询（BreaC）",
            "title": "xxx研究报告",
            "subtitle": "研究报告",
            "client": "委托方名称",
            "framework": "研究框架描述",
            "date": "2026-08-23",
            "tables": [
                {
                    "title": "IP打分表",
                    "headers": ["符号", "寓意", "识别度"],
                    "rows": [["财神", "9", "8"], ["关公", "8", "9"]]
                }
            ]
        }
        """
        with open(path, 'r', encoding='utf-8') as f:
            raw = json.load(f)

        tables = []
        for t in raw.get('tables', []):
            tables.append(DataTable(
                title=t['title'],
                headers=t['headers'],
                rows=t['rows'],
                col_widths_cm=t.get('col_widths_cm'),
                header_bg=t.get('header_bg', '1F3864'),
                zebra_bg=t.get('zebra_bg', 'EDF1F8'),
            ))

        # 解析颜色覆盖（hex 字符串 -> RGBColor）
        def _parse_color(hex_str: Optional[str]) -> Optional[RGBColor]:
            if not hex_str:
                return None
            h = hex_str.lstrip('#')
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        return cls(
            company=raw.get('company', '燃创咨询（BreaC）'),
            title=raw.get('title', ''),
            subtitle=raw.get('subtitle', ''),
            client=raw.get('client', ''),
            framework=raw.get('framework', ''),
            date=raw.get('date', ''),
            font=raw.get('font', DEFAULT_FONT),
            line_spacing=raw.get('line_spacing', 1.5),
            font_size_pt=raw.get('font_size_pt', 11),
            margin_top=raw.get('margin_top', 2.2),
            margin_bottom=raw.get('margin_bottom', 2.2),
            margin_left=raw.get('margin_left', 2.4),
            margin_right=raw.get('margin_right', 2.4),
            tables=tables,
            color_navy=_parse_color(raw.get('color_navy')),
            color_navy_mid=_parse_color(raw.get('color_navy_mid')),
            color_navy_light=_parse_color(raw.get('color_navy_light')),
            color_gold=_parse_color(raw.get('color_gold')),
        )

    def to_json(self, path: str):
        """导出配置到 JSON 文件（方便复用/版本管理）。"""
        raw = {
            'company': self.company,
            'title': self.title,
            'subtitle': self.subtitle,
            'client': self.client,
            'framework': self.framework,
            'date': self.date,
            'font': self.font,
            'line_spacing': self.line_spacing,
            'font_size_pt': self.font_size_pt,
            'margin_top': self.margin_top,
            'margin_bottom': self.margin_bottom,
            'margin_left': self.margin_left,
            'margin_right': self.margin_right,
            'tables': [
                {
                    'title': t.title,
                    'headers': t.headers,
                    'rows': t.rows,
                    'col_widths_cm': t.col_widths_cm,
                    'header_bg': t.header_bg,
                    'zebra_bg': t.zebra_bg,
                }
                for t in self.tables
            ],
        }
        os.makedirs(os.path.dirname(path) or '.', exist_ok=True)
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(raw, f, ensure_ascii=False, indent=2)

    # ---- 颜色访问器（优先使用自定义值，否则回退默认）----

    @property
    def c_navy(self) -> RGBColor:
        return self.color_navy or ThemeColors.NAVY

    @property
    def c_navy_mid(self) -> RGBColor:
        return self.color_navy_mid or ThemeColors.NAVY_MID

    @property
    def c_navy_light(self) -> RGBColor:
        return self.color_navy_light or ThemeColors.NAVY_LIGHT

    @property
    def c_gold(self) -> RGBColor:
        return self.color_gold or ThemeColors.GOLD


# =====================================================================
# 核心排版引擎
# =====================================================================

class ReportTypesetter:
    """通用报告排版引擎。

    使用方式：
        config = ReportConfig(title='xxx', ...)
        typesetter = ReportTypesetter(config)
        typesetter.convert('input.md', 'output.docx')
    """

    def __init__(self, config: ReportConfig):
        self.cfg = config
        self._table_inserted: set = set()  # 已插入的表格标题（防重复）

    # -----------------------------------------------------------------
    # 低级工具：字体 / 样式
    # -----------------------------------------------------------------

    @staticmethod
    def set_eastasia(style_or_run_el, font: str = DEFAULT_FONT):
        """给样式或 run 的 rPr 设置 eastAsia / ascii / hAnsi 字体。

        这是修复 Word 中中文显示为宋体/黑体的关键：必须同时设置
        w:rFonts 的 ascii、hAnsi、eastAsia 三个属性。
        """
        rpr = style_or_run_el.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), font)
        rfonts.set(qn('w:hAnsi'), font)
        rfonts.set(qn('w:eastAsia'), font)

    def _setup_styles(self, doc: Document):
        """配置文档样式：正文字体/行距/对齐 + 标题分级加粗配色。"""
        cfg = self.cfg

        # Normal 样式
        normal = doc.styles['Normal']
        normal.font.name = cfg.font
        normal.font.size = Pt(cfg.font_size_pt)
        normal.font.color.rgb = ThemeColors.DARK
        self.set_eastasia(normal.element, cfg.font)
        pf = normal.paragraph_format
        pf.line_spacing = cfg.line_spacing
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 标题样式（使用配置的颜色或默认藏青系）
        heading_cfg = [
            ('Title',     26, True, cfg.c_navy),
            ('Heading 1', 18, True, cfg.c_navy),
            ('Heading 2', 15, True, cfg.c_navy_mid),
            ('Heading 3', 13, True, cfg.c_navy_light),
            ('Heading 4', 12, True, ThemeColors.DARK),
        ]
        for sname, size, bold, color in heading_cfg:
            st = doc.styles[sname]
            st.font.name = cfg.font
            st.font.size = Pt(size)
            st.font.bold = bold
            st.font.color.rgb = color
            self.set_eastasia(st.element, cfg.font)

    def _setup_margins(self, doc: Document):
        """设置页边距。"""
        cfg = self.cfg
        for section in doc.sections:
            section.top_margin    = Cm(cfg.margin_top)
            section.bottom_margin = Cm(cfg.margin_bottom)
            section.left_margin   = Cm(cfg.margin_left)
            section.right_margin  = Cm(cfg.margin_right)

    # -----------------------------------------------------------------
    # 四件套之 1：封面页
    # -----------------------------------------------------------------

    def _add_cover(self, doc: Document):
        """插入封面页：公司名 + 主标题 + 副标题 + 分隔线 + 元信息 + 分页符。"""
        cfg = self.cfg
        font = cfg.font
        CENTER = WD_ALIGN_PARAGRAPH.CENTER

        # 顶部留白
        for _ in range(4):
            doc.add_paragraph()

        # 公司名（金色）
        p = doc.add_paragraph()
        p.alignment = CENTER
        r = p.add_run(cfg.company)
        r.font.name = font; r.font.size = Pt(14); r.bold = True
        r.font.color.rgb = cfg.c_gold
        self.set_eastasia(r._element, font)

        # 主标题（藏青大字）
        p = doc.add_paragraph()
        p.alignment = CENTER
        p.paragraph_format.space_before = Pt(36)
        r = p.add_run(cfg.title)
        r.font.name = font; r.font.size = Pt(28); r.bold = True
        r.font.color.rgb = cfg.c_navy
        self.set_eastasia(r._element, font)

        # 副标题（灰色）
        if cfg.subtitle:
            p = doc.add_paragraph()
            p.alignment = CENTER
            p.paragraph_format.space_before = Pt(18)
            r = p.add_run(cfg.subtitle)
            r.font.name = font; r.font.size = Pt(16)
            r.font.color.rgb = ThemeColors.GRAY
            self.set_eastasia(r._element, font)

        # 分隔线（金色）
        p = doc.add_paragraph()
        p.alignment = CENTER
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run('\u2550' * 14)  # ━
        r.font.size = Pt(11); r.font.color.rgb = cfg.c_gold

        # 元信息行（委托方 / 框架 / 日期）
        info_items = []
        if cfg.client:
            info_items.append(('委托方', cfg.client))
        if cfg.framework:
            info_items.append(('研究框架', cfg.framework))
        if cfg.date:
            info_items.append(('报告日期', cfg.date))

        for label, val in info_items:
            p = doc.add_paragraph()
            p.alignment = CENTER
            p.paragraph_format.space_before = Pt(14)
            r = p.add_run(f'{label}：')
            r.font.name = font; r.font.size = Pt(11); r.bold = True
            r.font.color.rgb = cfg.c_navy_mid
            self.set_eastasia(r._element, font)
            r = p.add_run(val)
            r.font.name = font; r.font.size = Pt(11)
            r.font.color.rgb = ThemeColors.DARK
            self.set_eastasia(r._element, font)

        # 分页符
        doc.add_page_break()

    # -----------------------------------------------------------------
    # 四件套之 3：数据表格
    # -----------------------------------------------------------------

    @staticmethod
    def _set_cell(cell, text: str, bold: bool = False, size=None,
                  color=None, bg: Optional[str] = None, font: str = DEFAULT_FONT):
        """设置单元格文本与格式。"""
        cell.text = ''
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing = 1.15
        r = para.add_run(text)
        r.font.name = font; r.font.size = size or Pt(9); r.bold = bold
        ReportTypesetter.set_eastasia(r._element, font)
        if color is not None:
            r.font.color.rgb = color
        if bg:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg)
            cell._tc.get_or_add_tcPr().append(shd)

    def _add_data_table(self, doc: Document, dt: DataTable):
        """将一个 DataTable 渲染为 python-docx 真表格并插入文档。"""
        font = self.cfg.font
        table = doc.add_table(rows=1, cols=len(dt.headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头行
        for i, h in enumerate(dt.headers):
            self._set_cell(
                table.rows[0].cells[i], h,
                bold=True, size=Pt(9.5), color=ThemeColors.WHITE,
                bg=dt.header_bg, font=font,
            )

        # 数据行（斑马纹）
        for r_idx, row in enumerate(dt.rows):
            cells = table.add_row().cells
            for c_idx, val in enumerate(row):
                bg = dt.zebra_bg if r_idx % 2 == 1 else None
                self._set_cell(cells[c_idx], str(val), size=Pt(9), bg=bg, font=font)
                if bg:
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), bg)
                    cells[c_idx]._tc.get_or_add_tcPr().append(shd)

        # 列宽
        if dt.col_widths_cm:
            for row in table.rows:
                for c_idx, w in enumerate(dt.col_widths_cm):
                    if c_idx < len(row.cells):
                        row.cells[c_idx].width = Cm(w)

        doc.add_paragraph()  # 表格后空行
        return table

    # -----------------------------------------------------------------
    # 四件套之 4：页脚页码
    # -----------------------------------------------------------------

    def _add_page_number_footer(self, doc: Document):
        """在第一节页脚居中插入 PAGE 域（自动页码）。"""
        footer = doc.sections[0].footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
        fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
        run.font.name = self.cfg.font; run.font.size = Pt(9); run.font.color.rgb = ThemeColors.GRAY

    # -----------------------------------------------------------------
    # 段落辅助
    # -----------------------------------------------------------------

    def _add_para(self, doc, text: str, style: str = 'Normal', size=None,
                  bold=False, color=None, align=None, space_after=None):
        """添加段落，支持内联 **bold** 标记解析。"""
        cfg = self.cfg
        font = cfg.font
        para = doc.add_paragraph(style=style)
        if align is not None:
            para.alignment = align
        if space_after is not None:
            para.paragraph_format.space_after = space_after

        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = para.add_run(part[2:-2])
                r.font.name = font; r.font.size = size or Pt(cfg.font_size_pt); r.bold = True
                self.set_eastasia(r._element, font)
                if color is not None:
                    r.font.color.rgb = color
            else:
                r = para.add_run(part)
                r.font.name = font; r.font.size = size or Pt(cfg.font_size_pt); r.bold = bold
                self.set_eastasia(r._element, font)
                if color is not None:
                    r.font.color.rgb = color
        return para

    # -----------------------------------------------------------------
    # Markdown 解析与 DOCX 转换主流程
    # -----------------------------------------------------------------

    def convert(self, md_path: str, docx_path: str):
        """将 Markdown 文件转换为排版的 DOCX 文件。

        Args:
            md_path: 输入 Markdown 文件路径
            docx_path: 输出 DOCX 文件路径
        """
        with open(md_path, encoding='utf-8') as f:
            content = f.read()
        self._convert_content(content, docx_path)

    def convert_str(self, md_content: str, docx_path: str):
        """从字符串内容转换为排版的 DOCX（便于测试/管道调用）。

        Args:
            md_content: Markdown 文本内容
            docx_path: 输出 DOCX 文件路径
        """
        self._convert_content(md_content, docx_path)

    def _convert_content(self, content: str, docx_path: str):

        doc = Document()

        # 四件套初始化
        self._setup_styles(doc)
        self._setup_margins(doc)
        self._add_page_number_footer(doc)
        self._add_cover(doc)

        # 逐行解析 Markdown
        lines = content.split('\n')
        i = 0
        first_h1_skipped = False

        while i < len(lines):
            line = lines[i].rstrip()

            # 跳过文档主标题（封面已展示）
            if not first_h1_skipped and line.startswith('# ') and not line.startswith('## '):
                first_h1_skipped = True
                i += 1
                continue

            # 检查是否到达某个预定义表格的插入位置
            inserted = self._try_insert_table_at_heading(doc, line, i, lines)
            if inserted:
                i = inserted
                continue

            # h1
            if line.startswith('# ') and not line.startswith('## '):
                doc.add_heading(line[2:].strip(), level=1)
                i += 1
                continue

            # 引用块
            if line.startswith('> '):
                i = self._parse_blockquote(doc, lines, i)
                continue

            # h2
            if line.startswith('## ') and not line.startswith('### '):
                doc.add_heading(line[3:].strip(), level=2)
                i += 1
                continue

            # h3
            if line.startswith('### ') and not line.startswith('#### '):
                doc.add_heading(line[4:].strip(), level=3)
                i += 1
                continue

            # h4
            if line.startswith('#### '):
                doc.add_heading(line[5:].strip(), level=4)
                i += 1
                continue

            # Markdown 表格（管道符语法）
            if (i + 1 < len(lines)
                    and lines[i].startswith('|')
                    and lines[i + 1].startswith('|---')):
                i = self._parse_md_table(doc, lines, i)
                continue

            # 水平线
            if line.strip() == '---':
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run('\u2500' * 40)  # ─
                r.font.color.rgb = self.cfg.c_gold; r.font.size = Pt(8)
                i += 1
                continue

            # 整行加粗
            if line.startswith('**') and line.rstrip().endswith('**'):
                self._add_para(doc, line.strip()[2:-2],
                               size=Pt(11.5), bold=True)
                i += 1
                continue

            # 无序列表
            if line.startswith('- ') or line.startswith('* '):
                i = self._parse_bullet_list(doc, lines, i)
                continue

            # 有序列表
            if re.match(r'^\d+\.\s', line):
                i = self._parse_numbered_list(doc, lines, i)
                continue

            # 普通段落
            if line.strip():
                self._add_para(doc, line.strip())

            i += 1

        # 保存
        os.makedirs(os.path.dirname(os.path.abspath(docx_path)), exist_ok=True)
        doc.save(docx_path)
        print(f'DOCX saved: {docx_path}')

    # -----------------------------------------------------------------
    # 子解析器：按需插入预定义数据表格
    # -----------------------------------------------------------------

    def _try_insert_table_at_heading(self, doc, line: str, i: int,
                                      lines: list) -> Optional[int]:
        """检查当前行是否匹配某张预定义表格的标题。

        如果匹配，在该位置渲染 DataTable 并返回新的行号；
        不匹配则返回 None，交还给主循环处理。
        """
        for dt in self.cfg.tables:
            if dt.title in self._table_inserted:
                continue
            # 匹配策略：当前 h2 行文本包含表格标题
            if line.startswith('## ') and dt.title in line:
                doc.add_heading(line[3:].strip(), level=2)
                self._table_inserted.add(dt.title)
                self._add_data_table(doc, dt)
                i += 1
                # 跳过该章节后续原始文本（已被结构化表格替代）
                while i < len(lines) and not (lines[i].startswith('# ') or
                       lines[i].rstrip().startswith('---')):
                    i += 1
                return i
        return None

    # -----------------------------------------------------------------
    # 子解析器：引用块
    # -----------------------------------------------------------------

    def _parse_blockquote(self, doc, lines: list, i: int) -> int:
        font = self.cfg.font
        block = []
        while i < len(lines) and lines[i].startswith('> '):
            block.append(lines[i][2:])
            i += 1
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.8)
        parts = re.split(r'(\*\*.*?\*\*)', '\n'.join(block))
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = para.add_run(part[2:-2])
                r.bold = True; r.font.size = Pt(10)
                r.font.color.rgb = ThemeColors.GRAY; r.font.name = font
                self.set_eastasia(r._element, font)
            else:
                r = para.add_run(part)
                r.font.size = Pt(10); r.font.color.rgb = ThemeColors.GRAY
                r.font.name = font
                self.set_eastasia(r._element, font)
        return i

    # -----------------------------------------------------------------
    # 子解析器：Markdown 表格
    # -----------------------------------------------------------------

    @staticmethod
    def _parse_md_table(doc, lines: list, i: int) -> int:
        """解析 | 列 | 分隔 | 的 Markdown 表格，转为 docx 表格。"""
        header_cells = [c.strip() for c in lines[i].split('|')[1:-1]]
        i += 2  # 跳过表头 + 分隔行
        rows = []
        while i < len(lines) and lines[i].startswith('|'):
            cells = [c.strip() for c in lines[i].split('|')[1:-1]]
            if len(cells) == len(header_cells):
                rows.append(cells)
            i += 1
        if header_cells and rows:
            # 复用 _add_data_table 逻辑（构造临时 DataTable）
            dt = DataTable(
                title='', headers=header_cells, rows=rows,
                header_bg='1F3864', zebra_bg='EDF1F8',
            )
            ReportTypesetter._add_data_table(
                ReportTypesetter.__new__(ReportTypesetter), doc, dt
            )
        return i

    # -----------------------------------------------------------------
    # 子解析器：无序列表
    # -----------------------------------------------------------------

    def _parse_bullet_list(self, doc, lines: list, i: int) -> int:
        font = self.cfg.font
        list_lines = []
        while (i < len(lines)
               and (lines[i].startswith('- ') or lines[i].startswith('* ')
                    or (lines[i].startswith('  ')
                        and not lines[i].strip().startswith('|')))):
            list_lines.append(lines[i])
            i += 1
        for ll in list_lines:
            t = ll.strip()
            if t.startswith('- '): t = t[2:]
            elif t.startswith('* '): t = t[2:]
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.line_spacing = self.cfg.line_spacing
            parts = re.split(r'(\*\*.*?\*\*)', t)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = para.add_run(part[2:-2])
                    r.bold = True; r.font.size = Pt(self.cfg.font_size_pt)
                    r.font.name = font
                    self.set_eastasia(r._element, font)
                else:
                    r = para.add_run(part)
                    r.font.size = Pt(self.cfg.font_size_pt); r.font.name = font
                    self.set_eastasia(r._element, font)
        return i

    # -----------------------------------------------------------------
    # 子解析器：有序列表
    # -----------------------------------------------------------------

    def _parse_numbered_list(self, doc, lines: list, i: int) -> int:
        font = self.cfg.font
        list_lines = []
        while i < len(lines) and re.match(r'^\d+\.\s', lines[i]):
            list_lines.append(lines[i])
            i += 1
        for ll in list_lines:
            t = re.sub(r'^\d+\.\s+', '', ll)
            para = doc.add_paragraph(style='List Number')
            para.paragraph_format.line_spacing = self.cfg.line_spacing
            parts = re.split(r'(\*\*.*?\*\*)', t)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = para.add_run(part[2:-2])
                    r.bold = True; r.font.size = Pt(self.cfg.font_size_pt)
                    r.font.name = font
                    self.set_eastasia(r._element, font)
                else:
                    r = para.add_run(part)
                    r.font.size = Pt(self.cfg.font_size_pt); r.font.name = font
                    self.set_eastasia(r._element, font)
        return i


# =====================================================================
# 便捷入口：从命令行运行（兼容 typeset_caishen.py 用法）
# =====================================================================

def main():
    """命令行入口：读取财神报告 MD，用通用模块排版输出 DOCX。"""
    md_path = '/Users/yifansmacmini/.openclaw/workspace/supermind/memory/财神诸神山海经首饰文化研究-20260823.md'
    out_path = '/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/caishen_typeset_test.docx'

    # 构造财神报告配置（演示通用模块如何参数化）
    config = ReportConfig(
        company='燃创咨询（BreaC）',
        title='中国财神·诸神·山海经\n首饰文化研究',
        subtitle='研究报告',
        client='逸凡（曼拾 / 燃创创始人）',
        framework='IP资产盘点 → 符号价值评估 → 商业化现状 → 消费者洞察 → 机会地图与创品 → 切入建议',
        date='2026-08-23',
    )

    ts = ReportTypesetter(config)
    ts.convert(md_path, out_path)


if __name__ == '__main__':
    main()
