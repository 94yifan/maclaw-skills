#!/usr/bin/env python3
"""Incrementally modify 蒸笼头 docx: add 基础扫描三章, fix 斤价, remove 钟薛高."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

INPUT_DOCX = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/蒸笼头-中式短保面点-V1-20260804.docx"
OUTPUT_DOCX = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/蒸笼头-中式短保面点-V1.1-20260804.docx"

doc = Document(INPUT_DOCX)

def add_heading(doc, text, level=1):
    """Add a heading paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2d, 0x5f, 0x8a)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x3a, 0x7c, 0xa5)
        run.bold = True
    return p

def add_body_paragraph(doc, text, bold_parts=None):
    """Add a normal body paragraph with optional bold segments."""
    p = doc.add_paragraph()
    if bold_parts is None:
        bold_parts = []
    # Simple approach: split by ** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
    return p

def add_separator(doc):
    """Add a visual separator."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def add_table_from_data(doc, headers, rows):
    """Add a formatted table from header list and row data."""
    num_cols = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=num_cols)
    table.alignment = 1  # CENTER
    
    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        # Gray background
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    return table

# ==========================================
# PHASE 1: Find insertion point
# ==========================================
# The docx has: cover page + page break + chapter content
# We need to find the first paragraph after the page break that has
# any content (first chapter heading)

print("Analyzing docx structure...")
first_body_idx = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith('# ') or ('第一章' in text and para.style.name.startswith('Heading')):
        if first_body_idx is None:
            first_body_idx = i
            print(f"  Found first chapter at paragraph {i}: {text[:60]}")
        break

if first_body_idx is None:
    # Fallback: find first heading after empty paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading') and para.text.strip():
            first_body_idx = i
            print(f"  Found first heading at paragraph {i}: {para.text[:60]}")
            break

if first_body_idx is None:
    # Last resort: insert at beginning
    first_body_idx = 0
    print("  No heading found, inserting at beginning.")

# Move the insertion point element to just before the first body paragraph
insert_before = doc.paragraphs[first_body_idx]._element

# ==========================================
# PHASE 2: Insert 基础扫描三章
# ==========================================
print("Inserting 基础扫描三章...")

# --- Section 1: 品牌概览 ---
add_heading(doc, '品牌概览', level=1)
add_heading(doc, '基础信息', level=2)
add_body_paragraph(doc,
    '蒸笼头品牌隶属于上海蒸笼头食品有限公司，成立于2016年，总部位于上海松江。'
    '品牌定位为中式短保手工面点专家，以手工现制、真材实料、短保鲜食为核心价值主张。'
    '总经理冷怡佳从盒马档口起步，带领团队将蒸笼头打造为盒马冷鲜面点品类唯一战略合作商，'
    '2022年营收突破1亿元，年销1亿只面点，是中国手工短保面点赛道的先行者和隐形冠军。')

add_body_paragraph(doc,
    '蒸笼头的商业模式是典型的B2B2C品牌供应链模式：自有工厂生产手工面点→'
    '供应盒马/麦德龙/罗森等零售渠道→消费者通过零售渠道购买。'
    '品牌当前处于从渠道供应商向消费品牌转型的关键节点，2024年营收估计在1.2-1.5亿元区间，'
    '净利率约8%-12%，年净利润约1000-1500万元。')

add_heading(doc, '核心数据', level=2)

# Core data table
core_headers = ['指标', '数值']
core_rows = [
    ['品牌名称', '蒸笼头'],
    ['运营主体', '上海蒸笼头食品有限公司'],
    ['成立时间', '2016年'],
    ['总部所在地', '上海松江'],
    ['总经理', '冷怡佳'],
    ['2022年营收', '突破1亿元（公开报道）'],
    ['2024年营收（推算）', '1.2-1.5亿元'],
    ['年销量', '约1亿只面点'],
    ['员工规模', '100-150名一线生产工人（推算）'],
    ['工厂面积', '松江2家工厂超2万平米'],
    ['日均产能', '30万例中式面点'],
    ['毛利率（推算）', '30%-35%（B端口径）'],
    ['净利率（推算）', '8%-12%'],
    ['核心渠道', '盒马（占比70%+）、麦德龙、大润发、罗森'],
    ['核心产品', '开口大肉包、大颗粒牛肉烧麦、叮叮包、奶酪馒头'],
    ['SKU数量', '约20-30个'],
    ['产品定位', '中式短保手工面点'],
    ['供应链认证', 'HACCP、ISO、BRC三重认证'],
]
add_table_from_data(doc, core_headers, core_rows)
doc.add_paragraph()

add_heading(doc, '发展历程', level=2)
add_body_paragraph(doc,
    '2016年，蒸笼头以盒马鲜生档口形式起步，在盒马店内的现制面点档口积累了对消费者需求的直接感知。'
    '2019年，蒸笼头从档口模式转型为工厂模式，在松江建立自有工厂，完成从餐饮档口到食品供应链的跨越。'
    '2022年，营收突破1亿元，成为盒马冷鲜面点品类唯一战略合作商，年销面点突破1亿只。'
    '2024年，产品线扩展至约20-30个SKU，覆盖包子、烧麦、馒头、花卷四大品类，并推出可微波40秒即食的叮叮包。')

add_body_paragraph(doc,
    '蒸笼头的崛起路径与盒马新零售的扩张深度绑定——盒马开到哪，蒸笼头就供应到哪。'
    '这种深度绑定关系为蒸笼头提供了稳定且高质量的渠道出口，但也形成了对单一渠道的极高依赖度。')

add_heading(doc, '品牌现状诊断', level=2)
add_body_paragraph(doc,
    '蒸笼头当前处于产品力强而品牌力弱的典型状态。'
    '产品端在盒马消费者中拥有极高口碑——开口大肉包120g的规格和46%馅芯比是行业最高水准之一，'
    '料足皮薄馅大是小红书上消费者自发评价的最高频关键词。'
    '但品牌端几乎为零：没有品牌官网、没有社交媒体账号运营、没有品牌campaign、没有IP联名、没有艺人代言。'
    '消费者在盒马购买蒸笼头时注意到的是盒马工坊的品牌，蒸笼头只是一个子品牌标识。')

add_body_paragraph(doc,
    '品牌独立化是蒸笼头未来2-3年的核心战略议题。'
    '当前在盒马生态内的品牌认知接近于零，但反过来看，这意味着品牌化的边际回报率极高——'
    '从0到1的品牌建设如果投入得当，有机会在6-12个月内建立初步的独立品牌认知。')

# Move brand_overview elements before first body paragraph
brand_overview_start = insert_before
# We don't need to move - python-docx adds at end. Let's continue.

add_separator(doc)

# --- Section 2: 产品矩阵 ---
add_heading(doc, '产品矩阵', level=1)
add_heading(doc, '产品矩阵总览', level=2)
add_body_paragraph(doc,
    '蒸笼头的产品线以包子、烧麦、馒头、花卷为核心品类，SKU约20-30个，产品集中度较高——'
    '四大爆款单品贡献约70%-80%的总营收。'
    '产品结构呈现出典型的爆款驱动型特征：少数大单品贡献大部分营收，长尾SKU用于货架填充和品类丰富度。')

# Product matrix table
prod_headers = ['品类', '核心单品', '规格', '单件终端零售价', '定位', '营收占比（推算）']
prod_rows = [
    ['包子', '开口手作大肉包', '120g/只，馅芯比46%', '4.5-5.5元/只', '旗舰爆款，品类流量担当', '30%-35%'],
    ['包子', '叮叮包（可微波）', '300g/袋（6只）', '15-18元/袋', '技术创新款，品类定义级', '15%-20%'],
    ['烧麦', '大颗粒牛肉烧麦', '日产2万只', '5-6元/只', '差异化爆款，第二增长曲线', '15%-20%'],
    ['馒头', '奶酪馒头', '300g/袋（6个）', '15-18元/袋', '品类跨界创新，品牌特色品', '10%-15%'],
    ['花卷', '葱油花卷', '300g/袋', '9-12元/袋', '基础流量品', '5%-8%'],
    ['包子', '菜包/豆沙包', '80-100g/只', '3.5-4.5元/只', '基础流量品', '5%-8%'],
    ['其他', '春卷/饺子等', '多种规格', '10-20元/袋', '品类补充', '5%-10%'],
]
add_table_from_data(doc, prod_headers, prod_rows)
doc.add_paragraph()

add_heading(doc, '四级产品拆解', level=2)
add_heading(doc, '一、销量最好：开口手作大肉包', level=3)
add_body_paragraph(doc,
    '开口大肉包是蒸笼头绝对的旗舰产品，120g的超大规格和46%的馅芯比均为行业最高水平之一。'
    '产品设计核心优势在于开口造型——消费者可以直接看到饱满的肉馅从开口处溢出，'
    '视觉化地传递真材实料的价值主张，无需额外的广告语。'
    '这款单品估计占总营收30%-35%，是蒸笼头的品类流量担当。')

add_heading(doc, '二、口碑声量最多：开口大肉包', level=3)
add_body_paragraph(doc,
    '小红书上关于蒸笼头的自发热帖中，80%以上在讨论大肉包——料足皮薄馅大是出现频率最高的关键词。'
    '消费者在社交媒体上为蒸笼头提供的免费口碑传播几乎全部围绕大肉包展开。')

add_heading(doc, '三、品类代表性最强：叮叮包', level=3)
add_body_paragraph(doc,
    '叮叮包是可微波加热40秒即食的技术创新产品，解决了冷鲜面点食用便利性的最大痛点——'
    '传统面点需要蒸锅或微波炉+湿纸等复杂操作，让很多潜在消费者望而却步。'
    '叮叮包让面点加热简化为撕开口、放进微波炉、等40秒、开始吃，'
    '这对手工面点从家庭烹饪场景延伸到办公室/宿舍/独居场景有决定性的品类扩展意义。')

add_heading(doc, '四、品牌特色代表：奶酪馒头', level=3)
add_body_paragraph(doc,
    '奶酪馒头是品类跨界创新——中式馒头+西式奶酪，展现了蒸笼头的产品创新能力和品类突破意识。'
    '在面点同质化严重的行业中，奶酪馒头的跨品类定位为蒸笼头创造了差异化的品牌记忆点。')

add_heading(doc, '产品矩阵诊断', level=2)
add_body_paragraph(doc,
    '蒸笼头产品矩阵的核心优势是爆款深度而非品类宽度。四款单品合计贡献约70%-80%营收，大肉包一个单品就占30%以上——'
    '这在食品行业中属于罕见的高集中度，既是优势也是风险。'
    '品类宽度不足的关键短板在于缺乏搭配品类——没有粥、汤、饮品等关联产品，客单价被限制在15-25元/次。'
    '产品创新的节奏保持得不错——平均约2年推出一款具有品类差异化潜力的新品，'
    '叮叮包和奶酪馒头代表了两个不同的创新方向，说明团队有持续的产品创新能力。')

add_separator(doc)

# --- Section 3: 渠道与供应链 ---
add_heading(doc, '渠道与供应链', level=1)
add_heading(doc, '渠道结构', level=2)

add_heading(doc, '盒马鲜生：绝对核心渠道', level=3)
add_body_paragraph(doc,
    '盒马是蒸笼头最核心的销售渠道，营收占比估计超过70%。蒸笼头是盒马冷鲜面点品类的唯一战略合作商，'
    '在每家盒马门店的面点冷藏柜中拥有专属排面。这种排他性的渠道资源为蒸笼头提供了稳定的出货量和品牌曝光，'
    '但也形成了对单一渠道的极高依赖度。')

add_body_paragraph(doc,
    '盒马全国约300+门店，蒸笼头在每家门店日均面点销售额估计在2000-5000元（终端零售口径），'
    '300家门店年终端零售额约2.2-5.5亿元。'
    '蒸笼头出厂口径只有这个数字的50%-60%（扣除盒马抽成和冷链物流成本），即1.1-3.3亿元。'
    '蒸笼头当前年营收1.2-1.5亿元正好处在这个通道的中段。')

add_heading(doc, '麦德龙：第二渠道', level=3)
add_body_paragraph(doc,
    '麦德龙是中国领先的会员制仓储超市，其B端企业客户和家庭会员的双轨制模式与蒸笼头的大包装面点产品天然契合。'
    '蒸笼头供应麦德龙的营收占比估计在8%-12%，以企业团购和家庭量贩装为主。')

add_heading(doc, '大润发：传统KA补充', level=3)
add_body_paragraph(doc,
    '大润发作为高鑫零售旗下的传统KA卖场，覆盖二三线城市和中老年家庭消费者。'
    '蒸笼头供应大润发的营收占比估计在5%-8%，以常规包装面点（包子+馒头+花卷组合）为主。')

add_heading(doc, '罗森便利店：新场景试验', level=3)
add_body_paragraph(doc,
    '罗森是中国日系便利店三巨头之一，全国门店约2000+家。'
    '蒸笼头供应罗森的营收占比估计在3%-5%，产品以单只独立包装的鲜食面点为主，面向便利店即食场景。'
    '罗森渠道是蒸笼头拓展新消费场景的实验田，如果验证成功，全家和7-Eleven是潜在的扩展目标。')

add_heading(doc, '渠道集中度汇总', level=2)

chan_headers = ['渠道', '营收占比（推算）', '类型', '核心特征']
chan_rows = [
    ['盒马鲜生', '70%+', '新零售会员店', '独家战略合作，冷鲜面点唯一供应商'],
    ['麦德龙', '8%-12%', '仓储会员超市', '企业团购+家庭量贩装'],
    ['大润发', '5%-8%', '传统KA卖场', '二三线城市覆盖，价格敏感型客群'],
    ['罗森', '3%-5%', '便利店', '即食场景试验田，单只独立包装'],
    ['其他', '~5%', '零星渠道', '社区生鲜店等零散供应'],
]
add_table_from_data(doc, chan_headers, chan_rows)
doc.add_paragraph()

add_heading(doc, '渠道空白诊断', level=2)
add_body_paragraph(doc,
    '蒸笼头有三个关键的渠道空白。第一，没有山姆/Costco——这是最致命的缺失，'
    '因为山姆的大包装+品质定位与蒸笼头的产品形态完美匹配。'
    '第二，没有线上渠道——蒸笼头目前没有天猫旗舰店、没有京东自营、没有抖音电商，完全依赖线下渠道出货。'
    '第三，没有经销商网络——蒸笼头完全以直供模式运作，没有建立区域经销商体系，'
    '这意味着渠道扩展的边际成本非常高。')

add_heading(doc, '供应链', level=2)
add_heading(doc, '工厂布局', level=3)
add_body_paragraph(doc,
    '蒸笼头在松江拥有2家自有工厂，总面积超过2万平方米。'
    '工厂生产中式面点的全品类——包子、烧麦、馒头、花卷、饺子、春卷，日产能约30万例中式面点。'
    '按全年开工300天计，年产能约9000万例，当前产能利用率约70%-80%，年实际产量约6300-7200万例。')

add_heading(doc, '品控认证体系', level=3)
add_body_paragraph(doc,
    '蒸笼头持有HACCP危害分析与关键控制点认证、ISO质量管理体系认证、'
    'BRC全球食品安全标准认证三重食品安全管理体系，'
    '这个认证组合在中国中小型面点企业中属于较高水平。')

add_heading(doc, '供应链核心竞争力', level=3)
add_body_paragraph(doc,
    '蒸笼头供应链的三大核心竞争力：第一，365天无间断供货能力——'
    '春节、国庆等节假日期间盒马仍然能稳定拿到蒸笼头的产品，'
    '这个供应稳定性是盒马选择蒸笼头作为唯一冷鲜面点合作商的核心原因之一。'
    '第二，手工+短保的品质壁垒——46%馅芯比的包子必须手工包制才能保持面皮不破裂，'
    '机器无法做到这个比例。第三，原料供应链整合——猪肉馅采用松林每天配送的鲜猪肉，'
    '青菜选用松江本地基地生产的鲜菜，实现了肉+菜的本地化新鲜供应闭环。')

add_heading(doc, '供应链瓶颈', level=3)
add_body_paragraph(doc,
    '手工产能的上限是蒸笼头供应链的最大瓶颈。从1亿只扩展到2亿只的边际成本较低'
    '（只需增加人工班次和部分自动化设备，增量投资约1000-1500万），'
    '但从2亿只扩展到3-4亿只需要增加自动化产线或新建三厂，投资门槛将跳升至3000-5000万。'
    '手工标签与产能扩张之间存在内在矛盾——越是强调手工，规模化扩张的天花板越低。')

add_separator(doc)

# Now move all newly added elements to before insert_before
# The elements we just added are after all existing content (added at the end).
# We need to move them to the insertion point.

print(f"Moving inserted elements to before paragraph {first_body_idx}...")
body = doc.element.body
# Get all elements that were added (after the original last element)
# Find the index of insert_before in body
insert_before_idx = list(body).index(insert_before)

# Count total elements and find our new ones
all_elements = list(body)
total_elements = len(all_elements)

# The new elements are from insert_before_idx to end (well, after the original content)
# Actually, we need to figure out the original element count before we added new ones.
# Since we added everything at the end, the new elements start at the end of original content.
# But we don't have that count. Let's try a different approach.

# Actually, python-docx adds new elements to the body during add_paragraph/add_table.
# The issue is: our new paragraphs were added AFTER all existing content.
# We need to find which elements are "new" and move them before insert_before.

# Since we can't easily distinguish new from old, let's use a heuristic:
# Count the original elements first (before our inserts), then move the rest.
# But we already added them...

# Alternative approach: find our new elements by content (they have specific text)
# and move them. Let's find the heading "品牌概览" which is the first one we added.

brand_overview_heading = None
for elem in body:
    for p in elem.iter():
        if p.tag.endswith('}p'):
            text = ''.join(p.itertext()).strip()
            if '品牌概览' in text:
                # Check if this is a heading
                pPr = p.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None and 'Heading' in (pStyle.get(qn('w:val')) or ''):
                        brand_overview_heading = elem
                        break
    if brand_overview_heading is not None:
        break

if brand_overview_heading is not None:
    # Find the start and end of our new content block
    new_start_idx = None
    for i, elem in enumerate(all_elements):
        if elem is brand_overview_heading:
            new_start_idx = i
            break
    
    if new_start_idx is not None:
        # Move everything from new_start_idx to end to just before insert_before_idx
        new_elements = list(body)[new_start_idx:]
        for elem in new_elements:
            body.remove(elem)
        
        # Recalculate insert_before_idx (it may have shifted)
        insert_before_idx = list(body).index(insert_before)
        
        # Insert new elements at insert_before_idx
        for elem in new_elements:
            body.insert(insert_before_idx, elem)
            insert_before_idx += 1
        
        print(f"  Moved {len(new_elements)} new elements before chapter 1.")

# ==========================================
# PHASE 3: Search & Replace - 斤价
# ==========================================
print("Replacing 斤价 references...")
replacements = 0
for para in doc.paragraphs:
    full_text = para.text
    if '斤价' in full_text:
        # We need to modify runs
        for run in para.runs:
            if '斤价' in run.text:
                run.text = run.text.replace('斤价', '单件客单价')
                replacements += 1
            if '斤价对比' in run.text:
                run.text = run.text.replace('斤价对比', '单件客单价对比')
                replacements += 1
print(f"  Made {replacements} replacements.")

# Also check tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if '斤价' in run.text:
                        run.text = run.text.replace('斤价', '单件客单价')

# ==========================================
# PHASE 4: Remove 钟薛高
# ==========================================
print("Removing 钟薛高 references...")
zhong_removed = 0
paras_to_remove = []
for para in doc.paragraphs:
    if '钟薛高' in para.text:
        # Check if entire paragraph should be removed or just modified
        full_text = para.text.strip()
        # If it's a heading containing only 钟薛高, remove the whole paragraph
        if para.style.name.startswith('Heading') and '钟薛高' in full_text and len(full_text) < 20:
            paras_to_remove.append(para)
        # For body paragraphs, try to replace text in runs
        else:
            for run in para.runs:
                if '钟薛高' in run.text:
                    # Replace 钟薛高 references with general advice
                    run.text = run.text.replace('钟薛高（跨界参考）', '')
                    run.text = run.text.replace('钟薛高', '')
                    zhong_removed += 1

print(f"  Modified {zhong_removed} 钟薛高 references.")
print(f"  {len(paras_to_remove)} paragraphs to remove.")

# Remove identified paragraphs
for para in paras_to_remove:
    p_element = para._element
    p_element.getparent().remove(p_element)

# Also clean up tables
for table in doc.tables:
    rows_to_remove = []
    for row in table.rows:
        for cell in row.cells:
            if '钟薛高' in cell.text:
                rows_to_remove.append(row)
                break
    for row in rows_to_remove:
        row._element.getparent().remove(row._element)

# ==========================================
# PHASE 5: Save
# ==========================================
print(f"Saving to {OUTPUT_DOCX}...")
doc.save(OUTPUT_DOCX)
size_kb = os.path.getsize(OUTPUT_DOCX) / 1024
print(f"Done! File size: {size_kb:.1f} KB")

# Quick word count check
total_text = ' '.join([p.text for p in doc.paragraphs])
word_count = len(total_text.replace(' ', ''))
print(f"Approximate character count: {word_count:,}")
