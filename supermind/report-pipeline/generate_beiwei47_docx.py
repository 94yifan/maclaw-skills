#!/usr/bin/env python3
"""Generate DOCX for 北纬47度 report with cover page, bold, tables, charts."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_MD = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/content/report_final_beiwei47.md"
CHARTS_DIR = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/charts"
OUTPUT_DOCX = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/beiwei47_final.docx"

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# --- COVER PAGE ---
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('北纬47度')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('品牌竞品研究咨询报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

line_p = doc.add_paragraph()
line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line_p.add_run('━' * 30)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for text in ['燃创咨询 BreaC Lab', '五维模型 V5 | Schema v1.3', '2026年7月18日']:
    run = meta_p.add_run(text + '\n')
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Page break after cover
doc.add_page_break()

# --- BODY ---
with open(INPUT_MD, 'r', encoding='utf-8') as f:
    lines = f.readlines()

chart_map = {
    '../charts/tmall_top_sellers_beiwei47.png': 'tmall_top_sellers_beiwei47.png',
    '../charts/price_range_beiwei47.png': 'price_range_beiwei47.png',
    '../charts/store_compare_beiwei47.png': 'store_compare_beiwei47.png',
    '../charts/category_coverage_beiwei47.png': 'category_coverage_beiwei47.png',
}

def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc, text, level=0, is_bold=False):
    """Add paragraph with bold handling."""
    p = doc.add_paragraph()
    if level > 0:
        p.style = doc.styles[f'Heading {level}']
    
    # Handle **bold** patterns
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(11)
        else:
            run = p.add_run(part)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(11)
    return p

def add_table_from_lines(doc, table_lines):
    """Convert markdown table lines to docx table with formatting."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(c.replace('-', '').replace(':', '').replace(' ', '') == '' for c in cells):
            continue  # skip separator line
        rows.append(cells)
    
    if not rows:
        return None
    
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold in cells
            parts = re.split(r'(\*\*.*?\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # Header row: gray background
            if i == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                for run in p.runs:
                    run.bold = True
            
            # Borders
            for edge in ['top', 'left', 'bottom', 'right']:
                set_cell_border(cell, **{edge: {'val': 'single', 'sz': 4, 'color': '808080'}})
    
    return table

def insert_image(doc, chart_filename):
    """Insert chart image centered."""
    img_path = os.path.join(CHARTS_DIR, chart_filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        doc.add_paragraph()  # spacing

# Parse and build document
i = 0
table_buffer = []
in_table = False
in_code_block = False

while i < len(lines):
    line = lines[i].rstrip()
    
    # Skip code blocks
    if line.startswith('```'):
        in_code_block = not in_code_block
        i += 1
        continue
    
    if in_code_block:
        i += 1
        continue
    
    # Handle tables
    if line.startswith('|') and not in_table:
        in_table = True
        table_buffer = [line]
        i += 1
        continue
    
    if in_table:
        if line.startswith('|'):
            table_buffer.append(line)
            i += 1
            # Check if next line is still table or end
            if i < len(lines) and lines[i].strip().startswith('|'):
                continue
            else:
                add_table_from_lines(doc, table_buffer)
                doc.add_paragraph()  # spacing
                table_buffer = []
                in_table = False
                continue
        else:
            add_table_from_lines(doc, table_buffer)
            doc.add_paragraph()
            table_buffer = []
            in_table = False
            # fall through to process current line
    
    # Handle images
    img_match = re.match(r'!\[.*\]\((.*)\)', line)
    if img_match:
        img_path = img_match.group(1)
        chart_file = chart_map.get(img_path)
        if chart_file:
            insert_image(doc, chart_file)
        i += 1
        continue
    
    # Handle headings
    if line.startswith('# '):
        add_formatted_paragraph(doc, line[2:], level=0)
        i += 1
        continue
    elif line.startswith('## '):
        add_formatted_paragraph(doc, line[3:], level=1)
        i += 1
        continue
    elif line.startswith('### '):
        add_formatted_paragraph(doc, line[4:], level=2)
        i += 1
        continue
    elif line.startswith('#### '):
        add_formatted_paragraph(doc, line[5:], level=3)
        i += 1
        continue
    
    # Handle horizontal rules
    if line.strip() == '---':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 40)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        i += 1
        continue
    
    # Handle quote blocks
    if line.startswith('> '):
        text = line[2:]
        p = add_formatted_paragraph(doc, text)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.italic = True
        i += 1
        continue
    
    # Empty line
    if not line.strip():
        i += 1
        continue
    
    # Regular paragraph
    add_formatted_paragraph(doc, line)
    i += 1

# Set heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.color.rgb = RGBColor(0x2d, 0x5f, 0x8a)
    elif level == 3:
        hs.font.size = Pt(12)
        hs.font.color.rgb = RGBColor(0x3a, 0x7c, 0xa5)

# Save
os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
doc.save(OUTPUT_DOCX)
print(f"DOCX saved: {OUTPUT_DOCX}")
print(f"File size: {os.path.getsize(OUTPUT_DOCX)} bytes ({os.path.getsize(OUTPUT_DOCX)/1024:.1f} KB)")
