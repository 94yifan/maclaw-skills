#!/usr/bin/env python3
"""Generate libernovo report docx from markdown chapters and chart PNGs."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

BASE = "/Users/yifansmacmini/.openclaw/workspace/supermind/output"
CONTENT = os.path.join(BASE, "content")
CHARTS = os.path.join(BASE, "charts")
OUTPUT = os.path.join(BASE, "libernovo_report_v1.docx")

doc = Document()

# --- style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

# --- helper ---
def add_md_content(doc, md_path):
    """Read markdown and add to doc preserving headings and paragraphs."""
    if not os.path.exists(md_path):
        doc.add_paragraph(f"[Content not found: {md_path}]")
        return
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_table = False
    table_rows = []
    
    for line in lines:
        stripped = line.rstrip()
        
        # Handle tables
        if stripped.startswith('|') and stripped.endswith('|'):
            if '---' in stripped:
                in_table = True
                continue
            if in_table:
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                # Render collected table
                col_count = max(len(r) for r in table_rows)
                # Header row
                table = doc.add_table(rows=1, cols=col_count, style='Light Grid Accent 1')
                hdr = table.rows[0].cells
                for j, cell_text in enumerate(table_rows[0]):
                    if j < col_count:
                        hdr[j].text = cell_text
                        for p in hdr[j].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)
                                run.font.bold = True
                # Data rows
                for row_data in table_rows[1:]:
                    row = table.add_row()
                    for j, cell_text in enumerate(row_data):
                        if j < col_count:
                            row.cells[j].text = cell_text
                            for p in row.cells[j].paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(9)
                doc.add_paragraph()  # spacer
                table_rows = []
                in_table = False
            
            # Headings
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith('---'):
                doc.add_paragraph('_' * 60)
            elif stripped.strip():
                p = doc.add_paragraph(stripped)
                # Bold **text** in paragraphs
                if '**' in stripped:
                    p.clear()
                    parts = re.split(r'(\*\*.*?\*\*)', stripped)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)

def add_chart(doc, chart_path, caption):
    """Add chart image with caption."""
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(f"[Chart not found: {chart_path}]")

# ====== TITLE PAGE ======
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("清闲 LiberNovo")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("品牌竞品研究报告")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x3A, 0x5F, 0x8A)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("人体工学椅行业 | 2026年7月").font.size = Pt(12)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta2.add_run("燃创品牌咨询 | 逸凡品牌咨询体系").font.size = Pt(11)

doc.add_page_break()

# ====== CHAPTERS ======
chapters = [
    ("ch1_core_findings.md", "第一章：核心发现与咨询窗口"),
    ("ch2_industry.md", "第二章：行业格局与竞品总矩阵"),
]

# Add Ch1 and Ch2
for filename, _ in chapters:
    path = os.path.join(CONTENT, filename)
    add_md_content(doc, path)
    doc.add_page_break()

# Add Ch3 with charts embedded before 3.6
path = os.path.join(CONTENT, "ch3_competitive.md")
if os.path.exists(path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_section = ""
    for line in lines:
        stripped = line.rstrip()
        
        # Detect when we reach 3.6 竞争模式归纳
        if stripped.startswith('## 3.6'):
            # Insert charts before 3.6
            doc.add_paragraph()
            doc.add_heading("图表：跨品牌数据总览", level=2)
            
            add_chart(doc, os.path.join(CHARTS, "chart_tmall_sales.png"),
                      "图1：天猫旗舰店爆款月销量对比（综合推断）")
            doc.add_paragraph()
            
            add_chart(doc, os.path.join(CHARTS, "chart_jd_sales.png"),
                      "图2：京东自营爆款累计销量对比（综合推断）")
            doc.add_paragraph()
            
            add_chart(doc, os.path.join(CHARTS, "chart_price.png"),
                      "图3：各品牌核心产品单件价对比（元/把）")
            doc.add_paragraph()
            
            add_chart(doc, os.path.join(CHARTS, "chart_repurchase.png"),
                      "图4：各品牌回头客/复购率对比（综合推断）")
            doc.add_paragraph()
            
            doc.add_paragraph("注：以上图表数据为基于行业研报和竞争分析的综合推断（△），非精确采集数据。天猫/京东店铺页面因反爬策略未直接获取销量和回头客等详细数据。")
            doc.add_page_break()
        
        # Render the line
        if stripped.startswith('# ') and not stripped.startswith('####'):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('---'):
            doc.add_paragraph('_' * 40)
        elif stripped.strip():
            p = doc.add_paragraph(stripped)
            if '**' in stripped:
                p.clear()
                parts = re.split(r'(\*\*.*?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

# Add Ch4, Ch5, Ch6
remaining = [
    ("ch4_deep_analysis.md", "第四章：本品深度分析"),
    ("ch5_gap.md", "第五章：本竞品差距对比"),
    ("ch6_recommendations.md", "第六章：咨询切入点与策略建议"),
]
for filename, _ in remaining:
    doc.add_page_break()
    path = os.path.join(CONTENT, filename)
    add_md_content(doc, path)

# ====== APPENDIX ======
doc.add_page_break()
doc.add_heading("附录：全报告数据核验表", level=1)

verification_data = [
    ("ch1", "永艺股份2025年营收", "48.82亿元", "新浪财经/广发证券研报", "✓ 已验证"),
    ("ch1", "永艺股份2025年归母净利润", "2.33亿元", "新浪财经/信达证券研报", "✓ 已验证"),
    ("ch1", "永艺内销线上收入增速", "+23%", "券商研报", "✓ 已验证"),
    ("ch1", "永艺升降桌收入增速", "+71%", "券商研报", "✓ 已验证"),
    ("ch2", "永艺外销占比", "74%", "2025年报", "✓ 已验证"),
    ("ch2", "永艺2026Q1营收", "11.42亿元", "新浪财经", "✓ 已验证"),
    ("ch3", "各品牌天猫爆款销量", "综合推断", "天猫+行业定位", "△ 综合评估"),
    ("ch3", "各品牌京东累计销量", "综合推断", "京东+行业定位", "△ 综合评估"),
    ("ch3", "各品牌核心产品单价", "综合推断", "天猫+行业定位", "△ 综合评估"),
    ("ch3", "各品牌回头客率", "综合推断", "天猫+行业经验", "△ 综合评估"),
    ("ch4", "清闲libernovo产品配置", "官网获取", "libernovo.com", "✓ 已验证"),
    ("ch5", "各品牌竞争定位", "综合分析", "多源综合", "△ 综合评估"),
]

table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
hdr = table.rows[0].cells
for j, text in enumerate(["章节", "数据点", "数值", "来源", "验证状态"]):
    hdr[j].text = text
    for p in hdr[j].paragraphs:
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.bold = True

for row_data in verification_data:
    row = table.add_row()
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        for p in row.cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Limitations
doc.add_heading("局限性声明", level=2)
limitations = [
    "电商数据（天猫/京东）因反爬策略阻挡未直接采集，销量和价格数据为基于行业定位的综合推断（△），非精确值",
    "财报数据截至2026年7月公开的最新报告（2025年报+2026Q1季报）",
    "社交数据（微博/小红书/抖音）因平台限制未系统采集",
    "品牌分析基于公开信息的综合判断，不代表品牌实际战略",
    "非上市品牌（清闲libernovo/西昊/黑白调/保友）的财务数据为商业机密，无公开来源",
]
for lim in limitations:
    doc.add_paragraph(lim, style='List Bullet')

# Sources
doc.add_heading("来源文献列表", level=2)
sources = [
    "[1] 广发证券, 永艺股份(603600)研报: 外销复苏&品类延伸 盈利修复加速, 2026-05-13",
    "[2] 信达证券, 永艺股份(603600)研报: 向内优化管理 向外拓展增长机遇, 2026-05-06",
    "[3] 华安证券, 永艺股份(603600)研报: 26Q1业绩向好 自主品牌建设卓有成效, 2026-04-28",
    "[4] 新浪财经, 永艺股份(603600)个股页面, https://finance.sina.com.cn/realstock/company/sh603600/nc.shtml, 2026-07-15",
    "[5] 永艺股份2025年年报 (via 新浪财经/券商研报摘要), 2026-04",
    "[6] 清闲LiberNovo官方网站, https://libernovo.com, 2026-07-15",
    "[7] 天风证券/浙商证券/华福证券/长江证券/华泰证券, 永艺股份多份研报, 2025-2026",
]
for src in sources:
    doc.add_paragraph(src, style='List Bullet')

# Save
doc.save(OUTPUT)
print(f"Report saved to: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT)} bytes")
