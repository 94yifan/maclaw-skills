#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将蒸馏 markdown 合并转换为 docx，%%H%%...%%/H%% 转为黄色高亮，**bold** 转为加粗。"""
import re, sys, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn

HL_START = "%%H%%"
HL_END = "%%/H%%"

def set_font(run, size=10.5, bold=False, color=None, name="微软雅黑"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_rich_paragraph(doc, text, style=None, size=10.5, bold_all=False, color=None,
                       space_before=0, space_after=6, align=None, indent=None):
    """支持 %%H%% 高亮与 **bold** 的富文本段落。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent is not None:
        pf.left_indent = Cm(indent)
    # 解析高亮块
    parts = []
    pos = 0
    while True:
        s = text.find(HL_START, pos)
        if s == -1:
            parts.append((text[pos:], False))
            break
        e = text.find(HL_END, s + len(HL_START))
        if e == -1:
            parts.append((text[pos:], False))
            break
        if s > pos:
            parts.append((text[pos:s], False))
        parts.append((text[s+len(HL_START):e], True))
        pos = e + len(HL_END)
    for seg, is_hl in parts:
        # 处理 **bold**
        seg_parts = re.split(r'(\*\*[^*]+\*\*)', seg)
        for sp in seg_parts:
            if not sp:
                continue
            if sp.startswith('**') and sp.endswith('**'):
                run = p.add_run(sp[2:-2])
                set_font(run, size=size, bold=True, color=color)
            else:
                run = p.add_run(sp)
                set_font(run, size=size, bold=bold_all, color=color)
            if is_hl:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(8)
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: (0x1F, 0x3B, 0x73), 2: (0x2E, 0x5A, 0xA8), 3: (0x3B, 0x6E, 0xBF)}
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level, (0,0,0)))
    return p

def parse_md(md_text, doc):
    lines = md_text.split('\n')
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith('```'):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            add_rich_paragraph(doc, line, size=9, space_after=2)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue
        m = re.match(r'^[-*]\s+(.*)', line)
        if m:
            add_rich_paragraph(doc, m.group(1), size=10.5, indent=0.5, space_after=3)
            i += 1
            continue
        m = re.match(r'^\d+[.、]\s+(.*)', line)
        if m:
            add_rich_paragraph(doc, m.group(1), size=10.5, indent=0.5, space_after=3)
            i += 1
            continue
        m = re.match(r'^>\s?(.*)', line)
        if m:
            add_rich_paragraph(doc, m.group(1), size=10, color=(0x55,0x55,0x55), indent=0.5, space_after=3)
            i += 1
            continue
        # 普通段落（可能跨多行，简单处理为一段）
        add_rich_paragraph(doc, line, size=10.5, space_after=6)
        i += 1

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    distilled_dir = os.path.join(base, 'distilled')
    files = [f for f in sorted(os.listdir(distilled_dir)) if f.endswith('.md')]
    if not files:
        print('NO_MD_FILES'); sys.exit(1)
    doc = Document()
    # 页面设置 A4
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    total_chars = 0
    for fname in files:
        path = os.path.join(distilled_dir, fname)
        md_text = open(path, encoding='utf-8').read()
        total_chars += len(re.sub(r'\s', '', re.sub(r'%%/??H%%', '', md_text)))
        parse_md(md_text, doc)
    out = os.path.join(base, 'AI-Agents-核心解读.docx')
    doc.save(out)
    print(f'DONE chars={total_chars} out={out}')

if __name__ == '__main__':
    main()
