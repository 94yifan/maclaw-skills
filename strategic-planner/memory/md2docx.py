#!/usr/bin/env python3
"""Convert a Markdown file (yangst-review-merged.md) to DOCX."""

import re
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT = "/Users/yifansmacmini/.openclaw/workspace/strategic-planner/memory/yangst-review-merged.md"
OUTPUT = "/Users/yifansmacmini/.openclaw/workspace/strategic-planner/memory/yangst-review-merged.docx"

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Heading styles
for level in range(1, 4):
    hstyle = doc.styles[f'Heading {level}']
    hfont = hstyle.font
    hfont.name = 'Microsoft YaHei'
    hstyle.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level == 1:
        hfont.size = Pt(18)
        hfont.bold = True
    elif level == 2:
        hfont.size = Pt(14)
        hfont.bold = True
    elif level == 3:
        hfont.size = Pt(12)
        hfont.bold = True

def parse_inline_bold(text):
    """Parse **bold** markers in inline text. Returns list of (text, is_bold) tuples."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    result = []
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            result.append((part[2:-2], True))
        elif part:
            result.append((part, False))
    return result

def add_rich_paragraph(para, text):
    """Add runs to a paragraph, handling **bold** markers."""
    segments = parse_inline_bold(text)
    for segment_text, is_bold in segments:
        run = para.add_run(segment_text)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if is_bold:
            run.bold = True

def set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def render_table(doc, table_alignments, table_rows_raw):
    """Render a markdown table into a DOCX table."""
    # Determine number of columns from first row
    first_row_cells = [c.strip() for c in table_rows_raw[0].split('|') if c.strip()]
    num_cols = len(first_row_cells)

    # Data rows (skip separator row)
    data_rows_raw = [r for r in table_rows_raw if not re.match(r'^[\s:|:\-]+$', r.strip('|').strip())]
    num_data_rows = len(data_rows_raw)
    if num_data_rows == 0:
        return

    # Build table
    tbl = doc.add_table(rows=num_data_rows, cols=num_cols)
    tbl.style = 'Table Grid'

    # Determine column widths proportional to their visual weight
    max_content = [0] * num_cols
    for row_data in data_rows_raw:
        cells = [c.strip() for c in row_data.split('|') if c.strip()]
        for i, c in enumerate(cells):
            if i < num_cols:
                max_content[i] = max(max_content[i], len(c))

    total = sum(max_content) or num_cols
    col_widths = [max(Inches(0.7), Inches(1.8 * (w / total))) if total > 0 else Inches(1.5) for w in max_content]

    # Column header from alignment row if present:
    # Detect separator row to know which row has alignments
    sep_row_idx = -1
    for i, r in enumerate(table_rows_raw):
        stripped = r.strip()
        if re.match(r'^[\|:\s\-]+$', stripped):
            sep_row_idx = i
            break

    for row_idx, row_data in enumerate(data_rows_raw):
        cells = [c.strip() for c in row_data.split('|') if c.strip()]
        for col_idx, cell_text in enumerate(cells):
            if col_idx >= num_cols:
                continue
            cell = tbl.cell(row_idx, col_idx)

            # Determine alignment from alignment row
            align = WD_ALIGN_PARAGRAPH.LEFT
            if sep_row_idx != -1 and row_idx == 0:
                # Header row
                pass
            elif sep_row_idx != -1 and row_idx > 0:
                # Check alignment from separator row for this column
                sep_parts = table_rows_raw[sep_row_idx].strip().split('|')
                if col_idx + 1 < len(sep_parts):
                    sep_seg = sep_parts[col_idx + 1].strip()
                    if sep_seg.startswith(':') and sep_seg.endswith(':'):
                        align = WD_ALIGN_PARAGRAPH.CENTER
                    elif sep_seg.endswith(':'):
                        align = WD_ALIGN_PARAGRAPH.RIGHT

            # Clear default paragraph
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = align

            # Parse **bold** in cell
            add_rich_paragraph(p, cell_text)

            # Header row formatting
            if sep_row_idx != -1 and row_idx == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "D9E2F3")

        # Set column widths
        for col_idx in range(num_cols):
            cell = tbl.cell(row_idx, col_idx)
            cell.width = col_widths[col_idx]

    doc.add_paragraph('')  # spacer


with open(INPUT, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Parse line by line, building table blocks
in_table = False
table_rows_raw = []
table_alignments = []
i = 0

while i < len(lines):
    line = lines[i]
    stripped = line.rstrip('\n').rstrip('\r')

    # Detect table start: line starts with |
    if stripped.startswith('|') and stripped.endswith('|') and '|' in stripped[1:-1]:
        table_rows_raw = [stripped]
        table_alignments = []
        i += 1
        # Collect all consecutive table lines
        while i < len(lines):
            next_line = lines[i].rstrip('\n').rstrip('\r')
            if next_line.startswith('|') and next_line.endswith('|') and '|' in next_line[1:-1]:
                table_rows_raw.append(next_line)
                i += 1
            else:
                break
        # Now render the table
        render_table(doc, table_alignments, table_rows_raw)
        in_table = False
        table_rows_raw = []
        continue

    # Heading detection (must be at start of line)
    heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2)
        para = doc.add_heading(text, level=level)
        # Ensure runs use the right font
        for run in para.runs:
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        i += 1
        continue

    # Horizontal rule
    if re.match(r'^-{3,}\s*$', stripped):
        doc.add_page_break()
        i += 1
        continue

    # Empty line
    if stripped == '':
        i += 1
        continue

    # Normal paragraph, might contain **bold**
    para = doc.add_paragraph()
    add_rich_paragraph(para, stripped)
    i += 1

doc.save(OUTPUT)
print(f"Saved to {OUTPUT}")

import os
size = os.path.getsize(OUTPUT)
print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")
