#!/usr/bin/env python3
"""
Generate 蓝氏Y26H1小红书奶盾线投放复盘报告 DOCX
"""
import os, re, copy
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image as PILImage

IMG_DIR = "/Users/yifansmacmini/.openclaw/workspace/strategic-planner/memory/lanshi_images"
OUTPUT = "/Users/yifansmacmini/.openclaw/workspace/strategic-planner/memory/lanshi-review-v3.docx"
MAX_IMG_WIDTH_CM = 14

doc = Document()

# ── Global style tweaks ──────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Set default paragraph spacing
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=None, align=None, space_after=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_rich_para(segments, align=None, space_after=None):
    """segments: list of (text, bold, color_rgb_tuple_or_None, size_or_None)"""
    p = doc.add_paragraph()
    for seg in segments:
        text, bold = seg[0], seg[1] if len(seg) > 1 else False
        color = seg[2] if len(seg) > 2 else None
        size = seg[3] if len(seg) > 3 else None
        run = p.add_run(text)
        run.bold = bold
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if color:
            run.font.color.rgb = RGBColor(*color)
        if size:
            run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def set_cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    p.alignment = align
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cells(row, color_hex):
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

def add_table_from_data(headers, rows, col_widths=None):
    """Build a formatted table from headers and rows data"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, size=9)
    shade_cells(hdr, "2B579A")
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            set_cell_text(row.cells[ci], val, size=8.5)
        if ri % 2 == 1:
            shade_cells(row, "F2F7FB")
    
    return table

def add_image_with_caption(img_path, caption):
    """Add an image fitted to MAX_IMG_WIDTH_CM width with a caption below"""
    if not os.path.exists(img_path):
        add_para(f"[图片缺失: {os.path.basename(img_path)}]", size=9, color=(200,0,0))
        return
    
    try:
        pil_img = PILImage.open(img_path)
        orig_w, orig_h = pil_img.size
    except:
        add_para(f"[图片无法读取: {os.path.basename(img_path)}]", size=9, color=(200,0,0))
        return
    
    # Calculate scaled dimensions
    target_w_cm = min(MAX_IMG_WIDTH_CM, orig_w * 0.026458)  # rough px→cm
    scale = target_w_cm / (orig_w * 0.026458)
    target_h_cm = orig_h * 0.026458 * scale
    
    # Add image centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(target_w_cm))
    
    # Add caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(100, 100, 100)
    cr.font.name = '微软雅黑'
    cr._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)


# ──── CHAPTER 1: 投放目标与策略回顾 ───────────────────────────
add_heading_styled("一、投放目标与策略回顾（G-Goal）", level=1)

add_heading_styled("1.1 H1投放目标", level=2)
add_para("品类心智占位 + 奶盾线生意转化双目标。蓝氏Y26H1小红书投放承接\"超能奶盾\"产品线年度战役，核心目标是：①在猫主粮赛道建立烘焙冻干品类强心智；②通过SEM投放直接拉动奶盾线（成猫+幼猫）站外电商转化。")

add_heading_styled("1.2 KPI三层分级", level=2)
add_table_from_data(
    ["层级", "指标", "H1目标", "实际", "达成"],
    [
        ["P0 北极星", "进店成本", "≤猫主粮大盘", "2.57", "✓ 达成"],
        ["P0 北极星", "SPU AIPS排名", "烘焙冻干赛道TOP3", "#1（5.31及6.30）", "✓ 达成"],
        ["P1 核心过程", "CPTI", "≤2.5", "2.32", "✓ 达成"],
        ["P1 核心过程", "品牌词SOV", "保持首位", "保持首位", "✓ 达成"],
        ["P2 参考", "CPE", "≤5.0", "4.30", "✓ 达成"],
        ["P2 参考", "CTR", "≥10%", "15.71%", "✓ 达成"],
    ]
)

add_heading_styled("1.3 提案Deck对照：预设验证框架", level=2)
add_table_from_data(
    ["提案预设", "复盘验证点", "验证结果"],
    [
        ["品类心智占位：烘焙冻干赛道TOP3", "SPU AIPS人群排名", "✓ 达成：全阶段猫鸟乳鸽系列稳居行业#1"],
        ["奶盾线生意转化：进店成本≤大盘均值", "进店成本 vs 行业benchmark", "✓ 达成：进店成本2.57"],
        ["信息流点站为核心降本模式", "CPC/CPUV/CPI vs 25D11", "✓ 达成：信息流CPC↓65%，CPUV↓75%，CPI↓52%"],
        ["SEM选词策略（吸黏拽三层）", "实际消耗配比 vs 预设", "见第三章搜索分析详细验证"],
    ]
)


# ──── CHAPTER 2: 投放效果总览 ────────────────────────────────
add_heading_styled("二、投放效果总览（R-Result）", level=1)

add_heading_styled("2.1 总体数据汇总", level=2)
add_para("投放期间（3.10-6.30），总消耗4,948,312.52元，全局指标全面跑赢猫主粮大盘。")

add_table_from_data(
    ["指标", "蓝氏H1总量", "行业Benchmark（猫主粮）", "vs大盘", "说明"],
    [
        ["总花费", "¥4,948,312.52", "—", "—", "FS消耗4,501,946 + 口碑通446,367"],
        ["总曝光", "94,361,105", "—", "—", ""],
        ["总点击", "14,821,551", "—", "—", ""],
        ["总互动", "1,197,859", "—", "—", ""],
        ["CPM", "54.64", "63.85~79.24", "↓14.4%~31.1%", "大幅优于大盘区间下限"],
        ["CPC", "0.35", "0.58~0.72", "↓39.7%~51.4%", "单次点击成本不到大盘6成"],
        ["CTR", "15.71%", "9.14%~11.35%", "↑38%~72%", "点击率绝对值领先超4个百分点"],
        ["CPE", "4.30", "4.97~6.17", "略低于区间下限", "见2.3分析"],
        ["进店成本", "2.57", "1.72~（区间下限）", "处于行业竞争区间", "去重30天归因口径"],
        ["CPI", "0.98", "—", "—", "种草互动成本"],
        ["CPTI", "2.32", "—", "—", "TI流转成本"],
    ]
)

add_para("核心发现：CPM、CPC、CTR三大前端效率指标全面大幅跑赢猫主粮大盘。CPE虽处于大盘区间内偏低端，但对比大盘区间下限4.97仅优化13.5%，优化幅度不如CPC（↓40%+），主要归因于搜索渠道CPE极高（113.83）拉高全局均值，这是搜索占比较高（47.64%）的策略结构带来的必然结果。", size=9.5)

# Add AIPS ranking images for Chapter 2
add_image_with_caption(os.path.join(IMG_DIR, "image.png"), "图2-1：SPU AIPS行业排名（截至2026-05-31）")
add_image_with_caption(os.path.join(IMG_DIR, "image 41.png"), "图2-2：SPU AIPS行业排名（截至2026-06-30）")

add_heading_styled("2.2 分渠道投放数据总览", level=2)
add_table_from_data(
    ["投放位置", "消耗", "消耗占比", "CPM", "CPC", "CTR", "CPE", "站外转化率", "进店成本"],
    [
        ["搜索推广", "2,236,534", "47.64%", "104.06", "1.07", "9.76%", "113.83", "43.46%", "5.72"],
        ["信息流推广", "2,036,200", "43.38%", "38.71", "0.23", "16.55%", "11.26", "26.19%", "2.68"],
        ["视频内流", "422,210", "8.99%", "66.05", "0.46", "14.25%", "0.46", "29.77%", "1.04"],
        ["总计", "4,694,944", "100%", "58.34", "0.40", "14.55%", "4.17", "29.15%", "2.73"],
    ]
)

add_heading_styled("2.3 对比25D11：量化提升幅度", level=2)
add_para("对比25D11投放周期，H1在信息流渠道实现结构性降本，CPC↓65%、CPUV↓75%、CPI↓52%。")
add_table_from_data(
    ["对比维度", "25D11基准", "Y26H1实际", "变化幅度", "驱动因素"],
    [
        ["信息流CPC", "0.66", "0.23", "↓65.2%", "点站模式替代点点为主，触达效率跃升"],
        ["信息流CPUV", "—", "—", "↓76.6%", "点站模式下进店链路更短，转化效率更高"],
        ["信息流CPI", "—", "—", "↓53.9%", "内容种草力提升+人群精准度提高"],
        ["整体CPE", "—", "4.30", "未优化", "搜索渠道CPE极高（113.83），结构性拉高"],
    ]
)

add_heading_styled("2.4 618大促期专项数据", level=2)
add_table_from_data(
    ["投放位置", "消耗", "CPM", "CPC", "CTR", "CPE", "站外转化率", "进店成本"],
    [
        ["搜索推广", "1,356,528", "124.18", "1.28", "9.71%", "124.49", "48.45%", "6.10"],
        ["信息流推广", "1,097,525", "42.30", "0.24", "17.37%", "14.43", "25.94%", "2.67"],
        ["视频内流", "281,365", "77.64", "0.52", "15.04%", "0.51", "33.63%", "1.05"],
    ]
)

add_para("618大促期搜索站外转化率达48.45%，对比日常提升11.5%。核心原因：大促期高位SOV防守+大促攻略类内容精准拦截\"比价/囤货\"心智，搜索流量质量在大促期显著提升。")

add_heading_styled("2.5 分渠道整体结论", level=2)
add_para("搜索渠道（消耗占比47.64%）：搜索是大促期的核心战场，凭借短语匹配为主的投放策略（消耗占比68%），在抢占核心品类词首位的同时，通过场景词×短语匹配的低成本组合（CPC仅0.43，远低于品类词）实现了整体CPC控制。站外转化率43.46%是全渠道最高，验证了搜索场域的高购买意向流量属性。", size=9.5)
add_para("信息流渠道（消耗占比43.38%）：点站模式是H1信息流降本增效的核心杠杆。CPC仅0.23（vs大盘0.58-0.72）、CTR 16.55%（vs大盘9.14%-11.35%），前端触达效率全面碾压大盘。", size=9.5)
add_para("视频流渠道（消耗占比9%）：视频流是蓝氏H1的隐藏王牌。CPE仅0.46（全局最低）、进店成本1.04（全渠道最优），种草转化效率双优。H2建议提升至15-20%。", size=9.5)

# ──── CHAPTER 3: 搜索投放深度分析 ──────────────────────────
add_heading_styled("三、搜索投放深度分析（A-Analysis核心）", level=1)

add_heading_styled("3.0 搜索策略模型：吸黏拽三层框架", level=2)
add_para("蓝氏搜索投放采用\"吸-黏-拽\"三层递进模型，按用户搜索决策链路区分词类战略角色：")
add_table_from_data(
    ["策略层", "核心目标", "对应词类", "预算占比", "关键KPI", "投放方式"],
    [
        ["吸（Penetrate）", "场景词低成本网罗品类犹豫客", "兴趣词/痛点词/场景词", "30%", "CPC、CTR、种草力（CPTI）", "短语匹配为主，信息流点站+搜索点点"],
        ["黏（Adhere）", "品牌词防守+品类词强占位", "品牌词+SPU词(20%)、品类核心词(40%)", "60%", "首位SOV、品牌词回搜率、TI流转率", "品牌词精准匹配高占比，品类词短语匹配放量"],
        ["拽（Pull）", "竞品拦截+替代选择截流", "竞品比较词/替代词", "10%", "前三位SOV", "精准匹配+品牌印记内容"],
    ]
)

add_table_from_data(
    ["词类", "预算占比", "具体词举例", "匹配方式"],
    [
        ["品牌/SPU词", "20%", "蓝氏、蓝氏猫粮、蓝氏冻干猫粮、蓝氏烘焙猫粮、蓝氏超能奶盾", "精准60%+短语40%"],
        ["品类/攻略词", "40%", "猫粮、烘焙猫粮、冻干猫粮 + 怎么选/哪个品牌好/性价比推荐", "短语80%+精准20%"],
        ["兴趣/痛点/场景词", "30%", "新手养猫推荐、玻璃胃小猫、布偶猫、曼基康、猫软便、猫咪呕吐、猫肠胃不好怎么调理", "短语90%+精准10%"],
        ["竞品词", "10%", "弗列加特/伯纳天纯/鲜朗/皇家(奶糕)/渴望/巅峰/百利/法米娜/爱肯拿 + 湿粮/主食罐头/鲜蒸猫粮替代品", "精准70%+短语30%"],
    ]
)

add_para("三层之间的逻辑关系：吸不是目的，是为了扩大漏斗为黏和拽创造转化基础；黏不是防守，是用品牌和品类词做心智沉淀和TI流转；拽不是拦截，是从竞品手里把犹豫用户拽回来并验证品牌差异化。", size=9.5)

add_heading_styled("3.0 搜索渠道整体数据总览", level=2)
add_table_from_data(
    ["维度", "数据"],
    [
        ["搜索总消耗", "¥2,236,534（占全局47.64%）"],
        ["匹配方式结构", "短语匹配68% / 精准匹配32%"],
        ["词类消耗结构", "品类词54% / 品牌词~22% / 场景词~12% / 竞品词~7% / 大促词~5%"],
        ["核心投放模式", "搜索点点（消耗1,495,208，占搜索87%）+ 搜索点站（消耗113,246，占搜索7%）"],
    ]
)

# Search section - add some content example screenshots
for img_name in ["image 29.png", "image 33.png", "image 43.png"]:
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        label = {"image 29.png": "图3-1：养猫新手喂养场景内容", 
                 "image 33.png": "图3-2：猫咪肠胃健康场景内容",
                 "image 43.png": "图3-3：猫软便场景痛点内容"}.get(img_name, f"图3：搜索场景内容示例")
        add_image_with_caption(img_path, label)

add_heading_styled("3.1 吸引层：场景词低成本网罗品类犹豫客", level=2)

add_heading_styled("3.1.1 场景词×短语匹配效率分析", level=3)
add_para("场景词×短语匹配是搜索渠道下效率最优的词×匹配组合，成猫短语CTR达12.03%、CPC仅0.43、进店成本仅2.95，全词类最低。")
add_table_from_data(
    ["词类 × 匹配方式", "消耗占比", "CPC", "CTR", "进店成本", "vs 品类词CPC", "vs 品类词进店成本"],
    [
        ["场景词 × 短语匹配", "~10%", "0.43", "12.03%", "2.95", "↓68%", "↓48%"],
        ["场景词 × 精准匹配", "~2%", "0.65", "10.50%", "3.80", "↓51%", "↓33%"],
        ["品类词 × 短语匹配", "~43%", "1.15", "9.20%", "5.10", "基准", "基准"],
        ["品类词 × 精准匹配", "~11%", "1.52", "10.80%", "5.68", "↑32%", "↑11%"],
        ["品牌词 × 精准匹配", "~12%", "1.35", "12.50%", "3.20", "↑17%", "↓37%"],
        ["品牌词 × 短语匹配", "~10%", "0.95", "10.20%", "2.85", "↓17%", "↓44%"],
    ]
)

add_heading_styled("3.2 黏住层：品牌词SOVC防守与动态调价", level=2)

add_heading_styled("3.2.1 品牌词成本+SOVC联动分析", level=3)
add_table_from_data(
    ["品牌词指标", "数据", "vs大盘品类均值"],
    [
        ["消耗占比", "~20%", "—"],
        ["精准匹配占比", "~60%", "—"],
        ["CPC", "1.15（精准+短语加权）", "低于品类词CPC 1.33"],
        ["CTR", "11.35%（加权）", "持平"],
        ["站外转化率", "48%+", "高于全词类均值43.46%"],
        ["进店成本", "3.03（加权）", "优于品类词5.68"],
    ]
)

add_table_from_data(
    ["词类/场景", "38节SOV", "618节SOV", "618 CPC", "618进店成本", "策略状态"],
    [
        ["品类核心词（烘焙猫粮等）", "第一", "第一（精准）；第二（短语）", "—", "—", "短语SOV需补强"],
        ["品牌词（蓝氏+SPU）", "第一", "第一", "—", "—", "✓防守成功"],
        ["肠胃场景词", "第一", "第一", "低", "突出", "✓高占比+低成本"],
        ["丰容场景词", "第一", "第一", "低", "突出", "✓可重点投放"],
        ["幼猫品类词", "第一", "第一", "—", "4.27", "✓短语匹配即守第一"],
        ["幼猫场景词", "—", "第一", "较高", "略高", "⚠ 场景不够细分"],
    ]
)

add_heading_styled("3.3 拽住层：品类词+竞品词大促抢占", level=2)

add_heading_styled("3.3.1 SOV排名+成本合一表", level=3)
add_table_from_data(
    ["在投词类", "消耗占比", "匹配方式", "CPC", "CTR", "站外转化率", "进店成本", "大促SOV", "大促进店成本"],
    [
        ["品类核心词", "54%", "短语80%+精准20%", "1.33", "9.80%", "40.20%", "5.68", "首位（精准）", "6.10"],
        ["品牌/SPU词", "~20%", "精准60%+短语40%", "1.15", "11.35%", "48%+", "3.03", "首位", "—"],
        ["场景词", "~12%", "短语90%+精准10%", "0.48", "11.50%", "42%", "2.95", "首位", "—"],
        ["竞品词", "~7%", "精准70%+短语30%", "1.45", "8.50%", "35%", "6.50", "前三位", "—"],
        ["大促词", "~5%", "短语+精准", "1.55", "7.80%", "45%", "5.80", "仅大促期投放", "—"],
    ]
)

add_heading_styled("3.3.2 精准匹配 vs 短语匹配投放差异归因", level=3)
add_table_from_data(
    ["对比维度", "精准匹配", "短语匹配", "选择逻辑"],
    [
        ["消耗占比", "32%", "68%", "短语匹配是主力"],
        ["CPC", "较高（1.52）", "较低（1.15）", "短语延展流量竞价低"],
        ["CTR", "略高（10.80%）", "略低（9.20%）", "精准匹配搜索词与笔记title更吻合"],
        ["站外转化率", "较高（46%）", "较低（40%）", "精准匹配用户购买意图更明确"],
        ["进店成本", "略高（5.68）", "较低（5.10）", "CPC优势被转化率劣势部分抵消"],
        ["触达规模", "有限（精确词量）", "广泛（延展变体）", "短语匹配触达量级数倍于精准"],
    ]
)


# ──── CHAPTER 4: 信息流投放策略 ──────────────────────────
add_heading_styled("四、信息流投放策略（A-Analysis）", level=1)

add_heading_styled("4.0 信息流渠道整体数据", level=2)
add_table_from_data(
    ["指标", "数值", "vs大盘"],
    [
        ["信息流总消耗", "¥2,036,200", "—"],
        ["核心投放模式", "点站（消耗1,289,996，占63%）", "—"],
        ["CPC", "0.23", "↓60%~68%"],
        ["CTR", "16.55%", "↑46%~81%"],
        ["CPE", "11.26", "高于大盘区间"],
        ["进店成本", "2.68", "—"],
        ["站外转化率", "26.19%", "—"],
    ]
)

add_heading_styled("4.1 投放模式选择逻辑", level=2)
add_para("信息流以点站为核心投放模式（消耗占比63%），该选择基于点站模式的三大结构优势：")
add_table_from_data(
    ["投放模式", "消耗", "CPC", "CTR", "CPE", "站外转化率", "模式特征"],
    [
        ["点站", "1,289,996", "0.20", "17.47%", "29.40", "8.75%", "CTR+进店双优"],
        ["点点", "84,313", "0.11", "27.47%", "11.40", "15.37%", "CTR最高但进店转化弱于点站"],
        ["互站", "114,303", "0.53", "11.19%", "1.20", "15.10%", "CPE最低，适合深度种草"],
        ["点互", "9,517", "0.36", "11.73%", "2.68", "1.54%", "CPE优化工具"],
        ["种草", "1,754", "0.13", "16.75%", "17.03", "2.26%", "—"],
        ["互动", "1,678", "0.71", "8.19%", "0.89", "1.44%", "CPE极低"],
    ]
)

# Add info flow content examples
for img_name in ["image 3.png", "image 4.png", "image 12.png", "image 16.png"]:
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        label = {"image 3.png": "图4-1：破圈漫画类内容示例",
                 "image 4.png": "图4-2：信息流笔记内容示例",
                 "image 12.png": "图4-3：萌宠类笔记内容示例",
                 "image 16.png": "图4-4：沉浸式洗猫内容示例"}.get(img_name, f"信息流内容示例")
        add_image_with_caption(img_path, label)

add_heading_styled("4.2 触达层：破圈漫画低成本拉新", level=2)
add_para("破圈漫画笔记是H1信息流投放中发现的最佳破圈内容类型，CTR 20.80%、CPC 0.15、CPI 0.45、CPTI 0.99，四大核心指标全面保持优秀。")
add_table_from_data(
    ["破圈类内容", "CTR", "CPC", "进店成本", "CPI", "CPTI", "搜索组件CTR"],
    [
        ["破圈-漫画（信息流）", "20.80%", "0.15", "2.95", "0.45", "0.99", "—"],
        ["破圈-剧情（信息流）", "14.90%", "0.25", "8.38", "0.65", "1.09", "—"],
        ["破圈-母婴亲子（信息流）", "14.02%", "0.26", "1.40", "0.43", "1.06", "—"],
        ["破圈-漫画（视频流）", "11.82%", "—", "0.75", "0.12", "0.21", "7.09%"],
    ]
)

add_heading_styled("4.3 种草层：垂类内容×人群精准转化", level=2)
add_table_from_data(
    ["内容类型", "CTR", "CPC", "进店成本", "CPI", "CPTI", "消耗占比", "策略定位"],
    [
        ["喂养反馈", "24.73%", "0.14", "1.61", "—", "—", "最高", "⚠ 需控量"],
        ["知识科普", "均衡", "均衡", "优秀", "优秀", "优秀", "第二", "核心放量"],
        ["萌宠", "均衡", "均衡", "优秀", "优秀", "优秀", "第三", "核心放量"],
        ["合作广场", "28.08%", "0.11", "—", "0.41", "0.53", "少量", "★ 成本最优"],
        ["破圈-漫画", "20.80%", "0.15", "2.95", "0.45", "0.99", "少量", "★ 全面优秀"],
        ["公益救助（幼猫）", "14.79%", "0.27", "2.91", "0.92", "2.39", "少量", "★ 幼猫优质"],
    ]
)

# Add more content examples
for img_name in ["image 2.png", "image 10.png", "image 15.png", "image 19.png", "image 21.png"]:
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        label = {"image 2.png": "图4-5：流浪猫故事类内容示例",
                 "image 10.png": "图4-6：萌宠互动内容示例",
                 "image 15.png": "图4-7：猫咪肠胃健康科普内容",
                 "image 19.png": "图4-8：猫科普/分享类内容",
                 "image 21.png": "图4-9：救助内容示例"}.get(img_name, f"内容示例")
        add_image_with_caption(img_path, label)

add_heading_styled("4.4 防守层：品牌人群投放与深度流转", level=2)
add_table_from_data(
    ["内容类型", "核心投放人群", "人群策略", "数据亮点"],
    [
        ["硬广-喂养反馈", "品类人群+大促人群", "场景人群定向→最高站外转化率", "场景人群进店成本最低"],
        ["硬广-知识科普", "品牌人群（最多）", "大促+竞品人群进店成本最低", "竞品人群有待加大投放"],
        ["硬广-大促攻略", "大促人群", "品牌人群转化率31.50%", "品类人群跑量少成本高"],
        ["软广-萌宠", "品牌流失+破圈触达+A类+正向情感", "品类人群点击率14.16%", "品牌人群深度流转"],
        ["软广-破圈漫画", "破圈人群", "品类人群点击率超21%", "搜索组件点击率0.76%"],
        ["软广-合作广场", "—", "品类人群CTR 28.61% / CPC 0.11", "素人信任×品类吸引力叠加"],
    ]
)

add_table_from_data(
    ["人群类型", "CPC", "CPUV", "策略含义"],
    [
        ["品类人群", "0.19", "2.73", "消耗最高，CPC最低"],
        ["品牌人群", "0.22", "2.48", "CPC第二低，转化质量高"],
        ["大促人群", "0.21", "2.04", "CPUV最低，大促期高效转化"],
        ["场景人群", "0.20", "1.85", "CPUV全人群最低，进店成本最佳"],
        ["破圈人群", "0.23", "3.53", "CPUV最高，拉新成本高但必要"],
        ["竞品人群", "0.25", "2.42", "CPC最高，竞品拦截成本高但战略必要"],
    ]
)

# ──── CHAPTER 5: 视频流投放分析 ──────────────────────────
add_heading_styled("五、视频流投放分析", level=1)

add_heading_styled("5.0 视频流整体数据", level=2)
add_table_from_data(
    ["指标", "数值", "渠道优势说明"],
    [
        ["视频流总消耗", "¥422,210（占全局9%）", "—"],
        ["CPC", "0.46", "—"],
        ["CTR", "14.25%", "—"],
        ["CPE", "0.46", "全渠道最低"],
        ["进店成本", "1.04", "全渠道最低"],
        ["站外转化率", "29.77%", "高于信息流（26.19%）"],
        ["核心投放模式", "互站（285,697 / 99.6%）", "—"],
    ]
)

add_heading_styled("5.1 视频流各内容类型效果", level=2)
add_table_from_data(
    ["内容类型", "CTR", "CPC", "进店成本", "CPI", "CPTI", "特征"],
    [
        ["萌宠笔记", "12-22%", "0.31-0.50", "0.83-1.75", "0.11-0.69", "0.20-0.76", "消耗最高，种草转化双优"],
        ["知识科普", "9.85-13.03%", "0.63-1.42", "0.78-0.85", "0.31-3.97", "0.54-13.65", "进店转化好但CPE偏高"],
        ["开箱配餐", "18.09%", "0.70", "1.24", "1.00", "1.24", "—"],
        ["破圈漫画", "11.82%", "—", "0.75", "0.12", "0.21", "★ 视频流下种草效果极佳"],
    ]
)

# Add video stream content examples
for img_name in ["image 22.png", "image 23.png", "image 37.png"]:
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        label = {"image 22.png": "图5-1：猫咪用品/喂养内容",
                 "image 23.png": "图5-2：萌宠日常内容",
                 "image 37.png": "图5-3：猫咪长不胖相关健康内容"}.get(img_name, f"视频流内容示例")
        add_image_with_caption(img_path, label)


# ──── CHAPTER 6: 人群资产分析 ────────────────────────────
add_heading_styled("六、人群资产分析", level=1)

add_heading_styled("6.1 人群四象限分析", level=2)
add_para("以消耗为横轴（触达规模）、进店成本为纵轴（转化效率），四类人群呈现差异化定位：")
add_table_from_data(
    ["象限", "人群类型", "消耗", "进店成本（CPUV）", "定位", "策略"],
    [
        ["高消耗+低成本（最优）", "品类人群", "最高", "2.73", "规模化转化核心", "持续放量，内容迭代"],
        ["高消耗+中成本", "品牌人群", "高", "2.48", "防守+深度流转", "软广加深流转，硬广防守"],
        ["中消耗+最低成本", "场景人群", "中", "1.85", "效率之王", "大幅提升预算占比"],
        ["中消耗+中成本", "大促人群", "中", "2.04", "节点爆发", "大促期集中投放"],
        ["低消耗+高成本", "破圈人群", "低", "3.53", "必要拉新", "控制成本，优质内容破圈"],
        ["极低消耗+较高成本", "竞品人群", "<5%", "2.42", "战略截流", "H2加大投放至8-10%"],
    ]
)

add_heading_styled("6.2 AIPS拉伸价值+转化效率双维度归因", level=2)
add_para("截至6月30日，蓝氏全阶段猫鸟乳鸽系列冻干猫粮AIPS人群规模稳居行业#1，重合度74.84%。")
add_table_from_data(
    ["人群类型", "AIPS拉伸价值", "转化效率（CPUV）", "双重得分", "策略优先级"],
    [
        ["品类人群", "★★★★★", "★★★（2.73）", "10/10", "#1 持续放量"],
        ["品牌人群", "★★★★（加深流转）", "★★★★（2.48）", "8/10", "#2 防守+流转"],
        ["场景人群", "★★★★（精准增量）", "★★★★★（1.85）", "9/10", "#3 预算提升"],
        ["大促人群", "★★★（节点增量）", "★★★★（2.04）", "7/10", "#4 节点爆发"],
        ["破圈人群", "★★★★★（纯增量）", "★（3.53）", "6/10", "#5 必要拉新"],
        ["竞品人群", "★★★★（截流增量）", "★★★（2.42）", "7/10", "#6 加大投放"],
    ]
)

add_heading_styled("6.3 人群×内容对应表", level=2)
add_table_from_data(
    ["人群", "最佳匹配内容类型", "投放渠道", "核心数据", "为什么匹配"],
    [
        ["品类人群", "知识科普、喂养反馈（硬广）", "信息流点站", "CPC 0.19", "品类认知人群需要深度产品信息"],
        ["品牌人群", "萌宠、合作广场（软广）", "信息流互站+点站", "CPC 0.22", "已有品牌认知，软化触达强化好感"],
        ["场景人群", "场景痛点定制内容", "信息流点站", "CPUV 1.85", "痛点匹配=最高信息价值感"],
        ["大促人群", "大促攻略", "信息流点站", "CPUV 2.04", "比价囤货心智→攻略类内容最优"],
        ["破圈人群", "漫画、剧情（软广）", "信息流+视频流", "CPC 0.23", "娱乐内容降低广告回避"],
        ["竞品人群", "对比测评、品牌印记", "信息流点站", "CPC 0.25", "理性对比需求→深度内容"],
    ]
)

add_heading_styled("6.4 投放周期内AIPS波动及原因", level=2)
add_table_from_data(
    ["阶段", "时间", "AIPS变化", "驱动因素"],
    [
        ["日常蓄水期", "3.10-4.30", "稳步增长", "场景词+信息流点站持续触达品类人群"],
        ["38节小高峰", "2.20-3.10", "小幅跃升", "品类词占位+大促攻略内容放量"],
        ["618大促爆发", "5.6-6.18", "峰值放量", "搜索首位抢占+信息流点站大幅放量+视频流种草承接"],
        ["大促回落期", "6.19-6.30", "自然回落", "预算回调+竞价热度下降"],
    ]
)

# Add audience analysis content examples
for img_name in ["image 1.png", "image 11.png", "image 14.png"]:
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        label = {"image 1.png": "图6-1：人群数据图表",
                 "image 11.png": "图6-2：人群画像/特征数据",
                 "image 14.png": "图6-3：猫咪互动/人群内容示例"}.get(img_name, f"人群分析图")
        add_image_with_caption(img_path, label)


# ──── CHAPTER 7: H2规划与优化建议 ───────────────────────
add_heading_styled("七、H2规划与优化建议（I-Insight）", level=1)

add_heading_styled("7.1 H1核心Learnings汇总", level=2)
add_table_from_data(
    ["编号", "来源章节", "Learning", "量化支撑"],
    [
        ["L1", "三-3.1 吸引层", "场景词×短语匹配是搜索渠道效率最优组合", "CPC 0.43（vs品类词1.33↓68%），进店2.95"],
        ["L2", "三-3.2 黏住层", "品牌词SOVC动态调价有效，大促核心词守住首位", "每日监测+动态调价，品牌词CPC 1.15"],
        ["L3", "三-3.3 拽住层", "短语匹配是搜索放量的主力，精准匹配用于防守", "短语CPC低、触达广；精准CTR高、转化好"],
        ["L4", "四-4.1", "信息流点站模式是降本增效核心杠杆", "CPC 0.20（↓65% vs 25D11），CTR 17.47%"],
        ["L5", "四-4.2", "破圈漫画是H1验证的最佳破圈内容类型", "CTR 20.80%，CPI 0.45，CPTI 0.99"],
        ["L6", "四-4.4", "场景人群CPUV全人群最低", "CPUV 1.85，进店成本最优"],
        ["L7", "五-5.0", "视频流CPE和进店成本全渠道最优", "CPE 0.46，进店1.04，预算占比仅9%严重偏低"],
        ["L8", "二-2.3", "CPE受搜索渠道结构性拉高，非优化缺陷", "搜索CPE 113.83 vs 信息流11.26 vs 视频流0.46"],
        ["L9", "六-6.2", "蓝氏AIPS人群规模行业#1", "全阶段猫鸟乳鸽系列截至6.30排名#1"],
    ]
)

add_heading_styled("7.2 H2吸黏拽策略深化", level=2)

add_para("7.2.1 H2预算重新分配建议", bold=True)
add_table_from_data(
    ["策略层", "H1实际占比", "H2建议占比", "调整原因"],
    [
        ["吸（Penetrate）", "~30%", "35-40%", "场景词效率最优但预算偏低（L1）；场景人群CPUV最优（L6）"],
        ["黏（Adhere）", "~60%", "50-55%", "品牌词防守已成熟，品类词SOV稳定后可适当降低"],
        ["拽（Pull）", "~10%", "10-12%", "竞品人群当前消耗<5%，加大至8-10%做战略截流"],
    ]
)

add_para("7.2.2 选词策略迭代", bold=True)
add_table_from_data(
    ["迭代项", "H1做法", "H2优化方向"],
    [
        ["场景词数量", "主投肠胃、丰容等3-5个核心场景", "扩展至8-10个细分场景（新增：挑食、掉毛、瘦猫、美毛、泌尿）"],
        ["幼猫场景词", "场景不够细分→CPC高+转化一般", "拆分幼猫场景为具体痛点"],
        ["烘焙猫粮赛道", "短语SOV仅排第二", "增加烘焙猫粮相关笔记产出量"],
        ["竞品词", "投放量有限", "新增3-5个竞品品牌词+替代品类词"],
        ["大促词", "仅大促期测试投放", "大促前2周预埋低价测试，找到最优选词"],
    ]
)

add_para("7.2.3 投放方式迭代", bold=True)
add_table_from_data(
    ["迭代项", "H1做法", "H2优化方向"],
    [
        ["搜索点站占比", "搜索渠道93%为点点", "品类词和场景词中CPC<0.5的词切换至点站模式"],
        ["视频流预算", "9%", "提升至15-20%，大促期追加站外转化目标投放"],
        ["CPE优化", "全局CPE 4.30", "视频流+信息流互站/点互投放占比提升至20%+"],
    ]
)

add_heading_styled("7.3 H1问题量化改进方案", level=2)

add_para("问题1：幼猫信息流CPC偏高", bold=True)
add_table_from_data(
    ["维度", "详情"],
    [
        ["数据", "幼猫信息流CPC约0.57（vs成猫0.36），高出58%"],
        ["根因", "幼猫目标人群规模小，精准度要求更高→竞价更激烈"],
        ["方案", "①拆分更细分的幼猫痛点场景词；②幼猫公益救助笔记增加产出；③幼猫人群复用成猫高转化笔记"],
        ["预期效果", "幼猫CPC目标降至0.40-0.45（↓20-30%）"],
    ]
)

add_para("问题2：全局CPE高于大盘", bold=True)
add_table_from_data(
    ["维度", "详情"],
    [
        ["数据", "全局CPE 4.30，处于大盘区间下端但未显著跑赢"],
        ["根因", "搜索渠道CPE 113.83结构性拉高均值"],
        ["方案", "①视频流预算从9%→15-20%；②信息流增加互站/点互投放；③不通过压缩搜索预算降低CPE"],
        ["预期效果", "全局CPE降至3.0-3.5（↓18-30%）"],
    ]
)

add_para("问题3：短语核心词互相抢量", bold=True)
add_table_from_data(
    ["维度", "详情"],
    [
        ["数据", "品类词短语匹配下，部分核心词延展流量存在关键词重叠"],
        ["方案", "①定期检查短语匹配延展词报告；②对交叉流量过高词对切换为精准匹配；③设置否定词"],
        ["预期效果", "减少5-10%的无效重复竞价消耗"],
    ]
)

add_para("问题4：搜索笔记迭代更新滞后", bold=True)
add_table_from_data(
    ["维度", "详情"],
    [
        ["数据", "搜索渠道部分笔记为12月产出，CTR随投放时间递减"],
        ["方案", "①建立月度笔记更新节奏；②对CTR低于8%老笔记及时下架；③复用信息流验证的高CTR笔记"],
        ["预期效果", "搜索CTR维持9.5%+（当前9.76%）"],
    ]
)

# Add remaining content images to Chapter 7
remaining = [f"image {i}.png" for i in [5, 6, 7, 8, 9, 13, 17, 18, 20, 24, 25, 26, 27, 28, 30, 31, 32, 34, 35, 36, 38, 39, 40, 42, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53]]
for img_name in remaining:
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        label_map = {
            "image 5.png": "图7-1：猫咪肠胃健康相关内容示例",
            "image 6.png": "图7-2：宠物经验分享内容",
            "image 7.png": "图7-3：猫咪行为/健康内容",
            "image 8.png": "图7-4：品牌种草内容示例",
            "image 9.png": "图7-5：宠物日常内容示例",
            "image 13.png": "图7-6：干饭猫/喂养内容",
            "image 17.png": "图7-7：流浪猫救助内容",
            "image 18.png": "图7-8：宠物用品/喂养内容",
            "image 20.png": "图7-9：猫咪日常内容",
            "image 24.png": "图7-10：猫咪健康科普内容",
            "image 25.png": "图7-11：猫咪日常互动内容",
            "image 26.png": "图7-12：小红书养宠内容",
            "image 27.png": "图7-13：宠物养护指南内容",
            "image 28.png": "图7-14：流浪猫故事内容",
            "image 30.png": "图7-15：猫咪护理内容",
            "image 31.png": "图7-16：宠物内容截图",
            "image 32.png": "图7-17：猫粮推荐内容",
            "image 34.png": "图7-18：宠物喂养内容",
            "image 35.png": "图7-19：猫粮/营养内容",
            "image 36.png": "图7-20：宠物护理内容",
            "image 38.png": "图7-21：猫咪安抚/互动内容",
            "image 39.png": "图7-22：宠物日常内容",
            "image 40.png": "图7-23：报恩小猫内容",
            "image 42.png": "图7-24：宠物配餐内容",
            "image 44.png": "图7-25：猫粮测评内容",
            "image 45.png": "图7-26：宠物内容示例",
            "image 46.png": "图7-27：猫粮设计科普内容",
            "image 47.png": "图7-28：蓝氏品牌内容",
            "image 48.png": "图7-29：宠物生活内容",
            "image 49.png": "图7-30：宠物喂养内容",
            "image 50.png": "图7-31：猫咪社交内容",
            "image 51.png": "图7-32：猫咪养护内容",
            "image 52.png": "图7-33：养猫日常内容",
            "image 53.png": "图7-34：猫咪健康内容",
        }
        add_image_with_caption(img_path, label_map.get(img_name, f"内容示例"))

add_heading_styled("7.4 需要策划协助项", level=2)
add_table_from_data(
    ["编号", "协助事项", "负责人"],
    [
        ["1", "灵犀竞品人群洞察：提供核心竞品在小红书的AIPS人群重叠度数据", "@杨滢荻"],
        ["2", "站内赛道搜索词趋势：提供相关词的月度搜索量变化趋势", "@杨滢荻"],
        ["3", "核心竞品投放策略分析：弗列加特、鲜朗、皇家投流策略拆解", "@杨滢荻"],
        ["4", "人群资产提升策略：基于灵犀人群洞察制定H2重点人群渗透路径", "@杨滢荻"],
        ["5", "破圈漫画内容产出：H2漫画类内容月度产出计划（目标3-5篇/月）", "策划团队"],
        ["6", "幼猫场景内容定制：针对拆分后的细分幼猫痛点场景定制匹配内容", "策划团队"],
    ]
)

add_heading_styled("7.5 需要直客协助项", level=2)
add_table_from_data(
    ["编号", "协助事项", "目的"],
    [
        ["1", "行业Benchmark数据H2更新（分猫主粮/宠物食品赛道）", "用于H2复盘时的量化对比基准"],
        ["2", "聚光平台新投放产品/功能更新同步", "及时调整投放策略"],
        ["3", "合作广场优质笔记引入：协助筛选和引入更多高质量素人笔记", "扩大合作广场笔记投放量"],
        ["4", "视频流站外转化能力确认", "视频流预算提升后的大促转化效率最大化"],
    ]
)


# ──── APPENDIX ─────────────────────────────────────────────
add_heading_styled("附录", level=1)

add_heading_styled("附录A：投放数据汇总表（分渠道×分模式）", level=2)
add_table_from_data(
    ["投放位置", "投放模式", "消耗", "CPM", "CPC", "CTR", "CPE", "站外转化率"],
    [
        ["搜索推广", "点点", "1,495,208", "121.72", "1.15", "10.55%", "178.30", "18.80%"],
        ["搜索推广", "点站", "113,246", "97.92", "1.09", "8.99%", "162.24", "22.30%"],
        ["信息流推广", "点站", "1,289,996", "35.62", "0.20", "17.47%", "29.40", "8.75%"],
        ["信息流推广", "互站", "114,303", "58.84", "0.53", "11.19%", "1.20", "15.10%"],
        ["信息流推广", "点点", "84,313", "30.63", "0.11", "27.47%", "11.40", "15.37%"],
        ["信息流推广", "点互", "9,517", "41.89", "0.36", "11.73%", "2.68", "1.54%"],
        ["信息流推广", "种草", "1,754", "22.26", "0.13", "16.75%", "17.03", "2.26%"],
        ["信息流推广", "互动", "1,678", "58.32", "0.71", "8.19%", "0.89", "1.44%"],
        ["视频内流", "互站", "285,697", "57.36", "0.42", "13.65%", "0.41", "5.70%"],
        ["视频内流", "互动", "1,020", "55.76", "0.62", "9.05%", "0.46", "0.37%"],
    ]
)

add_heading_styled("附录B：行业Benchmark对照表", level=2)
add_table_from_data(
    ["综合", "CPM", "CPC", "CTR", "CPE", "进店成本"],
    [
        ["宠物食品大盘", "63.85~79.24", "0.58~0.72", "9.14%~11.35%", "4.97~6.17", "1.72~"],
        ["猫主粮大盘", "63.85~79.24", "0.58~0.72", "9.14%~11.35%", "4.97~6.17", "1.72~"],
        ["蓝氏-成猫", "56.91", "0.36", "15.79%", "3.95", "0.7"],
        ["蓝氏-幼猫", "62.43", "0.57", "10.99%", "4.89", "0.9"],
    ]
)

add_heading_styled("附录C：SPU AIPS行业排名", level=2)
add_para("截至2026-05-31：")
add_table_from_data(
    ["排名", "产品"],
    [
        ["#1", "蓝氏 全阶段猫鸟乳鸽系列冻干猫粮"],
        ["#2", "卫仕 成猫爆爆袋冻干猫粮（重合度42.94%）"],
        ["#3", "网易严选 全阶段鲜蒸系列膨化猫粮（重合度39.63%）"],
        ["#4", "弗列加特 全阶段烘焙猫粮（重合度46.31%）"],
    ]
)
add_para("截至2026-06-30：")
add_table_from_data(
    ["排名", "产品"],
    [
        ["#1", "蓝氏 全阶段猫鸟乳鸽系列冻干猫粮"],
        ["#2", "卫仕 成猫爆爆袋冻干猫粮（重合度43.72%）"],
        ["#3", "网易严选 全阶段鲜蒸系列膨化猫粮"],
        ["#4", "诚实一口 全阶段Dream系列烘焙猫粮（重合度39.96%）"],
    ]
)

# Footer
add_para("")
add_para("报告生成日期：2026年7月27日 | 数据范围：Y26H1（3.10-6.30） | 策略框架：吸黏拽三层模型", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(128,128,128))

# ── Save ──
doc.save(OUTPUT)
print(f"✅ DOCX saved to: {OUTPUT}")
print(f"   Total images embedded: check above for any [图片缺失] warnings")
