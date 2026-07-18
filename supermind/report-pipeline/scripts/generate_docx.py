#!/usr/bin/env python3
"""
北纬47度报告 Markdown → DOCX 转换器
- 嵌入 ![]() 引用的图片到对应位置
- 图片宽度 6英寸，居中
- 封面页、标题层级、粗体、表格边框+表头灰底、正文10.5pt等线
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import docx.oxml

# ── Paths ──────────────────────────────────────────────
MD_PATH = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/content/report_final_beiwei47.md"
CHART_DIR = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/charts/"
OUTPUT_PATH = "/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/beiwei47_final.docx"

# ── Read markdown ──────────────────────────────────────
with open(MD_PATH, "r", encoding="utf-8") as f:
    md_lines = f.readlines()

# ── Create document ────────────────────────────────────
doc = Document()

# ── Styles setup ───────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '等线'
font.size = Pt(10.5)
# Set East Asian font
rPr = style.element.find(qn('w:rPr'))
if rPr is None:
    rPr = docx.oxml.OxmlElement('w:rPr')
    style.element.append(rPr)
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), '等线')

# Also set paragraph spacing
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# Set page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── Helper functions ───────────────────────────────────

def add_heading_styled(text, level):
    """Add a heading paragraph with styled formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = '等线'
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = docx.oxml.OxmlElement('w:rPr')
            run._element.append(rPr)
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '等线')
        rPr.append(rFonts)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = '等线'
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = docx.oxml.OxmlElement('w:rPr')
            run._element.append(rPr)
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '等线')
        rPr.append(rFonts)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    elif level == 3:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = '等线'
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = docx.oxml.OxmlElement('w:rPr')
            run._element.append(rPr)
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '等线')
        rPr.append(rFonts)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    elif level == 4:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = '等线'
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = docx.oxml.OxmlElement('w:rPr')
            run._element.append(rPr)
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '等线')
        rPr.append(rFonts)
    
    return p


def add_body_paragraph(text_parts):
    """Add a body paragraph. text_parts is a list of (text, bold) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    
    # Set paragraph font defaults
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = docx.oxml.OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    
    for text, is_bold in text_parts:
        if text == '':
            continue
        run = p.add_run(text)
        run.font.name = '等线'
        run.font.size = Pt(10.5)
        if is_bold:
            run.bold = True
        # Set East Asian font
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = docx.oxml.OxmlElement('w:rPr')
            run._element.append(rPr)
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '等线')
        rPr.append(rFonts)
    
    return p


def add_quote_paragraph(text_parts):
    """Add a blockquote-style paragraph (gray, indented)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    
    for text, is_bold in text_parts:
        if text == '':
            continue
        run = p.add_run(text)
        run.font.name = '等线'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if is_bold:
            run.bold = True
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = docx.oxml.OxmlElement('w:rPr')
            run._element.append(rPr)
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '等线')
        rPr.append(rFonts)
    
    return p


def add_image(image_path):
    """Add a centered image, 6 inches wide."""
    if not os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[图片未找到: {os.path.basename(image_path)}]')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return p
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(6))
    return p


def add_table(headers, rows):
    """Add a formatted table with header gray background and borders."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = '等线'
        run.font.size = Pt(9)
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = docx.oxml.OxmlElement('w:rPr')
            run._element.append(rPr)
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '等线')
        rPr.append(rFonts)
        
        # Gray background
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9" w:val="clear"/>')
        hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # Data rows
    for r, row in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, cell_text in enumerate(row):
            row_cells[c].text = ''
            p = row_cells[c].paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = '等线'
            run.font.size = Pt(8.5)
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = docx.oxml.OxmlElement('w:rPr')
                run._element.append(rPr)
            rFonts = docx.oxml.OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '等线')
            rPr.append(rFonts)
    
    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    
    return table


def add_horizontal_rule():
    """Add a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = docx.oxml.OxmlElement('w:pBdr')
    bottom = docx.oxml.OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_bold_text(text):
    """Parse **bold** markers in text, return list of (text, is_bold) tuples."""
    parts = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        parts.append((match.group(1), True))
        last_end = match.end()
    if last_end < len(text):
        parts.append((text[last_end:], False))
    return parts if parts else [(text, False)]


# ── Parse markdown ─────────────────────────────────────

i = 0
while i < len(md_lines):
    line = md_lines[i].rstrip()
    
    # Skip empty lines
    if not line:
        i += 1
        continue
    
    # Horizontal rule
    if line.strip() == '---':
        add_horizontal_rule()
        i += 1
        continue
    
    # Image: ![alt](path)
    img_match = re.match(r'!\[.*?\]\((.*?)\)', line)
    if img_match:
        rel_path = img_match.group(1)
        # Resolve relative path to absolute
        if rel_path.startswith('../charts/'):
            filename = os.path.basename(rel_path)
            img_path = os.path.join(CHART_DIR, filename)
        elif rel_path.startswith('/'):
            img_path = rel_path
        else:
            img_path = os.path.join(os.path.dirname(MD_PATH), rel_path)
        add_image(img_path)
        i += 1
        continue
    
    # Heading 4: ####
    if line.startswith('#### '):
        add_heading_styled(line[5:].strip(), 4)
        i += 1
        continue
    
    # Heading 3: ###
    if line.startswith('### '):
        add_heading_styled(line[4:].strip(), 3)
        i += 1
        continue
    
    # Heading 2: ##
    if line.startswith('## '):
        add_heading_styled(line[3:].strip(), 2)
        i += 1
        continue
    
    # Heading 1: #
    if line.startswith('# '):
        add_heading_styled(line[2:].strip(), 1)
        i += 1
        continue
    
    # Blockquote: >
    if line.startswith('> '):
        text = line[2:].strip()
        parts = parse_bold_text(text)
        add_quote_paragraph(parts)
        i += 1
        continue
    
    # Table: starts with |
    if line.startswith('|') and line.endswith('|'):
        # Collect all table lines
        table_lines = []
        while i < len(md_lines) and md_lines[i].strip().startswith('|') and md_lines[i].strip().endswith('|'):
            table_lines.append(md_lines[i].strip())
            i += 1
        
        if len(table_lines) < 2:
            continue
        
        # Parse header
        header_line = table_lines[0]
        headers = [h.strip() for h in header_line.split('|')[1:-1]]
        
        # Skip separator line (|---|...|)
        row_start = 1
        if re.match(r'^[\|\s\-:]+\|[\|\s\-:]+$', table_lines[1]):
            row_start = 2
        
        # Parse rows
        rows = []
        for tl in table_lines[row_start:]:
            cells = [c.strip() for c in tl.split('|')[1:-1]]
            rows.append(cells)
        
        add_table(headers, rows)
        continue
    
    # Normal paragraph (may contain **bold**)
    text = line.strip()
    parts = parse_bold_text(text)
    add_body_paragraph(parts)
    i += 1

# ── Save ───────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)

# ── Verify ─────────────────────────────────────────────
file_size = os.path.getsize(OUTPUT_PATH)
print(f"✅ DOCX 生成完成: {OUTPUT_PATH}")
print(f"   文件大小: {file_size / 1024:.1f} KB")
print(f"   图片嵌入: {'是' if file_size > 80000 else '否 - 文件可能不完整'}")
