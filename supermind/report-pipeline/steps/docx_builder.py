"""
Step 13: DOCX 生成模块 — 完整重建版。

使用纯 python-docx 组装品牌研究报告文档。
流程：
  创建 Document → 封面 → TOC → brand_overview → 各章节（硬编码标题） → 图表板块 → 附录 → save

章节映射（不 glob，直接文件名）：
  brand_overview     → content/brand_overview.md
  ch1                → content/ch1_findings.md
  ch2                → content/ch2_industry.md
  ch3                → ch3_competitive/deep_*.md + summary_brands.md + competition_patterns.md
  ch4                → content/ch4_deep/{focus_brand}_deep.md
  ch5                → content/ch5_gap.md
  ch6                → content/ch6_recommendations.md
  ch7                → content/ch7_sleep_insights.md
  appendix_innovation → content/innovation_strategy.md
  appendix_founder   → content/founder_research.md
  不输出 pre_research.md

依赖：python-docx + lxml（仅 parse 用，无 zipfile 操作）
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from steps.utils import (
    step_start, step_success, step_fail,
    save_json, load_json, save_text, load_markdown,
    verify_input_file, verify_output_file,
    content_dir, charts_dir, output_dir, BASE_DIR
)
from config import ReportSchema, ProjectConfig


# ── 核心函数 ────────────────────────────────────────────────

def assemble_docx(schema: ReportSchema, project_config: ProjectConfig) -> Path:
    """
    主入口：组装完整 docx 文档。
    纯 python-docx 构建，无 zipfile/lxml 操作。
    """
    step_start("docx_assembly", "DOCX 生成 — 纯 python-docx 重建版")

    out_dir = ensure_output_dir()
    docx_filename = project_config.get("output_settings.docx_filename",
                                        f"{project_config.project_name}_品牌研究报告.docx")
    docx_path = out_dir / docx_filename

    _validate_docx_naming(docx_filename)

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ═══════════════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_config.project_name)
    run.font.size = Pt(32)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{project_config.report_type}")
    run.font.size = Pt(22)

    subtitle_extra = project_config.industry if project_config.industry else ""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle_extra)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%Y年%m月%d日"))
    run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # TOC（手工目录）
    # ═══════════════════════════════════════════════
    _add_heading(doc, '目录', level=1)
    toc_items = _build_toc(schema, project_config)
    for item in toc_items:
        _add_body(doc, item)
    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 加载各章节内容
    # ═══════════════════════════════════════════════
    chapter_map = build_chapter_map(project_config)

    # ═══════════════════════════════════════════════
    # BRAND OVERVIEW
    # ═══════════════════════════════════════════════
    if chapter_map.get('brand_overview'):
        add_md_content_to_docx(doc, chapter_map['brand_overview'])
        doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 产品矩阵（基础扫描第二章）
    # ═══════════════════════════════════════════════
    if chapter_map.get('product_matrix'):
        _add_heading(doc, '产品矩阵', level=1)
        add_md_content_to_docx(doc, chapter_map['product_matrix'])
        doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 渠道与供应链（基础扫描第三章）
    # ═══════════════════════════════════════════════
    if chapter_map.get('channel_supply_chain'):
        _add_heading(doc, '渠道与供应链', level=1)
        add_md_content_to_docx(doc, chapter_map['channel_supply_chain'])
        doc.add_page_break()

    # ═══════════════════════════════════════════════
    # CH1: 核心发现与咨询窗口
    # ═══════════════════════════════════════════════
    _add_heading(doc, '第1章  核心发现与咨询窗口', level=1)
    ch1_text = chapter_map.get('ch1', '')
    if ch1_text:
        ch1_text = re.sub(r'^#\s+.*\n', '', ch1_text, count=1)
        add_md_content_to_docx(doc, ch1_text)
    else:
        _add_body(doc, '[内容待补充]')
    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # CH2: 行业格局与竞品总矩阵
    # ═══════════════════════════════════════════════
    _add_heading(doc, '第2章  行业格局与竞品总矩阵', level=1)
    ch2_text = chapter_map.get('ch2', '')
    if ch2_text:
        add_md_content_to_docx(doc, ch2_text)
    else:
        _add_body(doc, '[内容待补充]')
    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # CH3: 竞品多维度扫描
    # ═══════════════════════════════════════════════
    _add_heading(doc, '第3章  竞品多维度扫描', level=1)

    # 3.1 深度品牌分析
    _add_heading(doc, '3.1  深度品牌分析', level=2)
    deep_brand_text = chapter_map.get('ch3_deep_brands', '')
    if deep_brand_text:
        add_md_content_to_docx(doc, deep_brand_text)

    # 3.2 汇总品牌速览
    _add_heading(doc, '3.2  汇总品牌速览', level=2)
    summary_text = chapter_map.get('ch3_summary', '')
    if summary_text:
        add_md_content_to_docx(doc, summary_text)

    # 3.3 竞争模式归纳
    _add_heading(doc, '3.3  竞争模式归纳', level=2)
    patterns_text = chapter_map.get('ch3_patterns', '')
    if patterns_text:
        add_md_content_to_docx(doc, patterns_text)

    # 3.4 内容类型五分类分析（v2.0新增）
    content_types_text = chapter_map.get('ch3_content_types', '')
    if content_types_text:
        _add_heading(doc, '3.6  内容类型五分类分析', level=2)
        add_md_content_to_docx(doc, content_types_text)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # CHARTS SECTION（图表放在 ch3 与 ch4 之间）
    # ═══════════════════════════════════════════════
    c_dir = _get_charts_dir(project_config)
    chart_files = _collect_chart_files(c_dir)
    if chart_files:
        add_charts_to_docx(doc, chart_files)
    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # CH4: 本品深度分析
    # ═══════════════════════════════════════════════
    focus = project_config.focus_brand or "本品"
    _add_heading(doc, f'第4章  本品深度分析 —— {focus}', level=1)
    ch4_text = chapter_map.get('ch4', '')
    if ch4_text:
        add_md_content_to_docx(doc, ch4_text)
    else:
        _add_body(doc, '[内容待补充]')
    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # CH5: 差距对比
    # ═══════════════════════════════════════════════
    _add_heading(doc, '第5章  本竞品差距对比', level=1)
    ch5_text = chapter_map.get('ch5', '')
    if ch5_text:
        add_md_content_to_docx(doc, ch5_text)
    else:
        _add_body(doc, '[内容待补充]')
    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # CH6: 策略建议
    # ═══════════════════════════════════════════════
    _add_heading(doc, '第6章  咨询切入点与策略建议', level=1)
    ch6_text = chapter_map.get('ch6', '')
    if ch6_text:
        add_md_content_to_docx(doc, ch6_text)
    else:
        _add_body(doc, '[内容待补充]')
    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # CH7: 睡眠消费洞察（如适用）
    # ═══════════════════════════════════════════════
    ch7_text = chapter_map.get('ch7', '')
    if ch7_text and len(ch7_text.strip()) > 500:
        _add_heading(doc, '第7章  2026最新睡眠消费洞察', level=1)
        add_md_content_to_docx(doc, ch7_text)
        doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 附录 A：创品策略
    # ═══════════════════════════════════════════════
    innovation_text = chapter_map.get('appendix_innovation', '')
    if innovation_text:
        focus_brand = project_config.focus_brand or "品牌"
        _add_heading(doc, f'附录A：创品策略 —— {focus_brand}品牌创新方向', level=1)
        add_md_content_to_docx(doc, innovation_text)
        doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 附录 B：创始人研究
    # ═══════════════════════════════════════════════
    founder_text = chapter_map.get('appendix_founder', '')
    if founder_text:
        _add_heading(doc, '附录B：创始人研究', level=1)
        add_md_content_to_docx(doc, founder_text)

    # ═══════════════════════════════════════════════
    # 保存
    # ═══════════════════════════════════════════════
    doc.save(str(docx_path))

    verify_output_file(docx_path, "docx_assembly")
    step_success("docx_assembly", [str(docx_path)])
    print(f"  ✅ DOCX 生成完成: {docx_path.name} ({docx_path.stat().st_size / 1024:.0f} KB)")
    return docx_path


# ── 章节内容加载（build_chapter_map） ──────────────────────

def build_chapter_map(project_config: ProjectConfig) -> Dict[str, str]:
    """
    直接从已知文件名加载各章节内容。
    不 glob，不遍历子目录。返回 {key: markdown_text} 字典。
    排除 pre_research.md。
    """
    c_dir = _get_content_dir(project_config)
    ch3_dir = c_dir / "ch3_competitive"
    ch4_dir = c_dir / "ch4_deep"

    result = {}

    def _read(p: Path) -> str:
        if p.exists():
            return p.read_text(encoding="utf-8")
        return ""

    # brand_overview
    result['brand_overview'] = _read(c_dir / "brand_overview.md")
    # v2.1 基础扫描三章：产品矩阵 + 渠道供应链
    result['product_matrix'] = _read(c_dir / "product_matrix.md")
    result['channel_supply_chain'] = _read(c_dir / "channel_supply_chain.md")

    # ch1 — if not exists, generate default content
    ch1_path = c_dir / "ch1_findings.md"
    if ch1_path.exists():
        result['ch1'] = _read(ch1_path)
    else:
        # 生成默认内容
        result['ch1'] = (
            "# 核心发现与咨询窗口\n\n"
            "## 核心发现\n\n"
            "1. [待补充]\n\n"
            "## 咨询窗口\n\n"
            "[待补充]\n\n"
        )

    # ch2 — 行业格局 + 产业链地图
    ch2_base = _read(c_dir / "ch2_industry.md")
    ch2_chain = _read(c_dir / "ch2_chain_map.md")
    if not ch2_chain:
        # pipeline Step 5.1 生成名为 industry_chain_map.md
        ch2_chain = _read(c_dir / "industry_chain_map.md")
    if ch2_base and ch2_chain:
        result['ch2'] = ch2_base + "\n\n" + ch2_chain
    elif ch2_base:
        result['ch2'] = ch2_base
    elif ch2_chain:
        result['ch2'] = ch2_chain
    else:
        result['ch2'] = _read(c_dir / "ch2_industry_analysis.md")

    # ch3 深度品牌 — 合并所有 deep_*.md
    deep_parts = []
    if ch3_dir.exists():
        for f in sorted(ch3_dir.glob("deep_*.md")):
            if f.name.endswith("_prompt.md"):
                continue
            deep_parts.append(f.read_text(encoding="utf-8"))
    result['ch3_deep_brands'] = "\n\n".join(deep_parts)

    # ch3 汇总品牌
    summary_path = ch3_dir / "summary_brands.md"
    result['ch3_summary'] = _read(summary_path)

    # ch3 竞争模式
    patterns_path = ch3_dir / "competition_patterns.md"
    result['ch3_patterns'] = _read(patterns_path)

    # ch3 内容类型五分类（v2.0新增）
    content_types_path = c_dir / "ch3_content_types.md"
    result['ch3_content_types'] = _read(content_types_path)
    if not result['ch3_content_types']:
        # 也检查子目录命名
        result['ch3_content_types'] = _read(ch3_dir / "content_types.md")
    if not result['ch3_content_types']:
        # pipeline Step 7.1 生成名为 content_type_analysis.md
        result['ch3_content_types'] = _read(c_dir / "content_type_analysis.md")

    # ch4 — 深度品牌分析
    focus = project_config.focus_brand
    if focus:
        ch4_path = ch4_dir / f"{focus}_deep.md"
        if not ch4_path.exists():
            ch4_path = c_dir / "ch4_deep_analysis.md"
        result['ch4'] = _read(ch4_path)

    # ch5
    result['ch5'] = _read(c_dir / "ch5_gap.md")
    if not result['ch5']:
        result['ch5'] = _read(c_dir / "ch5_strategy.md")

    # ch6 — 策略建议 + 机会地图
    ch6_base = _read(c_dir / "ch6_strategy.md")
    ch6_opp = _read(c_dir / "ch6_opportunity_map.md")
    if ch6_base and ch6_opp:
        result['ch6'] = ch6_base + "\n\n" + ch6_opp
    elif ch6_base:
        result['ch6'] = ch6_base
    elif ch6_opp:
        result['ch6'] = ch6_opp
    else:
        result['ch6'] = _read(c_dir / "ch6_recommendations.md")

    # ch7 — 睡眠消费洞察（按需）
    result['ch7'] = _read(c_dir / "ch7_sleep_insights.md")

    # 附录 — 创品策略
    result['appendix_innovation'] = _read(c_dir / "ch10_innovation.md")
    if not result['appendix_innovation']:
        result['appendix_innovation'] = _read(c_dir / "innovation_strategy.md")

    # 附录 — 创始人研究
    result['appendix_founder'] = _read(c_dir / "ch11_founder.md")
    if not result['appendix_founder']:
        result['appendix_founder'] = _read(c_dir / "founder_research.md")

    return result


# ── Markdown → DOCX 渲染 ──────────────────────────────────

def add_md_content_to_docx(doc, md_text: str):
    """
    将 Markdown 文本渲染到 python-docx Document 中。
    规则：
    - 跳过（忽略）H1 标题（章节标题由外部硬编码提供）
    - H2/H3/H4 → add_heading (level 2/3/4)
    - **粗体** → Run.bold=True
    - | 分隔表格 → add_table + Table Grid 样式
    - - 开头段落 → List Bullet 样式段落
    - --- 分隔线 → 装饰段落
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    lines = md_text.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 空行
        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # 跳过顶部 H1（章节标题由外部提供）
        if line.startswith('# ') and not line.startswith('## '):
            i += 1
            continue

        # H2/H3/H4
        if line.startswith('#### '):
            _add_heading(doc, line[5:].replace('**', ''), level=4)
            i += 1
            continue
        if line.startswith('### '):
            _add_heading(doc, line[4:].replace('**', ''), level=3)
            i += 1
            continue
        if line.startswith('## '):
            _add_heading(doc, line[3:].replace('**', ''), level=2)
            i += 1
            continue

        # 分隔线
        if re.match(r'^-{3,}$', line):
            _add_body(doc, '─' * 50)
            i += 1
            continue

        # Markdown 表格：| col1 | col2 |
        if line.startswith('|') and line.endswith('|') and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if next_line.startswith('|') and '---' in next_line.replace(' ', ''):
                # 收集表格所有行
                table_rows = []
                while i < len(lines):
                    l = lines[i].strip()
                    if l.startswith('|') and l.endswith('|'):
                        cells = [c.strip() for c in l.split('|')[1:-1]]
                        table_rows.append(cells)
                        i += 1
                    else:
                        break

                if len(table_rows) >= 2:
                    # 跳过分隔行（|---|---|）
                    data_rows = table_rows[:1] + table_rows[2:] if len(table_rows) > 1 else table_rows
                    data_rows = [row for row in data_rows if any(c for c in row)]

                    if data_rows:
                        ncols = max(len(r) for r in data_rows)
                        table = doc.add_table(rows=len(data_rows), cols=ncols, style='Table Grid')
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        for ri, row in enumerate(data_rows):
                            for ci, cell_text in enumerate(row):
                                if ci < ncols:
                                    cell = table.cell(ri, ci)
                                    cell.text = cell_text
                                    # 表头加粗
                                    if ri == 0:
                                        for p in cell.paragraphs:
                                            for r in p.runs:
                                                r.bold = True
                                                r.font.size = Pt(9)
                                    else:
                                        for p in cell.paragraphs:
                                            for r in p.runs:
                                                r.font.size = Pt(9)
                        doc.add_paragraph()
                continue

        # 列表项
        if re.match(r'^[\-\*]\s+', line):
            text = re.sub(r'^[\-\*]\s+', '', line)
            _add_body(doc, text)
            i += 1
            continue

        # 数字列表
        if re.match(r'^\d+[\.\)]\s+', line):
            text = re.sub(r'^\d+[\.\)]\s+', '', line)
            _add_body(doc, text)
            i += 1
            continue

        # 普通段落（保留行内 **bold**）
        _add_body(doc, line)
        i += 1


# ── 图表版块 ──────────────────────────────────────────────

def add_charts_to_docx(doc, chart_files: List[Tuple[str, Path]]):
    """
    添加图表版块到文档。
    每张图表 = 图标题(H3) + 图片 + 数据来源。
    使用 add_picture 嵌入 PNG。
    """
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _add_heading(doc, '数据可视化：核心竞品对比', level=2)
    _add_body(doc, '以下图表基于天猫/京东旗舰店搜索实测数据生成。')
    doc.add_paragraph()

    for i, (title, img_path) in enumerate(chart_files):
        if not img_path.exists():
            print(f"  ⚠ 图表文件不存在，跳过: {img_path.name}")
            continue

        _add_heading(doc, f'图{i+1}：{title}', level=3)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))

        # 数据来源标注
        source_note = f"数据来源：{_chart_source_hint(img_path.name)}"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(source_note)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()


# ── 辅助函数 ──────────────────────────────────────────────

def _add_heading(doc, text: str, level: int = 1):
    """添加标题，去除可能的 ** 标记。"""
    h = doc.add_heading(text.replace('**', ''), level=level)
    return h


def _add_body(doc, text: str):
    """
    添加正文段落。
    处理行内 **bold** 标记 → Run.bold=True。
    """
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def _validate_docx_naming(filename: str) -> None:
    """验证 docx 文件名是否符合版本号规范。"""
    pattern = r'^.+?-.+-V\d+(\.\d+)?-\d{8}\.docx$'
    if not re.match(pattern, filename):
        print(f"  ⚠ 版本号命名警告：当前文件名「{filename}」不符合规范格式「品牌中文名-行业-V数字.数字-日期.docx」")
        print(f"     整数升版=结构改动（增删章节/补全模块），小数升版=文字改动（修bug/改措辞/调数据）")
    else:
        print(f"  ✓ 文件名版本号格式合规: {filename}")


def _build_toc(schema: ReportSchema, project_config: ProjectConfig) -> List[str]:
    """构建目录列表。"""
    items = [
        '品牌概览',
        '产品矩阵',
        '渠道与供应链',
        '第1章  核心发现与咨询窗口',
        '第2章  行业格局与竞品总矩阵',
        '第3章  竞品多维度扫描',
        '    3.1  深度品牌分析',
        '    3.2  汇总品牌速览',
        '    3.3  竞争模式归纳',
        '    3.6  内容类型五分类分析',
        '    数据可视化：核心竞品对比',
    ]
    focus = project_config.focus_brand or "本品"
    items.append(f'第4章  本品深度分析 —— {focus}')
    items.extend([
        '第5章  本竞品差距对比',
        '第6章  咨询切入点与策略建议',
    ])
    # ch7 如果有实质内容再加入
    ch7_path = content_dir(project_config) / "ch7_sleep_insights.md"
    if ch7_path.exists() and ch7_path.stat().st_size > 1000:
        items.append('第7章  2026最新睡眠消费洞察')
    items.extend([
        '附录A  创品策略',
        '附录B  创始人研究',
    ])
    return items


def _collect_chart_files(c_dir: Path) -> List[Tuple[str, Path]]:
    """
    收集图表文件。按预期顺序返回 [(title, path), ...]。
    优先使用 PNG，回退到 HTML。
    """
    expected = [
        ("天猫旗舰店爆款销售对比", "chart_brand_comparison_1"),
        ("京东自营爆款评价数对比", "chart_brand_comparison_2"),
        ("各品牌核心产品单件价对比", "chart_brand_comparison_3"),
        ("各品牌回头客率对比", "chart_brand_comparison_4"),
    ]

    results = []
    for title, chart_id in expected:
        png = c_dir / f"{chart_id}.png"
        html = c_dir / f"{chart_id}.html"
        if png.exists():
            results.append((title, png))
        elif html.exists():
            results.append((title, html))
    return results


def _chart_source_hint(filename: str) -> str:
    """根据文件名返回数据来源提示。"""
    hints = {
        "comparison_1": "天猫搜索实测",
        "comparison_2": "京东搜索实测",
        "comparison_3": "天猫/京东旗舰店定价",
        "comparison_4": "天猫旗舰店回头客标签",
    }
    for key, hint in hints.items():
        if key in filename:
            return hint
    return "电商平台"


def _get_content_dir(project_config) -> Path:
    """
    获取 content 目录。
    优先级：
    1. project_config.output_dir 指定的项目隔离目录
    2. output/ 下最近的项目名匹配目录
    3. 通用 content_dir()
    """
    project_name = project_config.project_name.replace(" ", "_").replace("品牌竞品研究", "")
    # 1. 项目隔离目录（配置指定）
    output_dir_raw = project_config._raw.get("output_dir", "")
    if output_dir_raw:
        project_content = BASE_DIR / output_dir_raw / "content"
        if project_content.exists():
            md_files = list(project_content.glob("*.md")) + list(project_content.glob("**/*.md"))
            if md_files:
                return project_content
    # 2. 搜索 output/ 下的最近匹配目录
    output_root = content_dir().parent
    matches = []
    for d in output_root.iterdir():
        if d.is_dir() and project_name.strip() in d.name:
            cdir = d / "content"
            if cdir.exists():
                md_files = list(cdir.glob("*.md")) + list(cdir.glob("**/*.md"))
                if md_files:
                    matches.append((d.stat().st_mtime, cdir))
    if matches:
        # 优先选择有实质内容的目录（>5000字符视为有真实内容），
        # 避免选中空占位目录；内容质量相同时再按 mtime 最新优先。
        def _has_content(cdir):
            md_files = list(cdir.glob('*.md')) + list(cdir.glob('**/*.md'))
            md_files = [f for f in md_files if not f.name.endswith('_prompt.md')]
            if not md_files:
                return (0, 0)
            total = sum(f.stat().st_size for f in md_files if f.stat().st_size > 200)
            return (1 if total > 5000 else 0, total)  # 5000 chars = 有真实内容
        matches.sort(key=lambda m: (*_has_content(m[1]), m[0]), reverse=True)
        return matches[0][1]
    # 3. 回退：通用目录
    return content_dir()


def _get_charts_dir(project_config) -> Path:
    """
    获取 charts 目录。
    注意：pipeline 各步骤写入 output/charts/（通用目录），
    因此也从通用目录读取。
    """
    return charts_dir()


def ensure_output_dir() -> Path:
    """确保 output/reports/ 目录存在（支持 project_config 隔离）。"""
    return output_dir("reports")


# ── 兼容层 ────────────────────────────────────────────────

def load_chapter_content(schema: ReportSchema, project_config: ProjectConfig) -> Dict[str, str]:
    """
    兼容旧接口：内部委托给 build_chapter_map。
    """
    return build_chapter_map(project_config)


def embed_charts_in_docx(*args, **kwargs):
    """
    废弃。其功能已合并到 add_charts_to_docx。
    调用时将打印警告但不抛出异常。
    """
    print("  ⚠ embed_charts_in_docx 已废弃（功能合并到 add_charts_to_docx），调用无效果")


def get_chart_files(project_config: ProjectConfig) -> List[Tuple[str, Path]]:
    """兼容旧接口。"""
    return _collect_chart_files(_get_charts_dir(project_config))
