#!/usr/bin/env python3
"""Assemble the libernovo complete report docx from markdown chapters + chart PNGs."""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUTPUT_DIR = Path("/Users/yifansmacmini/.openclaw/workspace/supermind/output")
CONTENT_DIR = OUTPUT_DIR / "content"
CHARTS_DIR = OUTPUT_DIR / "charts"

# Chapter files in order
CHAPTER_FILES = [
    ("ch1_core_findings.md", "第一章 核心发现与咨询窗口"),
    ("ch2_industry.md", "第二章 行业格局与竞品总矩阵"),
    ("ch3_competitive_01_sihoo.md", None),
    ("ch3_competitive_02_yongyi.md", None),
    ("ch3_competitive_03_hbada.md", None),
    ("ch3_competitive_04_ergonor.md", None),
    ("ch3_competitive_05_summary.md", None),
    ("ch4_deep_analysis.md", "第四章 本品深度分析：清闲/libernovo"),
    ("ch5_gap.md", "第五章 本竞品差距对比"),
    ("ch6_recommendations.md", "第六章 咨询切入点与策略建议"),
]

# Charts to embed (inserted before ch3_competitive_05_summary)
CHARTS = [
    ("chart_new_1_tmall.png", "图1：天猫旗舰店爆款销量对比（数据来源：天猫旗舰店，采集日期2026-07-15）"),
    ("chart_new_2_jd.png", "图2：京东自营爆款销量对比（数据来源：京东自营，采集日期2026-07-15；※标注为天猫数据参考）"),
    ("chart_new_3_price.png", "图3：各品牌爆款产品价格对比（数据来源：天猫旗舰店+京东自营，采集日期2026-07-15）"),
    ("chart_new_4_return.png", "图4：各品牌回头客率对比（数据来源：天猫旗舰店详情页，采集日期2026-07-15；※标注为综合推断）"),
]

def parse_markdown_to_blocks(text):
    """Parse markdown into blocks: (type, level, text)"""
    lines = text.strip().split('\n')
    blocks = []
    current_text = []
    current_type = 'paragraph'
    current_level = 0
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_text:
                blocks.append((current_type, current_level, '\n'.join(current_text)))
                current_text = []
            current_type = 'paragraph'
            current_level = 0
            continue
        
        # Check for table rows (| ... |)
        if stripped.startswith('|') and stripped.endswith('|'):
            if current_type != 'table':
                if current_text:
                    blocks.append((current_type, current_level, '\n'.join(current_text)))
                    current_text = []
                current_type = 'table'
                current_level = 0
            current_text.append(stripped)
            continue
        
        # Check for headings
        h_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if h_match:
            if current_text:
                blocks.append((current_type, current_level, '\n'.join(current_text)))
                current_text = []
            level = len(h_match.group(1))
            blocks.append(('heading', level, h_match.group(2)))
            current_type = 'heading'
            current_level = level
            continue
        
        # Regular paragraph text
        if current_type == 'table':
            if current_text:
                blocks.append((current_type, current_level, '\n'.join(current_text)))
                current_text = []
            current_type = 'paragraph'
            current_level = 0
        
        current_text.append(stripped)
    
    if current_text:
        blocks.append((current_type, current_level, '\n'.join(current_text)))
    
    return blocks

def add_paragraph_with_bold_leading(doc, text):
    """Add a paragraph, handling **bold** markers and | table separators."""
    # Skip separator lines in tables (|---|---|)
    if re.match(r'^\|[\s\-:|]+\|$', text):
        return
    
    # Check if this is a table row
    if text.startswith('|') and text.endswith('|'):
        cells = [c.strip() for c in text.split('|')[1:-1]]
        # Skip if it looks like a table
        if len(cells) >= 3:
            # For simplicity, render table rows as formatted paragraphs
            p = doc.add_paragraph()
            for i, cell in enumerate(cells):
                if i > 0:
                    p.add_run('  |  ')
                run = p.add_run(cell)
                run.font.size = Pt(10)
            return
    
    p = doc.add_paragraph()
    # Simple bold handling
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    
    # Set font
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'

def build_docx():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # --- Title page ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('清闲/libernovo 人体工学椅行业\n品牌研究报告')
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = '微软雅黑'
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run('基于五维分析模型的完整品牌竞争扫描')
    sub_run.font.size = Pt(14)
    sub_run.font.name = '微软雅黑'
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run('2026年7月  |  燃创品牌咨询')
    date_run.font.size = Pt(12)
    date_run.font.name = '微软雅黑'
    
    doc.add_page_break()
    
    # --- Process chapter files ---
    for filename, force_h1_title in CHAPTER_FILES:
        filepath = CONTENT_DIR / filename
        if not filepath.exists():
            print(f"WARNING: {filepath} not found, skipping")
            continue
        
        content = filepath.read_text(encoding='utf-8')
        blocks = parse_markdown_to_blocks(content)
        
        # Determine the H1 title for this chapter
        if force_h1_title:
            # Add chapter title as H1
            h = doc.add_heading(force_h1_title, level=1)
        else:
            # The first heading in the file might be H2 (for ch3 sub-sections)
            # We already have the main ch3 heading from the playbook,
            # so these files should just flow as content
            # But for the first file (sihoo), we need to add a section heading
            if "competitive" in filename:
                # These are sub-sections of ch3
                pass
        
        for block_type, block_level, block_text in blocks:
            if block_type == 'heading':
                if force_h1_title and block_level == 1:
                    # Skip H1 in sub-chapter files if we already added the main heading
                    # But keep it if it's a real section heading (like ch4, ch5, ch6)
                    if 'competitive' in filename:
                        # Downgrade H1 to H2 for ch3 sub-sections
                        doc.add_heading(block_text, level=2)
                    else:
                        doc.add_heading(block_text, level=block_level)
                else:
                    doc.add_heading(block_text, level=block_level)
            elif block_type == 'table':
                # Render table
                rows = []
                for row_text in block_text.split('\n'):
                    if re.match(r'^\|[\s\-:|]+\|$', row_text):
                        continue  # skip separator
                    cells = [c.strip() for c in row_text.split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    num_cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    for i, row in enumerate(rows):
                        for j, cell_text in enumerate(row):
                            if j < num_cols:
                                cell = table.cell(i, j)
                                cell.text = cell_text
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)
                    doc.add_paragraph()  # spacing after table
            else:
                # Regular paragraph
                for para_text in block_text.split('\n'):
                    para_text = para_text.strip()
                    if para_text:
                        add_paragraph_with_bold_leading(doc, para_text)
        
        # Insert charts before the summary/competition patterns section (ch3_competitive_04_ergonor is the last brand)
        if filename == "ch3_competitive_04_ergonor.md":
            doc.add_paragraph()
            chart_header = doc.add_heading('跨品牌数据总览', level=2)
            
            for chart_file, chart_caption in CHARTS:
                chart_path = CHARTS_DIR / chart_file
                if chart_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(chart_path), width=Inches(5.5))
                    
                    caption_p = doc.add_paragraph()
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = caption_p.add_run(chart_caption)
                    cap_run.font.size = Pt(9)
                    cap_run.italic = True
                    doc.add_paragraph()
                else:
                    print(f"WARNING: Chart {chart_path} not found")
        
        # Add page break between major chapters
        if filename in ("ch1_core_findings.md", "ch2_industry.md", "ch3_competitive_05_summary.md",
                        "ch4_deep_analysis.md", "ch5_gap.md"):
            doc.add_page_break()
    
    # --- Save ---
    output_path = OUTPUT_DIR / "libernovo_COMPLETE_report.docx"
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")
    
    # Get file size
    size_kb = output_path.stat().st_size / 1024
    print(f"File size: {size_kb:.1f} KB")
    
    return output_path, size_kb

if __name__ == "__main__":
    build_docx()
