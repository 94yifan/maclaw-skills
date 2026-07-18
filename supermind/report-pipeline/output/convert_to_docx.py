#!/usr/bin/env python3
"""Convert markdown report to DOCX, preserving **bold** formatting."""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_markdown_line(line):
    """Parse a markdown line into runs: (text, bold) tuples."""
    if not line:
        return [("", False)]
    
    # Split by ** pairs
    parts = re.split(r'(\*\*.*?\*\*)', line)
    runs = []
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            runs.append((part[2:-2], True))
        else:
            runs.append((part, False))
    return runs

def add_formatted_paragraph(doc, text, style_name=None, font_size=None, bold_all=False, 
                           alignment=None, space_after=None, space_before=None, color=None):
    """Add a paragraph with mixed bold/normal formatting."""
    p = doc.add_paragraph()
    
    if style_name:
        p.style = doc.styles[style_name]
    
    if alignment is not None:
        p.alignment = alignment
    
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    
    runs_data = parse_markdown_line(text)
    for run_text, is_bold in runs_data:
        if not run_text:
            continue
        run = p.add_run(run_text)
        if bold_all or is_bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
        # Set font for Chinese compatibility
        run.font.name = '微软雅黑'
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = r.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    return p

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle horizontal rules (---)
        if line.strip() == '---':
            # Add a thin horizontal line paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): '999999',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Handle blockquotes
        if line.startswith('>'):
            blockquote_lines = []
            while i < len(lines) and lines[i].rstrip().startswith('>'):
                blockquote_lines.append(lines[i].rstrip()[1:].strip())
                i += 1
            for bq_line in blockquote_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(bq_line)
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.italic = True
                run.font.name = '微软雅黑'
            continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|') and line.strip().endswith('|'):
            # Collect table lines
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            
            # Filter out separator rows (|---|---|)
            data_rows = []
            for tl in table_lines:
                if not re.match(r'^\|[\s\-:|]+\|$', tl):
                    data_rows.append(tl)
            
            if len(data_rows) >= 1:
                # Determine columns
                cols = len([c for c in data_rows[0].split('|') if c.strip()])
                table = doc.add_table(rows=len(data_rows), cols=cols)
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_line in enumerate(data_rows):
                    cells = [c.strip() for c in row_line.split('|')][1:-1]  # Remove leading/trailing empties
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # Clear default paragraph
                            cell.paragraphs[0].clear()
                            # Parse bold in cell
                            runs_data = parse_markdown_line(cell_text)
                            for run_text, is_bold in runs_data:
                                if not run_text:
                                    continue
                                run = cell.paragraphs[0].add_run(run_text)
                                run.font.size = Pt(9)
                                run.font.name = '微软雅黑'
                                if is_bold or row_idx == 0:
                                    run.bold = True
                
                # Add spacing after table
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        
        # Handle headings
        if line.startswith('# '):
            add_formatted_paragraph(doc, line[2:], bold_all=True, font_size=20,
                                   space_before=18, space_after=10)
            i += 1
            continue
        elif line.startswith('## '):
            add_formatted_paragraph(doc, line[3:], bold_all=True, font_size=16,
                                   space_before=16, space_after=8)
            i += 1
            continue
        elif line.startswith('### '):
            add_formatted_paragraph(doc, line[4:], bold_all=True, font_size=14,
                                   space_before=14, space_after=6)
            i += 1
            continue
        elif line.startswith('#### '):
            add_formatted_paragraph(doc, line[5:], bold_all=True, font_size=12,
                                   space_before=10, space_after=4)
            i += 1
            continue
        elif line.startswith('##### '):
            add_formatted_paragraph(doc, line[6:], bold_all=True, font_size=11,
                                   space_before=8, space_after=3)
            i += 1
            continue
        
        # Handle unordered lists
        if re.match(r'^[\-\*]\s+', line):
            # Collect all list items
            list_items = []
            while i < len(lines):
                cl = lines[i].rstrip()
                if re.match(r'^[\-\*]\s+', cl):
                    list_items.append(cl)
                    i += 1
                elif cl.strip() and not cl.startswith('#') and not cl.startswith('|') and not cl.startswith('>'):
                    # Continuation line (indented continuation)
                    if list_items:
                        list_items[-1] += ' ' + cl.strip()
                    else:
                        break
                    i += 1
                else:
                    break
            
            for item in list_items:
                # Remove leading - or * and space
                item_text = re.sub(r'^[\-\*]\s+', '', item)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.space_after = Pt(2)
                
                # Add bullet character
                bullet_run = p.add_run('• ')
                bullet_run.font.size = Pt(11)
                
                runs_data = parse_markdown_line(item_text)
                for run_text, is_bold in runs_data:
                    if not run_text:
                        continue
                    run = p.add_run(run_text)
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
                    if is_bold:
                        run.bold = True
            continue
        
        # Handle ordered lists
        if re.match(r'^\d+\.\s+', line):
            list_items = []
            while i < len(lines):
                cl = lines[i].rstrip()
                if re.match(r'^\d+\.\s+', cl):
                    list_items.append(cl)
                    i += 1
                elif cl.strip() and not cl.startswith('#') and not cl.startswith('|'):
                    if list_items:
                        list_items[-1] += ' ' + cl.strip()
                    else:
                        break
                    i += 1
                else:
                    break
            
            for idx, item in enumerate(list_items):
                item_text = re.sub(r'^\d+\.\s+', '', item)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.space_after = Pt(2)
                
                num_run = p.add_run(f'{idx+1}. ')
                num_run.font.size = Pt(11)
                
                runs_data = parse_markdown_line(item_text)
                for run_text, is_bold in runs_data:
                    if not run_text:
                        continue
                    run = p.add_run(run_text)
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
                    if is_bold:
                        run.bold = True
            continue
        
        # Handle empty lines
        if not line.strip():
            i += 1
            continue
        
        # Regular paragraph
        add_formatted_paragraph(doc, line, font_size=11, space_after=6)
        i += 1
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f'DOCX saved to {docx_path}')
    print(f'File size: {os.path.getsize(docx_path)} bytes')

if __name__ == '__main__':
    md_path = '/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/content/report_final.md'
    docx_path = '/Users/yifansmacmini/.openclaw/workspace/supermind/report-pipeline/output/reports/liumang_final.docx'
    convert_md_to_docx(md_path, docx_path)
